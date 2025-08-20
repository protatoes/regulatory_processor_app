#!/usr/bin/env python
# coding: utf-8

# In[1]:


import os
import re
from pathlib import Path
from typing import Optional, Dict, Tuple, List
from docx import Document
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx import Document
from docx.text.paragraph import Hyperlink
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd
import logging


# In[17]:


def can_reuse_replacement_text(mapping_row: pd.Series) -> bool:
    """
    Check if SmPC and PL replacement texts are the same.
    If they are, we can reuse the same replacement components.
    
    Parameters:
    -----------
    mapping_row : pd.Series
        Row from mapping table
        
    Returns:
    --------
    bool
        True if texts are identical and can be reused
    """
    smpc_text = str(mapping_row.get('National reporting system SmPC', '')).strip()
    pl_text = str(mapping_row.get('National reporting system PL', '')).strip()
    
    # Both must exist and be identical
    reusable = (smpc_text == pl_text and 
                smpc_text and 
                smpc_text.lower() != 'nan' and
                pl_text.lower() != 'nan')
    
    if reusable:
        print("✅ SmPC and PL replacement texts are identical - can reuse components")
    else:
        print("📝 SmPC and PL have different replacement texts - will generate separately")
    
    return reusable


def build_replacement_from_lines(mapping_row: pd.Series, section_type: str, 
                                country_delimiter: str = ',') -> List[Dict]:
    """
    Build replacement components from Line 1-N columns using the new structured approach.
    
    Parameters:
    -----------
    mapping_row : pd.Series
        Row from mapping table
    section_type : str
        Either 'SmPC' or 'PL'
        
    Returns:
    --------
    List[Dict]
        List of formatted text components ready for insertion
    """
def build_replacement_from_lines(mapping_row: pd.Series, section_type: str, 
                                country_delimiter: str = ',') -> List[Dict]:
    """
    Build replacement components from Line 1-N columns using positional country matching.
    
    Parameters:
    -----------
    mapping_row : pd.Series
        Row from mapping table
    section_type : str
        Either 'SmPC' or 'PL'
    country_delimiter : str
        Delimiter used to separate countries and their content (default: ',')
        
    Returns:
    --------
    List[Dict]
        List of formatted text components ready for insertion
    """
    print(f"🔨 Building replacement text from lines for {section_type}...")
    print(f"   Using delimiter: '{country_delimiter}'")
    
    components = []
    
    # Get all Line columns for this section (Line 1, Line 2, etc.)
    # Handle different column naming patterns
    if section_type == 'SmPC':
        line_columns = [col for col in mapping_row.index 
                       if col.startswith('Line ') and ('SmPC' in col or 'SmpC' in col)]  # Handle typo in Line 8
    else:  # PL
        line_columns = [col for col in mapping_row.index 
                       if col.startswith('Line ') and 'PL' in col]
    
    print(f"   - Looking for columns with pattern: Line X - {section_type}")
    print(f"   - Found {len(line_columns)} matching columns: {line_columns}")
    
    if not line_columns:
        print(f"⚠️  No line columns found for {section_type}")
        
        # Debug: Show what columns are available
        available_cols = [col for col in mapping_row.index if 'Line ' in col]
        print(f"   Available Line columns: {available_cols}")
        
        # Fallback to single-cell approach for PL if no line columns exist
        if section_type == 'PL':
            print("   → Attempting fallback to single-cell PL text...")
            single_cell_text = str(mapping_row.get('National reporting system PL', '')).strip()
            if single_cell_text and single_cell_text.lower() != 'nan':
                print(f"   → Found fallback text: '{single_cell_text[:50]}{'...' if len(single_cell_text) > 50 else ''}'")
                # Parse the single cell text (use existing logic or simple approach)
                components.append({
                    'text': '\n' + single_cell_text,
                    'bold': False,
                    'is_hyperlink': False
                })
                return components
        
        return components
    
    # Get hyperlinks for this section
    hyperlinks_col = f'Hyperlinks {section_type}'
    hyperlinks_str = str(mapping_row.get(hyperlinks_col, '')).strip()
    hyperlinks = [h.strip() for h in hyperlinks_str.split(',') if h.strip() and h.strip().lower() != 'nan']
    
    print(f"   - Found {len(line_columns)} line columns")
    print(f"   - Found {len(hyperlinks)} hyperlinks: {hyperlinks}")
    
    # Sort line columns by number (Line 1, Line 2, etc.)
    def extract_line_number(col_name):
        match = re.search(r'Line (\d+)', col_name)
        return int(match.group(1)) if match else 999
    
    sorted_columns = sorted(line_columns, key=extract_line_number)
    
    # Step 1: Get countries from Line 1
    line_1_col = None
    for col in sorted_columns:
        if extract_line_number(col) == 1:
            line_1_col = col
            break
    
    if not line_1_col:
        print("❌ No Line 1 found - cannot determine countries")
        return components
    
    line_1_text = str(mapping_row.get(line_1_col, '')).strip()
    if not line_1_text or line_1_text.lower() == 'nan':
        print("❌ Line 1 is empty - cannot determine countries")
        return components
    
    # Parse countries from Line 1
    countries = [c.strip() for c in line_1_text.split(country_delimiter) if c.strip()]
    print(f"   - Found {len(countries)} countries: {countries}")
    
    if not countries:
        print("❌ No countries found in Line 1")
        return components
    
    # Step 2: Collect content for each line (excluding Line 1)
    line_content = {}  # {line_number: [parts_for_each_country]}
    
    for line_col in sorted_columns:
        line_number = extract_line_number(line_col)
        if line_number == 1:  # Skip Line 1 (countries)
            continue
            
        line_text = str(mapping_row.get(line_col, '')).strip()
        print(f"   - Processing Line {line_number}: '{line_text[:50]}{'...' if len(line_text) > 50 else ''}'")
        
        # Skip completely empty cells or NaN
        if not line_text or line_text.lower() == 'nan':
            print(f"     → Skipping empty line {line_number}")
            continue
            
        # Handle whitespace-only cells as empty lines
        if line_text.isspace():
            print(f"     → Line {line_number} is whitespace-only (will add empty line)")
            # Add empty parts for all countries
            line_content[line_number] = ['\n'] * len(countries)
            continue
        
        # Split content by delimiter for each country
        parts = [p.strip() for p in line_text.split(country_delimiter)]
        
        # Pad with empty strings if fewer parts than countries
        while len(parts) < len(countries):
            parts.append('')
        
        line_content[line_number] = parts
        print(f"     → Split into {len(parts)} parts: {[p[:30] + '...' if len(p) > 30 else p for p in parts]}")
    
    # Step 3: Build separate blocks for each country
    print(f"\n🏗️  Building {len(countries)} country blocks...")
    
    # Add initial line break before first country block
    components.append({'text': '\n\n', 'bold': False, 'is_hyperlink': False})
    
    for country_idx, country_name in enumerate(countries):
        print(f"   - Building block for {country_name} (position {country_idx})")
        
        # Add country name (bold)
        components.append({
            'text': country_name,
            'bold': True,
            'is_hyperlink': False
        })
        components.append({'text': '\n', 'bold': False, 'is_hyperlink': False})
        
        # Add content lines for this country
        for line_number in sorted(line_content.keys()):
            parts = line_content[line_number]
            
            # Get the part for this country (if it exists)
            if country_idx < len(parts) and parts[country_idx]:
                content_text = parts[country_idx]
                
                # Handle whitespace-only content as line breaks
                if content_text == '\n' or content_text.isspace():
                    components.append({'text': '\n', 'bold': False, 'is_hyperlink': False})
                    continue
                
                print(f"     → Line {line_number}: '{content_text[:40]}{'...' if len(content_text) > 40 else ''}'")
                
                # Check if this content contains any URLs that should be hyperlinked
                url_processed = False
                for url in hyperlinks:
                    if url in content_text:
                        print(f"       → Found hyperlink '{url}' in content")
                        # Split the content around the URL
                        parts_split = content_text.split(url, 1)  # Split only on first occurrence
                        
                        # Add text before URL
                        if parts_split[0]:
                            components.append({
                                'text': parts_split[0],
                                'bold': False,
                                'is_hyperlink': False
                            })
                        
                        # Add the URL as hyperlink
                        components.append({
                            'text': url,
                            'bold': False,
                            'is_hyperlink': True
                        })
                        
                        # Add text after URL
                        if parts_split[1]:
                            components.append({
                                'text': parts_split[1],
                                'bold': False,
                                'is_hyperlink': False
                            })
                        
                        url_processed = True
                        break
                
                # If no URL found, add the whole content as regular text
                if not url_processed:
                    components.append({
                        'text': content_text,
                        'bold': False,
                        'is_hyperlink': False
                    })
                
                # Add line break after each content line
                components.append({'text': '\n', 'bold': False, 'is_hyperlink': False})
        
        # Add spacing between country blocks (except after the last one)
        if country_idx < len(countries) - 1:
            components.append({'text': '\n', 'bold': False, 'is_hyperlink': False})  # Extra line break
            print(f"     → Added spacing after {country_name} block")
    
    print(f"   ✅ Built {len(components)} text components for {len(countries)} countries")
    return components


def get_replacement_components(mapping_row: pd.Series, section_type: str, 
                             cached_components: Optional[List[Dict]] = None,
                             country_delimiter: str = ',') -> List[Dict]:
    """
    Get replacement components for a section, with reuse logic.
    
    Parameters:
    -----------
    mapping_row : pd.Series
        Row from mapping table
    section_type : str
        Either 'SmPC' or 'PL'
    cached_components : List[Dict], optional
        Previously generated components that can be reused
        
    Returns:
    --------
    List[Dict]
        List of formatted text components
    """
    # If we have cached components and texts are the same, reuse them
    if cached_components and can_reuse_replacement_text(mapping_row):
        print(f"♻️  Reusing cached components for {section_type}")
        return cached_components.copy()  # Return a copy to avoid mutation
    
    # Otherwise, build new components from lines
    return build_replacement_from_lines(mapping_row, section_type, country_delimiter)


def apply_formatted_replacement_v2(para: Paragraph, runs_to_replace: List[Run], components: List[Dict]):
    """
    Enhanced version of the replacement function with better hyperlink handling.
    """
    if not runs_to_replace or not components:
        print("   ⚠️  No runs to replace or no components provided")
        return

    print(f"   🔄 Applying replacement with {len(components)} components...")
    
    # Clear the first run and remove shading
    first_run = runs_to_replace[0]
    first_run.text = ''
    remove_shading_from_run(first_run)

    # Delete all subsequent runs in the replacement block
    for run in runs_to_replace[1:]:
        p = run._element.getparent()
        if p is not None:
            p.remove(run._element)

    # Insert the new formatted components
    current_element = first_run._element
    for i, comp in enumerate(components):
        if i == 0:
            run_to_modify = first_run
        else:
            new_run_element = OxmlElement('w:r')
            current_element.addnext(new_run_element)
            run_to_modify = Run(new_run_element, para)
            current_element = new_run_element
        
        run_to_modify.text = comp['text']
        
        # Apply formatting
        if comp.get('bold'):
            run_to_modify.bold = True
        if comp.get('is_hyperlink'):
            run_to_modify.font.color.rgb = RGBColor(0, 0, 255)
            run_to_modify.underline = True

    print("   ✅ Replacement applied successfully")


def remove_shading_from_run(run: Run):
    """
    Remove background shading from a run (gray highlighting).
    """
    try:
        # Access the run properties
        run_pr = run._element.get_or_add_rPr()
        
        # Remove shading element if it exists - fix namespace issue
        shading_elements = run_pr.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
        for shading in shading_elements:
            shading.getparent().remove(shading)
            
        print("     → Removed gray shading from run")
        
    except Exception as e:
        print(f"     ⚠️  Could not remove shading: {e}")
        # Non-critical error, continue processing


def run_annex_update_v2(doc: Document, mapping_row: pd.Series, section_type: str, 
                       cached_components: Optional[List[Dict]] = None,
                       country_delimiter: str = ',') -> Tuple[bool, Optional[List[Dict]]]:
    """
    Enhanced version of the annex update function that returns components for reuse.
    
    Parameters:
    -----------
    doc : Document
        The Word document to modify
    mapping_row : pd.Series
        Row from mapping table
    section_type : str
        Either 'SmPC' or 'PL'
    cached_components : List[Dict], optional
        Previously generated components for reuse
        
    Returns:
    --------
    Tuple[bool, List[Dict]]
        (success_flag, components_used) for potential reuse
    """
    print("\n" + "="*60)
    print(f"🔄 EXECUTING ENHANCED UPDATE FOR: Annex {('I' if section_type == 'SmPC' else 'IIIB')}")
    print("="*60)

    # Get the target string to find (what to replace)
    target_col = f'Original text national reporting - {section_type}'
    target_string = str(mapping_row.get(target_col, '')).strip()
    
    if ":" in target_string:
        target_string = target_string.split(':', 1)[-1].strip()

    print(f"   - Target string to replace: '{target_string[:100]}{'...' if len(target_string) > 100 else ''}'")

    if not target_string or target_string.lower() == 'nan':
        print("   ❌ No target string found in mapping file. Skipping.")
        return False, None

    # Get replacement components (either cached or newly built)
    components = get_replacement_components(mapping_row, section_type, cached_components, country_delimiter)
    
    if not components:
        print("   ❌ No replacement components generated. Skipping.")
        return False, None

    # Find and replace the text in the document
    found_replacement = False
    for para_idx, para in enumerate(doc.paragraphs):
        if target_string not in para.text:
            continue
            
        print(f"✅ Found target paragraph {para_idx}")
        print(f"   - Full paragraph text: '{para.text[:200]}{'...' if len(para.text) > 200 else ''}'")

        # Use existing logic to find content objects and runs to replace
        content_objects = get_paragraph_content(para)
        full_text = "".join(item['text'] for item in content_objects)
        
        print(f"   - Reconstructed text: '{full_text[:200]}{'...' if len(full_text) > 200 else ''}'")
        print(f"   - Target string: '{target_string}'")
        
        start_pos = full_text.find(target_string)
        
        if start_pos == -1:
            print(f"   ⚠️  Target string not found in reconstructed text, trying fuzzy match...")
            # Try to find a partial match or suggest the issue
            if target_string[:20] in full_text:
                print(f"   → Found partial match for first 20 chars")
            elif target_string[-20:] in full_text:
                print(f"   → Found partial match for last 20 chars")
            else:
                print(f"   → No partial matches found")
            continue
            
        end_pos = start_pos + len(target_string)
        print(f"   - Target text position: {start_pos} to {end_pos}")
        
        # Find objects to replace with precise boundary detection
        objects_to_replace = []
        char_count = 0
        print(f"   - Content objects analysis:")
        for i, item in enumerate(content_objects):
            item_len = len(item['text'])
            item_start = char_count
            item_end = char_count + item_len
            
            # Calculate overlap with target range
            overlap_start = max(item_start, start_pos)
            overlap_end = min(item_end, end_pos)
            overlap_length = max(0, overlap_end - overlap_start)
            
            # Only include if substantial overlap (>50% of item is within target range)
            if overlap_length > 0:
                overlap_percentage = overlap_length / item_len if item_len > 0 else 0
                item_within_target = (item_start >= start_pos and item_end <= end_pos)  # Completely within
                
                print(f"     {i+1}. {type(item['obj']).__name__}: '{item['text'][:50]}{'...' if len(item['text']) > 50 else ''}' ")
                print(f"        Position: {item_start}-{item_end}, Target: {start_pos}-{end_pos}")
                print(f"        Overlap: {overlap_length}/{item_len} chars ({overlap_percentage:.1%})")
                print(f"        Completely within target: {item_within_target}")
                
                # Include if: completely within target OR >80% overlap OR it's a hyperlink with any overlap
                should_include = (item_within_target or 
                                overlap_percentage > 0.8 or 
                                (hasattr(item['obj'], 'runs') and overlap_length > 0))  # Hyperlinks
                
                print(f"        → Include: {should_include}")
                
                if should_include:
                    objects_to_replace.append(item['obj'])
            else:
                print(f"     {i+1}. {type(item['obj']).__name__}: '{item['text'][:50]}{'...' if len(item['text']) > 50 else ''}' ")
                print(f"        Position: {item_start}-{item_end}, No overlap")
                
            char_count += item_len

        if not objects_to_replace:
            print("   ❌ Could not map target string to content objects")
            continue

        print(f"   - Will replace {len(objects_to_replace)} objects:")
        for i, obj in enumerate(objects_to_replace):
            obj_text = getattr(obj, 'text', 'No text attr')
            print(f"     {i+1}. {type(obj).__name__}: '{obj_text[:50]}{'...' if len(obj_text) > 50 else ''}'")

        # Convert to runs for replacement
        runs_for_replacement = []
        for item in objects_to_replace:
            if isinstance(item, Run):
                runs_for_replacement.append(item)
            elif hasattr(item, 'runs'):  # Hyperlink objects
                print(f"     → Converting hyperlink to {len(item.runs)} runs")
                runs_for_replacement.extend(item.runs)
            else:
                print(f"     ⚠️  Unknown object type: {type(item)}")

        print(f"   - Final runs to replace: {len(runs_for_replacement)}")

        # Apply the replacement
        apply_formatted_replacement_v2(para, runs_for_replacement, components)
        found_replacement = True
        break  # Assume one replacement per section

    if not found_replacement:
        print(f"⚠️  Target string not found in any paragraph")
        return False, components
        
    return True, components


def get_paragraph_content(paragraph: Paragraph) -> List[Dict]:
    """
    (From existing code) Deconstructs a paragraph into a list of its
    content objects, safely handling runs that might not contain text.
    Enhanced with better debugging.
    """
    content = []
    print(f"     🔍 Analyzing paragraph structure:")
    
    for i, child_element in enumerate(paragraph._p):
        text_content = ''  # Default to empty string
        
        if child_element.tag.endswith('r'):  # Run element
            run = Run(child_element, paragraph)
            # Safely get text, substituting None with an empty string
            text_content = run.text or ''
            content.append({'obj': run, 'text': text_content})
            print(f"       {i+1}. Run: '{text_content[:30]}{'...' if len(text_content) > 30 else ''}'")
            
        elif child_element.tag.endswith('hyperlink'):  # Hyperlink element
            try:
                link = Hyperlink(child_element, paragraph)
                # Safely get text from the hyperlink
                text_content = link.text or ''
                content.append({'obj': link, 'text': text_content})
                print(f"       {i+1}. Hyperlink: '{text_content[:30]}{'...' if len(text_content) > 30 else ''}' (runs: {len(link.runs) if hasattr(link, 'runs') else 'unknown'})")
            except Exception as e:
                print(f"       {i+1}. Hyperlink: Error reading ({e})")
                # Try to get text directly from element
                text_content = child_element.text or ''
                content.append({'obj': child_element, 'text': text_content})
        else:
            # Other element types
            element_name = child_element.tag.split('}')[-1] if '}' in child_element.tag else child_element.tag
            print(f"       {i+1}. {element_name}: (skipped)")
    
    print(f"     → Total content objects: {len(content)}")
    return content


def load_docx_file(file_path: str) -> Optional[Document]:
    """
    Safely load a docx file and return a Document object.
    """
    try:
        path = Path(file_path)
        
        if not path.exists():
            print(f"❌ Error: File not found: {file_path}")
            return None
            
        if not path.is_file():
            print(f"❌ Error: Path is not a file: {file_path}")
            return None
            
        if path.suffix.lower() not in ['.docx', '.doc']:
            print(f"⚠️  Warning: File may not be a Word document: {path.suffix}")
        
        doc = Document(file_path)
        print(f"✅ Successfully loaded: {path.name}")
        print(f"   - Paragraphs: {len(doc.paragraphs)}")
        print(f"   - Tables: {len(doc.tables)}")
        
        return doc
        
    except Exception as e:
        print(f"❌ Error loading document: {type(e).__name__}: {str(e)}")
        return None


def load_mapping_table(csv_path: str, encoding: str = 'utf-8') -> Optional[pd.DataFrame]:
    """
    Load the mapping table from CSV file with error handling.
    """
    try:
        path = Path(csv_path)
        
        if not path.exists():
            print(f"❌ Error: Mapping file not found: {csv_path}")
            return None
            
        df = pd.read_csv(csv_path, encoding=encoding)
        
        print(f"✅ Successfully loaded mapping table: {path.name}")
        print(f"   - Rows: {len(df)}")
        print(f"   - Columns: {len(df.columns)}")
        
        return df
        
    except UnicodeDecodeError:
        print(f"⚠️  Encoding error with {encoding}, trying 'latin-1'...")
        try:
            df = pd.read_csv(csv_path, encoding='latin-1')
            print(f"✅ Successfully loaded with latin-1 encoding")
            return df
        except Exception as e:
            print(f"❌ Error: {type(e).__name__}: {str(e)}")
            return None
            
    except Exception as e:
        print(f"❌ Error loading CSV: {type(e).__name__}: {str(e)}")
        return None


def get_country_code_mapping() -> Dict[str, Tuple[str, str]]:
    """
    Return a mapping of two-letter codes to (language, country).
    """
    return {
        'bg': ('Bulgarian', 'Bulgaria'),
        'hr': ('Croatian', 'Croatia'),
        'cs': ('Czech', 'Czech Republic'),
        'da': ('Danish', 'Denmark'),
        'nl': ('Dutch', 'Netherlands'),
        'en': ('English', 'Ireland'),
        'et': ('Estonian', 'Estonia'),
        'fi': ('Finnish', 'Finland'),
        'fr': ('French', 'France'),
        'de': ('German', 'Germany'),
        'el': ('Greek', 'Greece'),
        'hu': ('Hungarian', 'Hungary'),
        'is': ('Icelandic', 'Iceland'),
        'it': ('Italian', 'Italy'),
        'lv': ('Latvian', 'Latvia'),
        'lt': ('Lithuanian', 'Lithuania'),
        'mt': ('Maltese', 'Malta'),
        'no': ('Norwegian', 'Norway'),
        'pl': ('Polish', 'Poland'),
        'pt': ('Portuguese', 'Portugal'),
        'ro': ('Romanian', 'Romania'),
        'sk': ('Slovak', 'Slovakia'),
        'sl': ('Slovenian', 'Slovenia'),
        'es': ('Spanish', 'Spain'),
        'sv': ('Swedish', 'Sweden')
    }


def extract_country_code_from_filename(file_path: str) -> Optional[str]:
    """
    Extract country code from filename.
    """
    try:
        filename = Path(file_path).stem
        
        # Pattern for annotated files: ema-combined-h-XXXX-XX-annotated
        pattern1 = r'ema-combined-h-\d+-([a-z]{2})-annotated'
        match = re.search(pattern1, filename, re.IGNORECASE)
        
        if match:
            country_code = match.group(1).lower()
            print(f"📝 Found country code in filename: '{country_code}'")
            return country_code
        
        # Pattern for test files: ema-combined-h-XXXX-XX_Test
        pattern2 = r'ema-combined-h-\d+-([a-z]{2})[-_]'
        match = re.search(pattern2, filename, re.IGNORECASE)
        
        if match:
            country_code = match.group(1).lower()
            print(f"📝 Found country code in filename: '{country_code}'")
            return country_code
            
        print(f"⚠️  No country code found in filename: {filename}")
        return None
        
    except Exception as e:
        print(f"❌ Error extracting country code: {type(e).__name__}: {str(e)}")
        return None


def identify_document_country_and_language(file_path: str) -> Tuple[Optional[str], Optional[str], Optional[str]]:
    """
    Identify both country and language from a document filename.
    """
    print(f"\n🔍 Identifying country and language for: {Path(file_path).name}")
    
    country_code = extract_country_code_from_filename(file_path)
    
    if country_code:
        country_mapping = get_country_code_mapping()
        
        if country_code in country_mapping:
            language_name, country_name = country_mapping[country_code]
            print(f"✅ Identified: {country_code} → {language_name} ({country_name})")
            return country_code, language_name, country_name
        else:
            print(f"⚠️  Unknown country code: {country_code}")
            
    return country_code, None, None


def find_mapping_row_by_language_and_country(mapping_df: pd.DataFrame, 
                                            language_name: str, 
                                            country_name: Optional[str] = None) -> Optional[pd.Series]:
    """
    Find the appropriate mapping row based on language and optionally country.
    """
    special_cases = {
        'English': 'Ireland',
        'French': 'France',
        'German': 'Germany',
        'Dutch': 'Netherlands'
    }
    
    # First try exact match with country if provided
    if country_name:
        exact_match = mapping_df[
            (mapping_df['Language'].str.lower() == language_name.lower()) &
            (mapping_df['Country'].str.lower() == country_name.lower())
        ]
        if len(exact_match) == 1:
            print(f"✅ Found exact match: {language_name} - {country_name}")
            return exact_match.iloc[0]
    
    # Try language match only
    language_matches = mapping_df[mapping_df['Language'].str.lower() == language_name.lower()]
    
    if len(language_matches) == 0:
        print(f"❌ No mapping found for language: {language_name}")
        return None
    elif len(language_matches) == 1:
        print(f"✅ Found unique match: {language_name} - {language_matches.iloc[0]['Country']}")
        return language_matches.iloc[0]
    else:
        # Multiple matches - use special case defaults
        if language_name in special_cases:
            default_country = special_cases[language_name]
            default_match = language_matches[
                language_matches['Country'].str.contains(default_country, case=False)
            ]
            if len(default_match) >= 1:
                selected_country = default_match.iloc[0]['Country']
                print(f"✅ Multiple countries for {language_name}, defaulting to: {selected_country}")
                return default_match.iloc[0]
        
        print(f"⚠️  Multiple matches for {language_name}, auto-selecting: {language_matches.iloc[0]['Country']}")
        return language_matches.iloc[0]


def test_enhanced_workflow(file_path: str, mapping_csv_path: str):
    """
    Test the enhanced workflow with country code detection.
    """
    print("="*80)
    print("TESTING ENHANCED WORKFLOW")
    print("="*80)
    
    # Step 1: Load document
    doc = load_docx_file(file_path)
    if not doc:
        return None
    
    # Step 2: Identify country and language from filename
    country_code, language_name, country_name = identify_document_country_and_language(file_path)
    
    if not language_name:
        print("❌ Cannot proceed without language identification")
        return None
    
    print("\n" + "-"*40)
    
    # Step 3: Load mapping table
    mapping_df = load_mapping_table(mapping_csv_path)
    if mapping_df is None:
        return None
    
    print("\n" + "-"*40)
    
    # Step 4: Find appropriate mapping row
    mapping_row = find_mapping_row_by_language_and_country(mapping_df, language_name, country_name)
    
    if mapping_row is not None:
        print(f"\n✅ MAPPING FOUND:")
        print(f"   - Country: {mapping_row['Country']}")
        print(f"   - Language: {mapping_row['Language']}")
        print(f"   - Has SmPC text: {'Yes' if pd.notna(mapping_row['National reporting system SmPC']) else 'No'}")
        print(f"   - Has PL text: {'Yes' if pd.notna(mapping_row['National reporting system PL']) else 'No'}")
        
        # Show output filenames
        print(f"\n📄 Output files will be:")
        print(f"   - Annex I_EU_SmPC_{mapping_row['Language']}")
        print(f"   - Annex IIIB_EU_PL_{mapping_row['Language']}")
    
    return {
        'doc': doc,
        'file': Path(file_path).name,
        'country_code': country_code,
        'language_name': language_name,
        'country_name': country_name,
        'mapping_row': mapping_row,
        'mapping_country': mapping_row['Country'] if mapping_row is not None else None
    }


def process_document_with_enhanced_workflow(doc_path: str, mapping_csv_path: str, 
                                          output_path: Optional[str] = None,
                                          country_delimiter: str = ','):
    """
    Enhanced workflow that uses the new line-based replacement system.
    """
    print("="*80)
    print(f"STARTING ENHANCED WORKFLOW FOR: {Path(doc_path).name}")
    print("="*80)

    # Load document and mapping (reuse existing functions)
    result = test_enhanced_workflow(doc_path, mapping_csv_path)
    if not result or result['mapping_row'] is None:
        print("❌ ABORTING: Could not load document or find mapping row.")
        return None
    
    doc = result['doc']
    mapping_row = result['mapping_row']
    
    # Enhanced replacement with component reuse
    smpc_success, smpc_components = run_annex_update_v2(doc, mapping_row, 'SmPC', None, country_delimiter)
    pl_success, pl_components = run_annex_update_v2(doc, mapping_row, 'PL', smpc_components, country_delimiter)
    
    # Save document if changes made
    if (smpc_success or pl_success) and output_path:
        try:
            Path(output_path).parent.mkdir(parents=True, exist_ok=True)
            doc.save(output_path)
            print(f"\n💾 Document saved successfully to: {output_path}")
        except Exception as e:
            print(f"\n❌ Error saving document: {e}")
    else:
        print("\n" + "="*70)
        print("No changes were made, so the document was not saved.")
        print("="*70)
    
    return {
        'smpc_updated': smpc_success,
        'pl_updated': pl_success,
        'components_reused': smpc_components == pl_components if smpc_components and pl_components else False
    }


# ============= TESTING AND EXAMPLE USAGE =============

def test_line_parsing(mapping_csv_path: str, country_name: str = 'Ireland', 
                     section_type: str = 'SmPC', country_delimiter: str = ','):
    """
    Test function to see how the line parsing works for a specific country.
    """
    print("="*60)
    print(f"TESTING LINE PARSING FOR: {country_name} - {section_type}")
    print("="*60)
    
    # Load mapping table
    mapping_df = load_mapping_table(mapping_csv_path)
    if mapping_df is None:
        return None
    
    # Find the row for the specified country
    country_row = mapping_df[mapping_df['Country'].str.contains(country_name, case=False)]
    
    if len(country_row) == 0:
        print(f"❌ No mapping found for country: {country_name}")
        return None
    
    mapping_row = country_row.iloc[0]
    print(f"✅ Found mapping for: {mapping_row['Country']} - {mapping_row['Language']}")
    
    # Test the line parsing
    components = build_replacement_from_lines(mapping_row, section_type, country_delimiter)
    
    print(f"\n📝 GENERATED COMPONENTS:")
    print("-" * 40)
    for i, comp in enumerate(components):
        formatting = []
        if comp.get('bold'):
            formatting.append('BOLD')
        if comp.get('is_hyperlink'):
            formatting.append('HYPERLINK')
        
        display_text = repr(comp['text']) if comp['text'] in ['\n', '\n\n'] else comp['text']
        format_str = f" [{', '.join(formatting)}]" if formatting else ""
        
        print(f"{i+1:2d}. {display_text}{format_str}")
    
    # Test reuse logic
    can_reuse = can_reuse_replacement_text(mapping_row)
    print(f"\n♻️  Can reuse components for other section: {can_reuse}")
    
    return components


def debug_target_text_detection(doc_path: str, mapping_csv_path: str, section_type: str = 'SmPC'):
    """
    Debug function to analyze target text detection issues.
    """
    print("="*80)
    print(f"DEBUGGING TARGET TEXT DETECTION: {section_type}")
    print("="*80)
    
    # Load document and mapping
    result = test_enhanced_workflow(doc_path, mapping_csv_path)
    if not result or result['mapping_row'] is None:
        print("❌ Could not load document or mapping")
        return
    
    doc = result['doc']
    mapping_row = result['mapping_row']
    
    # Get target string
    target_col = f'Original text national reporting - {section_type}'
    target_string = str(mapping_row.get(target_col, '')).strip()
    
    if ":" in target_string:
        target_string = target_string.split(':', 1)[-1].strip()
    
    print(f"🔍 TARGET STRING: '{target_string}'")
    print(f"   Length: {len(target_string)} characters")
    
    # Search through all paragraphs
    matches_found = 0
    for para_idx, para in enumerate(doc.paragraphs):
        para_text = para.text
        
        # Check for exact match
        if target_string in para_text:
            matches_found += 1
            print(f"\n✅ EXACT MATCH in paragraph {para_idx}:")
            print(f"   Full text: '{para_text}'")
            print(f"   Length: {len(para_text)} characters")
            
            # Analyze the paragraph structure
            content_objects = get_paragraph_content(para)
            reconstructed = "".join(item['text'] for item in content_objects)
            
            print(f"   Reconstructed: '{reconstructed}'")
            print(f"   Match: {target_string in reconstructed}")
            
            # Find the position
            start_pos = para_text.find(target_string)
            end_pos = start_pos + len(target_string)
            print(f"   Position: {start_pos} to {end_pos}")
            print(f"   Before: '{para_text[:start_pos]}'")
            print(f"   Target: '{para_text[start_pos:end_pos]}'")  
            print(f"   After: '{para_text[end_pos:]}'")
            
        # Check for partial matches
        elif len(target_string) > 20:
            start_partial = target_string[:20]
            end_partial = target_string[-20:]
            
            if start_partial in para_text or end_partial in para_text:
                print(f"\n⚠️  PARTIAL MATCH in paragraph {para_idx}:")
                print(f"   Full text: '{para_text[:100]}{'...' if len(para_text) > 100 else ''}'")
                if start_partial in para_text:
                    print(f"   → Contains start: '{start_partial}'")
                if end_partial in para_text:
                    print(f"   → Contains end: '{end_partial}'")
    
    print(f"\n📊 SUMMARY: Found {matches_found} exact matches")
    
    if matches_found == 0:
        print("\n🔍 SEARCHING FOR SIMILAR TEXT...")
        # Look for paragraphs that might contain similar text
        keywords = target_string.split()[:3]  # First 3 words
        for para_idx, para in enumerate(doc.paragraphs):
            if any(keyword.lower() in para.text.lower() for keyword in keywords):
                print(f"   Paragraph {para_idx} contains keywords: '{para.text[:100]}{'...' if len(para.text) > 100 else ''}'")


def show_available_countries(mapping_csv_path: str):
    """
    Display all available countries in the mapping table for testing.
    """
    mapping_df = load_mapping_table(mapping_csv_path)
    if mapping_df is None:
        return
    
    print("📋 AVAILABLE COUNTRIES:")
    print("-" * 40)
    for idx, row in mapping_df.iterrows():
        print(f"  - {row['Country']} ({row['Language']})")


def debug_column_names(mapping_csv_path: str, country_name: str = 'Ireland'):
    """
    Debug function to see exactly what column names exist in the mapping file.
    """
    print("="*60)
    print(f"DEBUGGING COLUMN NAMES FOR: {country_name}")
    print("="*60)
    
    mapping_df = load_mapping_table(mapping_csv_path)
    if mapping_df is None:
        return
    
    # Find the row for the specified country
    country_row = mapping_df[mapping_df['Country'].str.contains(country_name, case=False)]
    if len(country_row) == 0:
        print(f"❌ No mapping found for country: {country_name}")
        return
    
    mapping_row = country_row.iloc[0]
    print(f"✅ Found mapping for: {mapping_row['Country']} - {mapping_row['Language']}")
    
    # Show all column names
    print(f"\n📋 ALL COLUMN NAMES ({len(mapping_row.index)} total):")
    print("-" * 60)
    for i, col in enumerate(mapping_row.index, 1):
        print(f"{i:2d}. {col}")
    
    # Show Line columns specifically
    line_cols = [col for col in mapping_row.index if 'Line ' in col]
    print(f"\n📋 LINE COLUMNS FOUND ({len(line_cols)} total):")
    print("-" * 60)
    for col in line_cols:
        value = str(mapping_row[col])[:50]
        print(f"  - {col}")
        print(f"    Value: '{value}{'...' if len(str(mapping_row[col])) > 50 else ''}'")
    
    # Test pattern matching
    print(f"\n🔍 PATTERN MATCHING TEST:")
    print("-" * 60)
    smpc_matches = [col for col in mapping_row.index 
                   if col.startswith('Line ') and ('SmPC' in col or 'SmpC' in col)]
    pl_matches = [col for col in mapping_row.index 
                 if col.startswith('Line ') and 'PL' in col]
    
    print(f"  SmPC pattern matches: {smpc_matches}")
    print(f"  PL pattern matches: {pl_matches}")


def show_available_countries(mapping_csv_path: str):
    """
    Display all available countries in the mapping table for testing.
    """
    mapping_df = load_mapping_table(mapping_csv_path)
    if mapping_df is None:
        return
    
    print("📋 AVAILABLE COUNTRIES:")
    print("-" * 40)
    for idx, row in mapping_df.iterrows():
        print(f"  - {row['Country']} ({row['Language']})")


# In[18]:


PATHS = {'mapping_table': Path('../data/Mapping Test.csv'),
        'Data_root': Path('../data/'),
        'en_test' : Path('../data/Docs/ema-combined-h-4844-en_Test.docx'),
        'es_test': Path('../data/Docs/ema-combined-h-4844-es_Test.docx')}


# In[21]:


# Create an output directory if it doesn't exist
output_dir = Path('../outputs/')
output_dir.mkdir(exist_ok=True)

# process_document_with_new_workflow(
#     PATHS['en_test'],
#     PATHS['mapping_table'],
#     output_dir / 'english_output.docx' # Use a new output name
# )


# In[25]:


# Show available countries for testing
print("Available countries for testing:")
show_available_countries(PATHS['mapping_table'])

print("\n" + "="*80)

# Test line parsing for Ireland SmPC
# components = test_line_parsing(PATHS['mapping_table'], 'Ireland', 'SmPC')

print("\n" + "="*80)

# Run full workflow
result = process_document_with_enhanced_workflow(PATHS['en_test'], PATHS['mapping_table'],output_dir / 'english_output.docx')


# In[ ]:





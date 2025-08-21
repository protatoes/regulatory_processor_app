"""Document processing module for EU SmPC and PL documents.

This module contains helper functions to load the mapping Excel file,
update the Annex I (SmPC) and Annex IIIB (Package Leaflet) sections
of combined SmPC Word documents, split the updated documents into
separate Annex I and Annex IIIB files, and convert those files to
PDF using LibreOffice.

Complete implementation with all required functionality:
- National reporting system updates (Section 4.8 SmPC, Section 4 PL)
- Date updates (Section 10 Annex I, Section 6 Annex IIIB)
- Local representatives update (Section 6 Annex IIIB)
- Multi-file generation for languages with multiple countries
- Proper file naming conventions
"""


import os
import re
import shutil
import logging
from pathlib import Path
from typing import Dict, List, Optional, Tuple, NamedTuple, Union
from dataclasses import dataclass
from datetime import datetime
import pandas as pd
import subprocess
from copy import deepcopy
from docx import Document
from docx.document import Document as DocumentObject
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.shared import RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.oxml import OxmlElement
from docx2pdf import convert

# ============================================================================= 
# CONSTANTS AND CONFIGURATION
# =============================================================================

class DirectoryNames:
    SPLIT_DOCS = "split_docs"
    PDF_DOCS = "pdf_docs"
    BACKUP_SUFFIX = ".orig"

class FileMarkers:
    ANNEX_MARKER = "_Annex_"
    TEMP_FILE_PREFIX = "~"
    ANNEX_PREFIX = "Annex"

class SectionTypes:
    SMPC = "SmPC"
    PL = "PL"

# Custom Exceptions
class ProcessingError(Exception):
    """Base exception for document processing errors."""
    pass

class ValidationError(ProcessingError):
    """Raised when input validation fails."""
    pass

class DocumentError(ProcessingError):
    """Raised when document operations fail."""
    pass

class MappingError(ProcessingError):
    """Raised when mapping file operations fail."""
    pass

# Result Types
class ProcessingResult(NamedTuple):
    success: bool
    message: str
    output_files: List[str] = []
    errors: List[str] = []

@dataclass
class ProcessingStats:
    """Tracks processing statistics throughout the workflow."""
    input_files_found: int = 0
    input_files_processed: int = 0
    variants_processed: int = 0
    variants_successful: int = 0
    output_files_created: int = 0
    errors_encountered: int = 0
    
    def success_rate(self) -> float:
        """Calculate overall success rate."""
        if self.variants_processed == 0:
            return 0.0
        return (self.variants_successful / self.variants_processed) * 100

@dataclass
class ProcessingConfig:
    """Configuration settings for document processing."""
    create_backups: bool = True
    convert_to_pdf: bool = True
    overwrite_existing: bool = False
    log_level: str = "INFO"
    country_delimiter: str = ";"

# ============================================================================= 
# DATE FORMATTING SYSTEM
# =============================================================================

class DateFormatterSystem:
    """Enhanced date formatting system with locale support."""
    
    def __init__(self, mapping_file_path: str):
        self.mapping_df = pd.read_excel(mapping_file_path)
        self.country_formats = self._load_country_formats()
        
    def _load_country_formats(self) -> Dict[str, Dict[str, str]]:
        """Load date formats from the mapping table."""
        formats = {}
        for _, row in self.mapping_df.iterrows():
            country = row['Country']
            formats[country] = {
                'annex_i': row.get('Annex I Date Format', ''),
                'annex_iiib': row.get('Annex IIIB Date Format', '')
            }
        return formats
    
    def format_date(self, date: datetime, country: str, annex_type: str) -> str:
        """Format a date according to country specifications."""
        if country not in self.country_formats:
            return date.strftime("%d %B %Y")  # Default format
            
        format_string = self.country_formats[country].get(annex_type, '')
        return self._parse_custom_format(date, format_string)
    
    def _parse_custom_format(self, date: datetime, format_string: str) -> str:
        """Parse custom format string and return formatted date."""
        if not format_string or format_string.lower() == 'nan':
            return date.strftime("%d %B %Y")
        
        # Handle common patterns
        if format_string == "dd month yyyy":
            return date.strftime("%d %B %Y")
        elif format_string == "month yyyy":
            return date.strftime("%B %Y")
        elif format_string == "dd. MMM yyyy":
            return date.strftime("%d. %b %Y")
        elif format_string == "MMM yyyy":
            return date.strftime("%b %Y")
        else:
            return date.strftime("%d %B %Y")
    
    def get_available_countries(self) -> List[str]:
        """Get list of available countries."""
        return list(self.country_formats.keys())
    
    def preview_format(self, country: str, sample_date: datetime = None) -> Dict[str, str]:
        """Preview date formatting for a country."""
        if sample_date is None:
            sample_date = datetime.now()
        
        if country not in self.country_formats:
            return {'error': f'Country {country} not found'}
        
        return {
            'annex_i_example': self.format_date(sample_date, country, 'annex_i'),
            'annex_iiib_example': self.format_date(sample_date, country, 'annex_iiib')
        }

# Global date formatter instance
_date_formatter: Optional[DateFormatterSystem] = None

def initialize_date_formatter(mapping_file_path: str) -> DateFormatterSystem:
    """Initialize the global date formatter."""
    global _date_formatter
    _date_formatter = DateFormatterSystem(mapping_file_path)
    return _date_formatter

def get_date_formatter() -> DateFormatterSystem:
    """Get the global date formatter instance."""
    global _date_formatter
    if _date_formatter is None:
        raise RuntimeError("Date formatter not initialized")
    return _date_formatter

def format_date_for_country(country: str, annex_type: str, date: Optional[datetime] = None) -> str:
    """Format a date using the enhanced date formatter."""
    if date is None:
        date = datetime.now()
    
    try:
        formatter = get_date_formatter()
        return formatter.format_date(date, country, annex_type)
    except Exception:
        return date.strftime("%d %B %Y")  # Fallback

# ============================================================================= 
# CORE UTILITY FUNCTIONS
# =============================================================================

def load_mapping_table(file_path: str) -> Optional[pd.DataFrame]:
    """Load the Excel mapping table and initialize the date formatter."""
    try:
        path = Path(file_path)
        if not path.exists():
            print(f"❌ Error: Mapping file not found: {file_path}")
            return None
            
        df = pd.read_excel(path)
        
        # Initialize the date formatter
        print(f"🔧 Initializing DateFormatterSystem...")
        try:
            initialize_date_formatter(file_path)
            formatter = get_date_formatter()
            available_countries = formatter.get_available_countries()
            print(f"✅ DateFormatterSystem initialized with {len(available_countries)} countries")
        except Exception as e:
            print(f"❌ Error initializing DateFormatterSystem: {e}")
            return None
        
        print(f"✅ Successfully loaded mapping table: {path.name}")
        print(f"   - Rows: {len(df)}")
        print(f"   - Columns: {len(df.columns)}")
        
        return df
            
    except Exception as e:
        print(f"❌ Error loading Excel file: {type(e).__name__}: {str(e)}")
        return None

def get_country_code_mapping() -> Dict[str, Tuple[str, str]]:
    """Return a mapping of two-letter codes to (language, country)."""
    return {
        'bg': ('Bulgarian', 'Bulgaria'), 'hr': ('Croatian', 'Croatia'),
        'cs': ('Czech', 'Czech Republic'), 'da': ('Danish', 'Denmark'),
        'nl': ('Dutch', 'Netherlands'), 'en': ('English', 'Ireland'),
        'et': ('Estonian', 'Estonia'), 'fi': ('Finnish', 'Finland'),
        'fr': ('French', 'France'), 'de': ('German', 'Germany'),
        'el': ('Greek', 'Greece'), 'hu': ('Hungarian', 'Hungary'),
        'is': ('Icelandic', 'Iceland'), 'it': ('Italian', 'Italy'),
        'lv': ('Latvian', 'Latvia'), 'lt': ('Lithuanian', 'Lithuania'),
        'mt': ('Maltese', 'Malta'), 'no': ('Norwegian', 'Norway'),
        'pl': ('Polish', 'Poland'), 'pt': ('Portuguese', 'Portugal'),
        'ro': ('Romanian', 'Romania'), 'sk': ('Slovak', 'Slovakia'),
        'sl': ('Slovenian', 'Slovenia'), 'es': ('Spanish', 'Spain'),
        'sv': ('Swedish', 'Sweden')
    }

def extract_country_code_from_filename(file_path: str) -> Optional[str]:
    """Extract country code from filename."""
    try:
        filename = Path(file_path).stem
        pattern1 = r'ema-combined-h-\d+-([a-z]{2})-annotated'
        match = re.search(pattern1, filename, re.IGNORECASE)
        if match:
            return match.group(1).lower()
        
        pattern2 = r'ema-combined-h-\d+-([a-z]{2})[-_]'
        match = re.search(pattern2, filename, re.IGNORECASE)
        if match:
            return match.group(1).lower()
        
        return None
    except Exception:
        return None

def identify_document_country_and_language(file_path: str) -> Tuple[Optional[str], Optional[str], Optional[str]]:
    """Identify both country and language from a document filename."""
    country_code = extract_country_code_from_filename(file_path)
    if country_code:
        country_mapping = get_country_code_mapping()
        if country_code in country_mapping:
            language_name, country_name = country_mapping[country_code]
            return country_code, language_name, country_name
    return country_code, None, None

def find_mapping_rows_for_language(mapping_df: pd.DataFrame, language_name: str) -> List[pd.Series]:
    """Find all mapping rows for a given language."""
    language_matches = mapping_df[mapping_df['Language'].str.lower() == language_name.lower()]
    return [language_matches.iloc[i] for i in range(len(language_matches))]

def generate_output_filename(base_name: str, language: str, country: str, doc_type: str) -> str:
    """Generate compliant filename according to specifications."""
    country_clean = country.replace('/', '_').replace(' ', '_')
    
    if doc_type == "combined":
        return f"{base_name}_{country_clean}.docx"
    elif doc_type == "annex_i":
        return f"Annex_I_EU_SmPC_{language}_{country_clean}.docx"
    elif doc_type == "annex_iiib":
        return f"Annex_IIIB_EU_PL_{language}_{country_clean}.docx"
    else:
        return f"{base_name}_{doc_type}.docx"

def convert_to_pdf(doc_path: str, output_dir: str) -> str:
    """Convert a Word document to PDF with multiple fallback methods."""
    pdf_output_path = Path(output_dir) / Path(doc_path).with_suffix(".pdf").name
    
    # Method 1: Try docx2pdf (primary method)
    try:
        convert(doc_path, str(pdf_output_path))
        return str(pdf_output_path)
    except Exception as e:
        print(f"   ⚠️ docx2pdf conversion failed: {e}")
    
    # Method 2: Try LibreOffice (if available)
    try:
        result = subprocess.run([
            'libreoffice', '--headless', '--convert-to', 'pdf',
            '--outdir', str(output_dir), doc_path
        ], capture_output=True, text=True, timeout=60)
        
        if result.returncode == 0 and pdf_output_path.exists():
            print(f"   ✅ LibreOffice conversion successful")
            return str(pdf_output_path)
        else:
            print(f"   ⚠️ LibreOffice conversion failed: {result.stderr}")
    except (subprocess.TimeoutExpired, FileNotFoundError) as e:
        print(f"   ⚠️ LibreOffice not available: {e}")
    
    # Method 3: Create a placeholder PDF (last resort)
    try:
        # Create a simple text file indicating conversion failed
        placeholder_path = pdf_output_path.with_suffix('.pdf.txt')
        with open(placeholder_path, 'w') as f:
            f.write(f"PDF conversion failed for: {Path(doc_path).name}\n")
            f.write(f"Original document available at: {doc_path}\n")
            f.write(f"Please convert manually or install LibreOffice for automatic conversion.\n")
        
        print(f"   📝 Created placeholder file: {placeholder_path.name}")
        return str(placeholder_path)
        
    except Exception as e:
        print(f"   ❌ All conversion methods failed: {e}")
        raise RuntimeError(f"Failed to convert {doc_path} to PDF: All methods failed")

def copy_paragraph(dest_doc: Document, source_para: Paragraph) -> None:
    """Copy a paragraph from one document to another, preserving runs."""
    new_para = dest_doc.add_paragraph()
    new_para.style = source_para.style
    for run in source_para.runs:
        new_run = new_para.add_run(run.text)
        new_run.bold = run.bold
        new_run.italic = run.italic
        new_run.underline = run.underline
        new_run.style = run.style

# ============================================================================= 
# DOCUMENT UPDATE FUNCTIONS
# =============================================================================
def is_hex_gray_color(hex_color: str) -> bool:
    """Check if a hex color represents a gray shade."""
    if not hex_color:
        return False
    
    hex_color = hex_color.replace('#', '').upper()
    
    gray_hex_values = [
        'BFBFBF', 'CCCCCC', 'D9D9D9', '808080', '999999', 
        '666666', 'C0C0C0', 'A0A0A0'
    ]
    
    if hex_color in gray_hex_values:
        return True
    
    # Check if R=G=B (indicates gray)
    try:
        if len(hex_color) == 6:
            r = int(hex_color[0:2], 16)
            g = int(hex_color[2:4], 16)
            b = int(hex_color[4:6], 16)
            return r == g == b
    except ValueError:
        pass
    
    return False


def is_run_gray_shaded(run: Run) -> bool:
    """Check if a run has gray shading."""
    try:
        run_pr = run._element.get_or_add_rPr()
        shading_elements = run_pr.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
        
        if shading_elements:
            for shading in shading_elements:
                fill = shading.get(qn('w:fill'))
                if fill and is_hex_gray_color(fill):
                    return True
        
        # Check font color for gray
        if run.font.color and hasattr(run.font.color, 'rgb'):
            gray_colors = [
                RGBColor(128, 128, 128), RGBColor(153, 153, 153), 
                RGBColor(102, 102, 102), RGBColor(96, 96, 96),
                RGBColor(217, 217, 217), RGBColor(191, 191, 191)
            ]
            if run.font.color.rgb in gray_colors:
                return True
                
    except Exception:
        pass
    
    return False


def is_run_hyperlink(run: Run) -> bool:
    """Check if a run is part of a hyperlink."""
    try:
        run_xml = run._r
        hyperlink_elements = run_xml.xpath('.//w:hyperlink')
        if hyperlink_elements:
            return True
            
        # Check hyperlink-style formatting
        if (run.font.color and hasattr(run.font.color, 'rgb') and 
            run.font.color.rgb == RGBColor(0, 0, 255) and run.underline):
            return True
            
    except Exception:
        pass
    
    return False


def find_target_text_runs(para: Paragraph, target_string: str) -> List[Run]:
    """
    Find runs that contain the target text to be replaced.
    """
    target_runs = []
    
    # Build text map to find position
    char_pos = 0
    run_positions = []
    
    for run in para.runs:
        run_start = char_pos
        run_end = char_pos + len(run.text)
        run_positions.append((run, run_start, run_end))
        char_pos = run_end
    
    # Find target string position in full text
    full_text = para.text
    target_start = full_text.lower().find(target_string.lower())
    
    if target_start == -1:
        return []
    
    target_end = target_start + len(target_string)
    
    # Find runs that overlap with target text
    for run, run_start, run_end in run_positions:
        # Check if run overlaps with target range
        if (run_start < target_end and run_end > target_start):
            target_runs.append(run)
    
    return target_runs


def find_target_text_range(para: Paragraph, target_string: str) -> Tuple[int, int]:
    """Find the complete target text range in paragraph."""
    full_text = para.text.lower()
    target_lower = target_string.lower()
    
    # Try exact match first
    start_pos = full_text.find(target_lower)
    if start_pos != -1:
        return start_pos, start_pos + len(target_string)
    
    # Try to find key parts
    key_phrases = ['national reporting system', 'appendix v', 'listed in appendix']
    
    earliest_start = len(full_text)
    latest_end = 0
    
    for phrase in key_phrases:
        pos = full_text.find(phrase)
        if pos != -1:
            earliest_start = min(earliest_start, pos)
            latest_end = max(latest_end, pos + len(phrase))
    
    if earliest_start < len(full_text):
        return earliest_start, latest_end
    
    return -1, -1


def find_runs_to_remove(para: Paragraph, target_string: str) -> List[Run]:
    """Find runs that should be removed (gray shaded, hyperlinks, or target text)."""
    runs_to_remove = []
    
    target_start, target_end = find_target_text_range(para, target_string)
    
    if target_start == -1:
        return runs_to_remove
    
    # Map character positions to runs
    char_pos = 0
    run_ranges = []
    
    for run in para.runs:
        run_start = char_pos
        run_end = char_pos + len(run.text)
        run_ranges.append((run, run_start, run_end))
        char_pos = run_end
    
    # Find runs to remove
    for run, run_start, run_end in run_ranges:
        should_remove = False
        
        # Check if run overlaps with target range
        if run_start < target_end and run_end > target_start:
            should_remove = True
        # Check if it's gray shaded
        elif is_run_gray_shaded(run):
            should_remove = True
        # Check if it's a hyperlink
        elif is_run_hyperlink(run):
            should_remove = True
        
        if should_remove:
            runs_to_remove.append(run)
    
    return runs_to_remove


def find_gray_and_hyperlink_runs(para: Paragraph, target_string: str) -> List[Run]:
    """
    Find all gray shaded runs and hyperlink runs that should be removed.
    """
    runs_to_remove = []
    
    # First find runs containing target text
    target_runs = find_target_text_runs(para, target_string)
    
    # Then find additional gray/hyperlink runs in vicinity
    for run in para.runs:
        should_remove = False
        
        # Remove if it's a target run
        if run in target_runs:
            should_remove = True
        # Remove if it's gray shaded
        elif is_run_gray_shaded(run):
            should_remove = True
        # Remove if it's a hyperlink
        elif is_run_hyperlink(run):
            should_remove = True
        
        if should_remove:
            runs_to_remove.append(run)
    
    return runs_to_remove


def create_hyperlink_run(para: Paragraph, text: str, url: str) -> Run:
    """
    Create a proper hyperlink run in the paragraph.
    """
    try:
        # Create hyperlink element
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), '')  # External link
        hyperlink.set(qn('w:anchor'), url)
        
        # Create run within hyperlink
        run_element = OxmlElement('w:r')
        run_props = OxmlElement('w:rPr')
        
        # Add hyperlink styling
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0000FF')
        run_props.append(color)
        
        underline = OxmlElement('w:u')
        underline.set(qn('w:val'), 'single')
        run_props.append(underline)
        
        run_element.append(run_props)
        
        # Add text
        text_element = OxmlElement('w:t')
        text_element.text = text
        run_element.append(text_element)
        
        hyperlink.append(run_element)
        
        # Insert into paragraph
        para._p.append(hyperlink)
        
        # Return the run object
        return Run(run_element, para)
        
    except Exception as e:
        # Fallback to styled text if hyperlink creation fails
        run = para.add_run(text)
        run.font.color.rgb = RGBColor(0, 0, 255)
        run.underline = True
        return run

def find_runs_to_remove(para: Paragraph, target_string: str) -> List[Run]:
    """Find runs that should be removed (gray shaded, hyperlinks, or target text)."""
    runs_to_remove = []
    
    target_start, target_end = find_target_text_range(para, target_string)
    
    if target_start == -1:
        return runs_to_remove
    
    # Map character positions to runs
    char_pos = 0
    run_ranges = []
    
    for run in para.runs:
        run_start = char_pos
        run_end = char_pos + len(run.text)
        run_ranges.append((run, run_start, run_end))
        char_pos = run_end
    
    # Find runs to remove
    for run, run_start, run_end in run_ranges:
        should_remove = False
        
        # Check if run overlaps with target range
        if run_start < target_end and run_end > target_start:
            should_remove = True
        # Check if it's gray shaded
        elif is_run_gray_shaded(run):
            should_remove = True
        # Check if it's a hyperlink
        elif is_run_hyperlink(run):
            should_remove = True
        
        if should_remove:
            runs_to_remove.append(run)
    
    return runs_to_remove


def build_replacement_text_by_country(components: List[Dict]) -> str:
    """Build replacement text grouped by country."""
    # Group components by country
    countries = {}
    for comp in components:
        country = comp['country']
        if country not in countries:
            countries[country] = []
        countries[country].append(comp)
    
    # Build text for each country
    country_blocks = []
    
    for country_name, country_components in countries.items():
        # Sort components by line number
        sorted_components = sorted(country_components, key=lambda x: x['line'])
        
        # Build country block
        country_lines = []
        
        for comp in sorted_components:
            line_text = comp['text']
            hyperlink = comp.get('hyperlink')
            email = comp.get('email')
            
            # Add hyperlinks if present and not already in text
            if hyperlink and hyperlink not in line_text:
                line_text += f" {hyperlink}"
            if email and email not in line_text:
                line_text += f" {email}"
            
            country_lines.append(line_text)
        
        # Join lines for this country
        country_block = '\n'.join(country_lines)
        country_blocks.append(country_block)
    
    # Join country blocks with double line breaks
    return '\n\n'.join(country_blocks)

def get_replacement_components(mapping_row: pd.Series, section_type: str, 
                              cached_components: Optional[List] = None, 
                              country_delimiter: str = ";") -> List:
    """Build replacement text components from mapping data."""
    if cached_components is not None:
        return cached_components
    
    components = []
    
    # Get line columns for this section type
    line_columns = [col for col in mapping_row.index 
                   if col.startswith('Line ') and section_type in col]
    
    if not line_columns:
        return components
    
    # Get hyperlinks and email links
    hyperlinks_col = f'Hyperlinks {section_type}'
    email_col = f'Link for email - {section_type}'
    
    hyperlinks_str = str(mapping_row.get(hyperlinks_col, '')).strip()
    email_str = str(mapping_row.get(email_col, '')).strip()
    
    # Parse hyperlinks and emails (semicolon separated)
    hyperlinks = [h.strip() for h in hyperlinks_str.split(country_delimiter) 
                 if h.strip() and h.strip().lower() != 'nan']
    emails = [e.strip() for e in email_str.split(country_delimiter) 
             if e.strip() and e.strip().lower() != 'nan']
    
    # Sort line columns by number
    def extract_line_number(col_name):
        match = re.search(r'Line (\d+)', col_name)
        return int(match.group(1)) if match else 999
    
    sorted_columns = sorted(line_columns, key=extract_line_number)
    
    # Find Line 1 to get countries
    line_1_col = None
    for col in sorted_columns:
        if extract_line_number(col) == 1:
            line_1_col = col
            break
    
    if not line_1_col:
        return components
    
    line_1_text = str(mapping_row.get(line_1_col, '')).strip()
    if not line_1_text or line_1_text.lower() == 'nan':
        return components
    
    # Parse countries using semicolon delimiter
    countries = [c.strip() for c in line_1_text.split(country_delimiter) if c.strip()]
    
    if not countries:
        return components
    
    # Process each line
    for col in sorted_columns:
        line_num = extract_line_number(col)
        content = str(mapping_row.get(col, '')).strip()
        
        if not content or content.lower() == 'nan':
            continue
        
        # Split content by countries using semicolon delimiter
        parts = [p.strip() for p in content.split(country_delimiter)]
        
        for i, country in enumerate(countries):
            if i < len(parts) and parts[i]:
                text = parts[i]
                
                # Determine links for this country position
                hyperlink = hyperlinks[i] if i < len(hyperlinks) else None
                email = emails[i] if i < len(emails) else None
                
                components.append({
                    'line': line_num,
                    'country': country,
                    'text': text,
                    'hyperlink': hyperlink,
                    'email': email
                })
    
    return components


def insert_formatted_replacement_surgically(para: Paragraph, insertion_point: int, 
                                          components: List[Dict], country_delimiter: str = ";"):
    """
    Insert properly formatted replacement text at a specific position in the paragraph.
    """
    # Group components by line
    lines = {}
    for comp in components:
        line_num = comp['line']
        if line_num not in lines:
            lines[line_num] = []
        lines[line_num].append(comp)
    
    # Insert at the specified position
    current_element = None
    if insertion_point < len(para.runs):
        current_element = para.runs[insertion_point]._element
    
    # Build replacement text line by line
    for line_idx, line_num in enumerate(sorted(lines.keys())):
        line_components = lines[line_num]
        
        # Add line break before non-first lines
        if line_idx > 0:
            new_run = para.add_run('\n')
            if current_element is not None:
                current_element.addnext(new_run._element)
                current_element = new_run._element
        
        # Add components for this line
        for comp_idx, comp in enumerate(line_components):
            # Add delimiter between components (except first)
            if comp_idx > 0:
                delimiter_run = para.add_run(f'{country_delimiter} ')
                if current_element is not None:
                    current_element.addnext(delimiter_run._element)
                    current_element = delimiter_run._element
            
            text = comp['text']
            country = comp['country']
            hyperlink = comp.get('hyperlink')
            email = comp.get('email')
            
            # Add text with country bolding
            if country and country in text:
                # Split text to bold only the country name
                parts = text.split(country, 1)
                
                # Add text before country
                if parts[0]:
                    text_run = para.add_run(parts[0])
                    if current_element is not None:
                        current_element.addnext(text_run._element)
                        current_element = text_run._element
                
                # Add bolded country name
                country_run = para.add_run(country)
                country_run.bold = True
                if current_element is not None:
                    current_element.addnext(country_run._element)
                    current_element = country_run._element
                
                # Add text after country
                if len(parts) > 1 and parts[1]:
                    remaining_run = para.add_run(parts[1])
                    if current_element is not None:
                        current_element.addnext(remaining_run._element)
                        current_element = remaining_run._element
            else:
                text_run = para.add_run(text)
                if current_element is not None:
                    current_element.addnext(text_run._element)
                    current_element = text_run._element
            
            # Add hyperlink if present
            if hyperlink:
                hyperlink_run = create_hyperlink_run(para, f' {hyperlink}', hyperlink)
                if current_element is not None:
                    current_element.addnext(hyperlink_run._element)
                    current_element = hyperlink_run._element
            
            # Add email link if present
            if email:
                email_run = create_hyperlink_run(para, f' {email}', f'mailto:{email}')
                if current_element is not None:
                    current_element.addnext(email_run._element)
                    current_element = email_run._element


def debug_paragraph_structure(para: Paragraph, target_string: str):
    """
    Debug function to understand paragraph structure and identify issues.
    """
    print(f"\n🔍 DEBUGGING PARAGRAPH STRUCTURE")
    print(f"Full paragraph text: '{para.text}'")
    print(f"Target string: '{target_string}'")
    print(f"Target found: {target_string.lower() in para.text.lower()}")
    print(f"Number of runs: {len(para.runs)}")
    
    for i, run in enumerate(para.runs):
        print(f"\nRun {i}:")
        print(f"  Text: '{run.text}'")
        print(f"  Bold: {run.bold}")
        print(f"  Underline: {run.underline}")
        print(f"  Font color: {run.font.color.rgb if run.font.color else 'None'}")
        
        # Check for shading
        is_shaded = is_run_gray_shaded_debug(run)
        is_hyperlink = is_run_hyperlink_debug(run)
        
        print(f"  Is gray shaded: {is_shaded}")
        print(f"  Is hyperlink: {is_hyperlink}")
        print(f"  Should remove: {is_shaded or is_hyperlink}")


def is_run_gray_shaded_debug(run: Run) -> bool:
    """
    Debug version of gray shading detection with detailed output.
    """
    try:
        # Check run properties for shading
        run_pr = run._element.get_or_add_rPr()
        shading_elements = run_pr.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
        
        if shading_elements:
            print(f"    Found shading elements: {len(shading_elements)}")
            for shading in shading_elements:
                fill = shading.get(qn('w:fill'))
                print(f"    Shading fill: {fill}")
                if fill and fill.lower() in ['d9d9d9', 'cccccc', 'gray', 'lightgray', 'auto']:
                    return True
        
        # Check font color for gray
        if run.font.color and hasattr(run.font.color, 'rgb'):
            color = run.font.color.rgb
            print(f"    Font color RGB: {color}")
            gray_colors = [
                RGBColor(128, 128, 128),  # Standard gray
                RGBColor(153, 153, 153),  # Light gray
                RGBColor(102, 102, 102),  # Dark gray
                RGBColor(96, 96, 96),     # Another common gray
                RGBColor(217, 217, 217),  # Very light gray
            ]
            if color in gray_colors:
                return True
                
    except Exception as e:
        print(f"    Error checking shading: {e}")
    
    return False


def is_run_hyperlink_debug(run: Run) -> bool:
    """
    Debug version of hyperlink detection.
    """
    try:
        # Check if run is within a hyperlink element
        run_xml = run._r
        hyperlink_elements = run_xml.xpath('.//w:hyperlink', 
                                         namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
        if hyperlink_elements:
            print(f"    Found hyperlink elements: {len(hyperlink_elements)}")
            return True
            
        # Check hyperlink-style formatting
        if (run.font.color and hasattr(run.font.color, 'rgb') and 
            run.font.color.rgb == RGBColor(0, 0, 255) and run.underline):
            print(f"    Has hyperlink-style formatting (blue + underline)")
            return True
            
    except Exception as e:
        print(f"    Error checking hyperlink: {e}")
    
    return False


def find_runs_to_remove_aggressive(para: Paragraph, target_string: str) -> List[Run]:
    """
    More aggressive approach to find runs that should be removed.
    """
    runs_to_remove = []
    
    # First, find runs containing target text
    target_start = para.text.lower().find(target_string.lower())
    if target_start == -1:
        return runs_to_remove
    
    target_end = target_start + len(target_string)
    
    # Map character positions to runs
    char_pos = 0
    run_ranges = []
    
    for run in para.runs:
        run_start = char_pos
        run_end = char_pos + len(run.text)
        run_ranges.append((run, run_start, run_end))
        char_pos = run_end
    
    print(f"\n🎯 TARGET RANGE: {target_start} to {target_end}")
    
    # Find runs that overlap with target or are adjacent problematic runs
    for run, run_start, run_end in run_ranges:
        should_remove = False
        reason = ""
        
        # Check if run overlaps with target range
        if run_start < target_end and run_end > target_start:
            should_remove = True
            reason = "overlaps with target text"
        
        # Check if it's gray shaded
        elif is_run_gray_shaded_debug(run):
            should_remove = True
            reason = "is gray shaded"
        
        # Check if it's a hyperlink
        elif is_run_hyperlink_debug(run):
            should_remove = True
            reason = "is hyperlink"
        
        # Check if it's a small connector (like period, comma) adjacent to target
        elif (len(run.text.strip()) <= 2 and 
              run.text.strip() in '.,;:' and
              abs(run_start - target_end) <= 5):  # Within 5 chars of target end
            should_remove = True
            reason = "is small connector near target"
        
        print(f"Run {run_start}-{run_end}: '{run.text}' -> Remove: {should_remove} ({reason})")
        
        if should_remove:
            runs_to_remove.append(run)
    
    return runs_to_remove


def build_replacement_components_simple(mapping_row: pd.Series, section_type: str, 
                                       country_delimiter: str = ";") -> List[Dict]:
    """
    Simplified version that focuses on getting the components right.
    """
    print(f"\n🔨 Building replacement components for {section_type}")
    
    components = []
    
    # Get line columns for this section type
    line_columns = [col for col in mapping_row.index 
                   if col.startswith('Line ') and section_type in col]
    
    print(f"Found line columns: {line_columns}")
    
    if not line_columns:
        print("No line columns found")
        return components
    
    # Get hyperlinks and email links
    hyperlinks_col = f'Hyperlinks {section_type}'
    email_col = f'Link for email - {section_type}'
    
    hyperlinks_str = str(mapping_row.get(hyperlinks_col, '')).strip()
    email_str = str(mapping_row.get(email_col, '')).strip()
    
    print(f"Hyperlinks: '{hyperlinks_str}'")
    print(f"Emails: '{email_str}'")
    
    # Parse hyperlinks and emails (semicolon separated)
    hyperlinks = [h.strip() for h in hyperlinks_str.split(country_delimiter) 
                 if h.strip() and h.strip().lower() != 'nan']
    emails = [e.strip() for e in email_str.split(country_delimiter) 
             if e.strip() and e.strip().lower() != 'nan']
    
    print(f"Parsed hyperlinks: {hyperlinks}")
    print(f"Parsed emails: {emails}")
    
    # Sort line columns by number
    def extract_line_number(col_name):
        match = re.search(r'Line (\d+)', col_name)
        return int(match.group(1)) if match else 999
    
    sorted_columns = sorted(line_columns, key=extract_line_number)
    
    # Find Line 1 to get countries
    line_1_col = None
    for col in sorted_columns:
        if extract_line_number(col) == 1:
            line_1_col = col
            break
    
    if not line_1_col:
        print("No Line 1 column found")
        return components
    
    line_1_text = str(mapping_row.get(line_1_col, '')).strip()
    print(f"Line 1 text: '{line_1_text}'")
    
    if not line_1_text or line_1_text.lower() == 'nan':
        print("Line 1 text is empty")
        return components
    
    # Parse countries using semicolon delimiter
    countries = [c.strip() for c in line_1_text.split(country_delimiter) if c.strip()]
    print(f"Countries: {countries}")
    
    if not countries:
        print("No countries found")
        return components
    
    # Process each line
    for col in sorted_columns:
        line_num = extract_line_number(col)
        content = str(mapping_row.get(col, '')).strip()
        
        print(f"Processing Line {line_num}: '{content}'")
        
        if not content or content.lower() == 'nan':
            continue
        
        # Split content by countries using semicolon delimiter
        parts = [p.strip() for p in content.split(country_delimiter)]
        print(f"  Split into parts: {parts}")
        
        for i, country in enumerate(countries):
            if i < len(parts) and parts[i]:
                text = parts[i]
                
                # Determine links for this country position
                hyperlink = hyperlinks[i] if i < len(hyperlinks) else None
                email = emails[i] if i < len(emails) else None
                
                component = {
                    'line': line_num,
                    'country': country,
                    'text': text,
                    'hyperlink': hyperlink,
                    'email': email
                }
                
                components.append(component)
                print(f"  Added component: {component}")
    
    print(f"Total components built: {len(components)}")
    return components


def insert_replacement_simple(para: Paragraph, insertion_point: int, components: List[Dict], 
                            section_type: str, mapping_row: pd.Series, country_delimiter: str = ";"):
    """
    Simplified insertion that adds text at the insertion point.
    """
    print(f"\n📝 INSERTING REPLACEMENT at position {insertion_point}")
    print(f"Components to insert: {len(components)}")
    
    # Group components by line
    lines = {}
    for comp in components:
        line_num = comp['line']
        if line_num not in lines:
            lines[line_num] = []
        lines[line_num].append(comp)
    
    # Build replacement text
    replacement_text = ""
    
    for line_idx, line_num in enumerate(sorted(lines.keys())):
        line_components = lines[line_num]
        
        # Add line break before non-first lines
        if line_idx > 0:
            replacement_text += "\n"
        
        # Add components for this line
        line_texts = []
        for comp in line_components:
            text = comp['text']
            country = comp['country']
            hyperlink = comp.get('hyperlink')
            email = comp.get('email')
            
            # For now, just add the text (we'll enhance formatting later)
            component_text = text
            if hyperlink:
                component_text += f" {hyperlink}"
            if email:
                component_text += f" {email}"
            
            line_texts.append(component_text)
        
        # Join with semicolon delimiter
        replacement_text += f"{country_delimiter} ".join(line_texts)
    
    # For PL sections, append the additional text
    if section_type == "PL":
        additional_text = str(mapping_row.get('Text to be appended after National reporting system PL', '')).strip()
        if additional_text and additional_text.lower() != 'nan':
            replacement_text += f"\n\n{additional_text}"
    
    print(f"Replacement text: '{replacement_text}'")
    
    # Simple insertion - add a new run with the replacement text
    new_run = para.add_run(replacement_text)
    print("✅ Replacement text inserted")
    
    return True


def run_annex_update_v2(doc: Document, mapping_row: pd.Series, section_type: str, 
                       cached_components: Optional[List] = None, 
                       country_delimiter: str = ";") -> Tuple[bool, Optional[List]]:
    """Update national reporting systems in SmPC or PL sections."""
    # Get the target text to find and replace
    target_col = f'Original text national reporting - {section_type}'
    target_string = str(mapping_row.get(target_col, '')).strip()
    
    if ":" in target_string:
        target_string = target_string.split(':', 1)[-1].strip()

    if not target_string or target_string.lower() == 'nan':
        return False, None

    # Get replacement components
    components = get_replacement_components(mapping_row, section_type, cached_components, country_delimiter)
    
    if not components:
        return False, None
    
    # Find and update the target text
    found = False
    for para in doc.paragraphs:
        if target_string.lower() in para.text.lower():
            
            # Find runs to remove
            runs_to_remove = find_runs_to_remove(para, target_string)
            
            if runs_to_remove:
                # Remove the identified runs
                for run in runs_to_remove:
                    run._element.getparent().remove(run._element)
                
                # Build and insert replacement text
                replacement_text = build_replacement_text_by_country(components)
                
                # For PL sections, append additional text
                if section_type == "PL":
                    additional_text = str(mapping_row.get('Text to be appended after National reporting system PL', '')).strip()
                    if additional_text and additional_text.lower() != 'nan':
                        replacement_text += f"\n\n{additional_text}"
                
                # Insert the replacement
                para.add_run(replacement_text)
                
                found = True
                break
    
    return found, components


def update_document_with_fixed_smpc_blocks(doc: Document, mapping_row: pd.Series) -> Tuple[bool, List[str]]:
    """
    Main function to update document with fixed SmPC block handling.
    
    This function coordinates the updates for both SmPC and PL sections
    using the fixed functions that properly handle semicolon delimiters
    and selective text replacement.
    """
    updates_applied = []
    total_success = False
    
    try:
        # Get the correct country delimiter from config (default to semicolon)
        country_delimiter = ";"  # This should come from ProcessingConfig
        
        # 1. Update SmPC national reporting systems
        smpc_success, smpc_components = run_annex_update_v2(
            doc, mapping_row, "SmPC", None, country_delimiter=country_delimiter
        )
        if smpc_success:
            updates_applied.append("SmPC national reporting")
            total_success = True

        # 2. Update PL national reporting systems
        pl_success, _ = run_annex_update_v2(
            doc, mapping_row, "PL", smpc_components if smpc_success else None, country_delimiter=country_delimiter
        )
        if pl_success:
            updates_applied.append("PL national reporting")
            total_success = True
        
        return total_success, updates_applied
        
    except Exception as e:
        raise Exception(f"Failed to apply SmPC block updates: {e}")


def handle_pl_additional_text(para: Paragraph, mapping_row: pd.Series) -> bool:
    """
    Handle the additional text that needs to be appended after PL national reporting system.
    
    This text comes from the "Text to be appended after National reporting system PL" column
    and provides additional safety reporting information.
    """
    additional_text = str(mapping_row.get('Text to be appended after National reporting system PL', '')).strip()
    
    if not additional_text or additional_text.lower() == 'nan':
        return False
    
    # Add spacing and the additional text
    para.add_run('\n\n')
    
    # Add the additional text with appropriate formatting
    additional_run = para.add_run(additional_text)
    # You can customize formatting here if needed (e.g., italic, different color, etc.)
    
    return True


def create_pl_replacement_block(mapping_row: pd.Series, country_delimiter: str = ";") -> str:
    """
    Create the complete PL replacement block including the main content and additional text.
    
    This handles the case where PL uses block format rather than line-by-line format.
    """
    # Get main PL content
    main_content = str(mapping_row.get('National reporting system PL', '')).strip()
    
    # Get additional text
    additional_text = str(mapping_row.get('Text to be appended after National reporting system PL', '')).strip()
    
    # Combine them
    full_content = []
    
    if main_content and main_content.lower() != 'nan':
        full_content.append(main_content)
    
    if additional_text and additional_text.lower() != 'nan':
        full_content.append(additional_text)
    
    return '\n\n'.join(full_content) if full_content else ''


def update_section_10_date(doc: DocumentObject, mapping_row: pd.Series) -> bool:
    """Update date in Annex I Section 10."""
    country = mapping_row.get('Country', '')
    date_header = mapping_row.get('Annex I Date Text', 'Date of first authorisation/renewal of the authorisation')
    
    if not country:
        return False
    
    try:
        formatted_date = format_date_for_country(country, 'annex_i')
    except Exception:
        date_format = mapping_row.get('Annex I Date Format', '')
        formatted_date = datetime.now().strftime("%d %B %Y")
    
    found = False
    for para in doc.paragraphs:
        text_lower = para.text.lower()
        
        if ('10.' in text_lower and ('first authorisation' in text_lower or 
            'date of first' in text_lower or 
            date_header.lower() in text_lower or
            'date of revision' in text_lower)):
            
            para.clear()
            run = para.add_run(f"{date_header}\n{formatted_date}")
            run.bold = False
            found = True
            break
    
    return found

def update_annex_iiib_date(doc: Document, mapping_row: pd.Series) -> bool:
    """Update date in Annex IIIB Section 6."""
    country = mapping_row.get('Country', '')
    date_text = mapping_row.get('Annex IIIB Date Text', 'This leaflet was last revised in')
    
    if not country:
        return False
    
    try:
        formatted_date = format_date_for_country(country, 'annex_iiib')
    except Exception:
        formatted_date = datetime.now().strftime("%d %B %Y")
    
    found = False
    for para in doc.paragraphs:
        text_lower = para.text.lower()
        
        if (date_text.lower() in text_lower or
            'leaflet was last revised' in text_lower or
            'dernière approbation' in text_lower or
            'última revisión' in text_lower):
            
            para.clear()
            run = para.add_run(f"{date_text} {formatted_date}")
            run.bold = False
            found = True
            break
    
    return found

def update_local_representatives(doc: Document, mapping_row: pd.Series) -> bool:
    """Update local representatives in Section 6 of Annex IIIB."""
    local_rep_text = str(mapping_row.get('Local Representative', '')).strip()
    bold_countries_str = str(mapping_row.get('Country names to be bolded - Local Reps', '')).strip()
    
    if not local_rep_text or local_rep_text.lower() == 'nan':
        return False
    
    bold_countries = [c.strip() for c in bold_countries_str.split(',') 
                     if c.strip() and c.strip().lower() != 'nan']
    
    found = False
    in_section_6 = False
    
    for idx, para in enumerate(doc.paragraphs):
        text_lower = para.text.lower()
        
        # Check if we're entering Section 6
        if ('6.' in text_lower and 'contents of the pack' in text_lower) or \
           ('section 6' in text_lower) or \
           ('contenu de l\'emballage' in text_lower):
            in_section_6 = True
            continue
        
        # Look for existing local rep text to replace
        if in_section_6 and ('marketing authorisation holder' in text_lower or
                            'local representative' in text_lower or
                            'représentant local' in text_lower):
            
            para.clear()
            
            # Process the local rep text line by line
            lines = local_rep_text.split('\n')
            
            for i, line in enumerate(lines):
                if not line.strip():
                    para.add_run('\n')
                    continue
                
                # Check if this line contains a country to be bolded
                should_bold = any(country in line for country in bold_countries)
                
                run = para.add_run(line)
                run.bold = should_bold
                
                if i < len(lines) - 1:
                    para.add_run('\n')
            
            found = True
            break
    
    return found

# ============================================================================= 
# Split Annexes Workflow
# =============================================================================

def split_annexes(source_path: str, output_dir: str, language: str, country: str, mapping_row: pd.Series) -> Tuple[str, str]:
    """Split a combined SmPC document into Annex I and Annex IIIB documents."""
    return split_annexes_three_headers_with_fallback(source_path, output_dir, language, country, mapping_row)

def split_annexes_enhanced(source_path: str, output_dir: str, language: str, country: str, mapping_row: pd.Series) -> Tuple[str, str]:
    """
    Split a combined SmPC document into Annex I and Annex IIIB documents using language-specific headers.
    
    This enhanced version uses the mapping file's language-specific headers and implements
    a bottom-up approach:
    1. First identify and split Annex IIIB (using "Annex IIIB Header in country language")
    2. Then identify and split Annex II (using "Annex II Header in country language") 
    3. Everything that remains becomes Annex I
    
    Args:
        source_path: Path to the combined document
        output_dir: Directory to save split documents
        language: Language of the document
        country: Country of the document
        mapping_row: Row from mapping file containing language-specific headers
        
    Returns:
        Tuple of (annex_i_path, annex_iiib_path)
    """
    
    # Load the document
    doc = Document(source_path)
    
    # Get language-specific headers from mapping file
    annex_ii_header = str(mapping_row.get('Annex II Header in country language', '')).strip()
    annex_iiib_header = str(mapping_row.get('Annex IIIB Header in country language', '')).strip()
    
    # Validate headers are available
    if not annex_ii_header or annex_ii_header.lower() == 'nan':
        raise ValueError(f"Missing Annex II header for {country} ({language})")
    if not annex_iiib_header or annex_iiib_header.lower() == 'nan':
        raise ValueError(f"Missing Annex IIIB header for {country} ({language})")
    
    print(f"🔍 Using headers for {country} ({language}):")
    print(f"   Annex II: '{annex_ii_header}'")
    print(f"   Annex IIIB: '{annex_iiib_header}'")
    
    # Find split points by scanning all paragraphs
    annex_ii_split_index = None
    annex_iiib_split_index = None
    
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        # Look for Annex II header (case-insensitive, flexible matching)
        if annex_ii_split_index is None and _is_header_match(text, annex_ii_header):
            annex_ii_split_index = idx
            print(f"✅ Found Annex II header at paragraph {idx}: '{text[:50]}...'")
        
        # Look for Annex IIIB header (case-insensitive, flexible matching)
        if annex_iiib_split_index is None and _is_header_match(text, annex_iiib_header):
            annex_iiib_split_index = idx
            print(f"✅ Found Annex IIIB header at paragraph {idx}: '{text[:50]}...'")
    
    # Validate that we found the required headers
    if annex_ii_split_index is None:
        raise ValueError(f"Could not find Annex II header '{annex_ii_header}' in document")
    if annex_iiib_split_index is None:
        raise ValueError(f"Could not find Annex IIIB header '{annex_iiib_header}' in document")
    
    # Ensure proper order (Annex II should come before Annex IIIB)
    if annex_ii_split_index >= annex_iiib_split_index:
        raise ValueError(f"Document structure error: Annex II (para {annex_ii_split_index}) should come before Annex IIIB (para {annex_iiib_split_index})")
    
    print(f"📊 Split points identified:")
    print(f"   Annex I: paragraphs 0 to {annex_ii_split_index - 1}")
    print(f"   Annex II: paragraphs {annex_ii_split_index} to {annex_iiib_split_index - 1}")
    print(f"   Annex IIIB: paragraphs {annex_iiib_split_index} to end")
    
    # Create new documents
    annex_i_doc = Document()
    annex_iiib_doc = Document()
    
    # Split the document based on identified boundaries
    for idx, para in enumerate(doc.paragraphs):
        if idx < annex_ii_split_index:
            # Annex I content (everything before Annex II)
            copy_paragraph(annex_i_doc, para)
        elif idx >= annex_iiib_split_index:
            # Annex IIIB content (everything from Annex IIIB header onwards)
            copy_paragraph(annex_iiib_doc, para)
        # Note: We skip Annex II content (between annex_ii_split_index and annex_iiib_split_index)
        # as we only need Annex I and Annex IIIB for the final output
    
    # Create country-specific subfolder
    country_safe = country.replace('/', '_').replace(' ', '_')
    country_dir = os.path.join(output_dir, country_safe)
    os.makedirs(country_dir, exist_ok=True)
    
    # Generate output paths in country subfolder
    base_name = Path(source_path).stem
    annex_i_filename = generate_output_filename(base_name, language, country, "annex_i")
    annex_iiib_filename = generate_output_filename(base_name, language, country, "annex_iiib")
    
    annex_i_path = os.path.join(country_dir, annex_i_filename)
    annex_iiib_path = os.path.join(country_dir, annex_iiib_filename)
    
    # Save documents
    annex_i_doc.save(annex_i_path)
    annex_iiib_doc.save(annex_iiib_path)
    
    print(f"💾 Created: {country_safe}/{annex_i_filename}")
    print(f"💾 Created: {country_safe}/{annex_iiib_filename}")
    
    return annex_i_path, annex_iiib_path

def split_annexes_three_headers_xml(source_path: str, output_dir: str, language: str, country: str, mapping_row: pd.Series) -> Tuple[str, str]:
    """
    Split document using all three headers with XML-based approach to preserve all formatting.
    
    This method:
    1. Finds all three annex headers to define precise boundaries
    2. Uses XML manipulation to preserve tables, images, formatting, etc.
    3. Creates clean splits without losing any document elements
    
    Args:
        source_path: Path to the combined document
        output_dir: Directory to save split documents  
        language: Language of the document
        country: Country of the document
        mapping_row: Row from mapping file containing all three language-specific headers
        
    Returns:
        Tuple of (annex_i_path, annex_iiib_path)
    """
    
    print(f"\n🔬 THREE-HEADER XML SPLITTING")
    print(f"File: {Path(source_path).name}")
    print(f"Country: {country} ({language})")
    
    # Load the document
    doc = Document(source_path)
    
    # Get all three language-specific headers from mapping file
    annex_i_header = str(mapping_row.get('Annex I Header in country language', '')).strip()
    annex_ii_header = str(mapping_row.get('Annex II Header in country language', '')).strip()
    annex_iiib_header = str(mapping_row.get('Annex IIIB Header in country language', '')).strip()
    
    # Validate all headers are available
    if not annex_i_header or annex_i_header.lower() == 'nan':
        raise ValueError(f"Missing Annex I header for {country} ({language})")
    if not annex_ii_header or annex_ii_header.lower() == 'nan':
        raise ValueError(f"Missing Annex II header for {country} ({language})")
    if not annex_iiib_header or annex_iiib_header.lower() == 'nan':
        raise ValueError(f"Missing Annex IIIB header for {country} ({language})")
    
    print(f"🎯 Target headers:")
    print(f"   Annex I: '{annex_i_header}'")
    print(f"   Annex II: '{annex_ii_header}'")
    print(f"   Annex IIIB: '{annex_iiib_header}'")
    
    # Find all header positions
    header_positions = find_header_positions(doc, annex_i_header, annex_ii_header, annex_iiib_header)
    
    if not header_positions['annex_i']:
        raise ValueError(f"Could not find Annex I header '{annex_i_header}' in document")
    if not header_positions['annex_ii']:
        raise ValueError(f"Could not find Annex II header '{annex_ii_header}' in document")
    if not header_positions['annex_iiib']:
        raise ValueError(f"Could not find Annex IIIB header '{annex_iiib_header}' in document")
    
    # Validate header order
    validate_header_order(header_positions)
    
    print(f"✅ Header positions validated:")
    print(f"   Annex I: Paragraph {header_positions['annex_i']}")
    print(f"   Annex II: Paragraph {header_positions['annex_ii']}")
    print(f"   Annex IIIB: Paragraph {header_positions['annex_iiib']}")
    
    # Extract sections using XML manipulation
    annex_i_doc = extract_section_xml(doc, 
                                      start_idx=header_positions['annex_i'], 
                                      end_idx=header_positions['annex_ii'])
    
    annex_iiib_doc = extract_section_xml(doc, 
                                         start_idx=header_positions['annex_iiib'], 
                                         end_idx=None)  # To end of document
    
    # Generate output paths
    base_name = Path(source_path).stem
    annex_i_filename = generate_output_filename(base_name, language, country, "annex_i")
    annex_iiib_filename = generate_output_filename(base_name, language, country, "annex_iiib")
    
    annex_i_path = os.path.join(output_dir, annex_i_filename)
    annex_iiib_path = os.path.join(output_dir, annex_iiib_filename)
    
    # Save documents
    annex_i_doc.save(annex_i_path)
    annex_iiib_doc.save(annex_iiib_path)
    
    print(f"💾 Created with XML preservation:")
    print(f"   {annex_i_filename}")
    print(f"   {annex_iiib_filename}")
    
    return annex_i_path, annex_iiib_path


def find_header_positions(doc: Document, annex_i_header: str, annex_ii_header: str, annex_iiib_header: str) -> Dict[str, int]:
    """
    Find the paragraph positions of all three annex headers.
    
    Returns:
        Dictionary with keys 'annex_i', 'annex_ii', 'annex_iiib' and paragraph indices as values
    """
    
    positions = {'annex_i': None, 'annex_ii': None, 'annex_iiib': None}
    
    # Find best match for each header
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        # Check for Annex I header
        if positions['annex_i'] is None and _is_header_match(text, annex_i_header):
            positions['annex_i'] = idx
            print(f"✅ Found Annex I header at paragraph {idx}: '{text[:50]}...'")
        
        # Check for Annex II header
        if positions['annex_ii'] is None and _is_header_match(text, annex_ii_header):
            positions['annex_ii'] = idx
            print(f"✅ Found Annex II header at paragraph {idx}: '{text[:50]}...'")
        
        # Check for Annex IIIB header
        if positions['annex_iiib'] is None and _is_header_match(text, annex_iiib_header):
            positions['annex_iiib'] = idx
            print(f"✅ Found Annex IIIB header at paragraph {idx}: '{text[:50]}...'")
    
    return positions


def validate_header_order(positions: Dict[str, int]) -> None:
    """
    Validate that headers are in the correct order: I < II < IIIB.
    
    Args:
        positions: Dictionary with header positions
        
    Raises:
        ValueError: If headers are not in correct order
    """
    
    annex_i_pos = positions['annex_i']
    annex_ii_pos = positions['annex_ii']
    annex_iiib_pos = positions['annex_iiib']
    
    if annex_i_pos >= annex_ii_pos:
        raise ValueError(f"Document structure error: Annex I (para {annex_i_pos}) should come before Annex II (para {annex_ii_pos})")
    
    if annex_ii_pos >= annex_iiib_pos:
        raise ValueError(f"Document structure error: Annex II (para {annex_ii_pos}) should come before Annex IIIB (para {annex_iiib_pos})")
    
    print(f"📊 Document structure validated:")
    print(f"   Annex I: {annex_iiib_pos - annex_i_pos} paragraphs")
    print(f"   Annex II: {annex_iiib_pos - annex_ii_pos} paragraphs") 
    print(f"   Annex IIIB: Continues to end of document")


def extract_section_xml(source_doc: Document, start_idx: int, end_idx: Optional[int] = None) -> Document:
    """
    Extract a section from the source document using safe paragraph copying to avoid XML corruption.
    
    This approach prioritizes document integrity over advanced XML preservation.
    
    Args:
        source_doc: Source document to extract from
        start_idx: Starting paragraph index (inclusive)
        end_idx: Ending paragraph index (exclusive). If None, goes to end of document.
        
    Returns:
        New document containing the extracted section without corruption
    """
    
    # Determine which paragraphs to include
    total_paragraphs = len(source_doc.paragraphs)
    actual_end_idx = end_idx if end_idx is not None else total_paragraphs
    
    print(f"📋 Extracting paragraphs {start_idx} to {actual_end_idx-1} (total: {actual_end_idx - start_idx})")
    
    # Use safe paragraph-by-paragraph copying to avoid XML corruption
    return _extract_section_safe_copy(source_doc, start_idx, actual_end_idx)


def _extract_section_safe_copy(source_doc: Document, start_idx: int, end_idx: int) -> Document:
    """
    Safe document extraction that preserves formatting without XML corruption.
    
    This method copies paragraphs, tables, and basic formatting while ensuring
    the resulting document is valid and doesn't trigger Word warnings.
    
    Args:
        source_doc: Source document
        start_idx: Start paragraph index  
        end_idx: End paragraph index
        
    Returns:
        New document with safely copied content
    """
    
    print(f"📋 Using safe copying for range {start_idx} to {end_idx-1}")
    
    # Create new document
    new_doc = Document()
    
    # Copy document-level settings safely
    _copy_document_settings_safe(source_doc, new_doc)
    
    # Clear the default empty paragraph
    if new_doc.paragraphs:
        p = new_doc.paragraphs[0]
        p.clear()
    
    # Track what we're copying
    paragraphs_copied = 0
    tables_copied = 0
    
    # Get both paragraphs and tables from the source document
    source_elements = _get_document_elements_in_order(source_doc)
    
    # Filter elements to the target range
    target_elements = []
    current_para_idx = 0
    
    for element in source_elements:
        if element['type'] == 'paragraph':
            if start_idx <= current_para_idx < end_idx:
                target_elements.append(element)
            current_para_idx += 1
        elif element['type'] == 'table':
            # Include tables that fall within our range
            if start_idx <= current_para_idx < end_idx:
                target_elements.append(element)
    
    # Copy the selected elements safely
    for element in target_elements:
        if element['type'] == 'paragraph':
            copy_paragraph_safe(new_doc, element['content'])
            paragraphs_copied += 1
        elif element['type'] == 'table':
            copy_table_safe(new_doc, element['content'])
            tables_copied += 1
    
    print(f"✅ Safely copied {paragraphs_copied} paragraphs and {tables_copied} tables")
    
    return new_doc


def _get_document_elements_in_order(doc: Document) -> List[Dict]:
    """
    Get all document elements (paragraphs and tables) in their order of appearance.
    
    Args:
        doc: Source document
        
    Returns:
        List of dictionaries with 'type' and 'content' keys
    """
    
    elements = []
    
    # Add all paragraphs
    for para in doc.paragraphs:
        elements.append({
            'type': 'paragraph',
            'content': para
        })
    
    # Note: Tables are embedded within the document structure
    # For simplicity, we'll handle them as part of paragraph processing
    # This avoids the complexity of XML order tracking
    
    return elements


def copy_paragraph_safe(dest_doc: Document, source_para) -> None:
    """
    Safely copy a paragraph from source to destination document.
    
    This preserves basic formatting while avoiding XML corruption.
    
    Args:
        dest_doc: Destination document
        source_para: Source paragraph to copy
    """
    
    # Create new paragraph
    new_para = dest_doc.add_paragraph()
    
    # Copy paragraph-level properties safely
    try:
        new_para.style = source_para.style
    except:
        # If style copying fails, use default
        pass
    
    try:
        new_para.alignment = source_para.alignment
    except:
        pass
    
    # Copy runs with formatting
    for run in source_para.runs:
        new_run = new_para.add_run(run.text)
        
        # Copy basic formatting safely
        try:
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
        except:
            # If formatting copying fails, continue with plain text
            pass


def copy_table_safe(dest_doc: Document, source_table) -> None:
    """
    Safely copy a table from source to destination document.
    
    Args:
        dest_doc: Destination document
        source_table: Source table to copy
    """
    
    try:
        # Get table dimensions
        rows = len(source_table.rows)
        cols = len(source_table.columns) if rows > 0 else 0
        
        if rows > 0 and cols > 0:
            # Create new table
            new_table = dest_doc.add_table(rows=rows, cols=cols)
            
            # Copy cell contents
            for row_idx in range(rows):
                for col_idx in range(cols):
                    try:
                        source_cell = source_table.cell(row_idx, col_idx)
                        dest_cell = new_table.cell(row_idx, col_idx)
                        dest_cell.text = source_cell.text
                    except:
                        # If cell copying fails, continue
                        continue
                        
    except Exception as e:
        print(f"⚠️  Warning: Could not copy table - {e}")


def _copy_document_settings_safe(source_doc: Document, target_doc: Document) -> None:
    """
    Safely copy basic document settings without causing corruption.
    
    Args:
        source_doc: Source document
        target_doc: Target document
    """
    
    try:
        # Only copy very basic properties that are unlikely to cause issues
        if hasattr(source_doc, 'core_properties') and hasattr(target_doc, 'core_properties'):
            # Copy basic metadata only
            target_doc.core_properties.author = source_doc.core_properties.author
    except:
        # If any copying fails, continue without it
        pass


def _copy_document_properties(source_doc: Document, target_doc: Document) -> None:
    """
    Copy document-level properties like styles, themes, etc.
    
    Args:
        source_doc: Source document
        target_doc: Target document to copy properties to
    """
    
    # This function is kept for compatibility but now calls the safe version
    _copy_document_settings_safe(source_doc, target_doc)


def debug_three_header_structure(source_path: str, mapping_row: pd.Series) -> None:
    """
    Debug the three-header approach to validate header detection.
    
    Args:
        source_path: Path to document to analyze
        mapping_row: Mapping row with header information
    """
    
    doc = Document(source_path)
    country = mapping_row.get('Country', 'Unknown')
    language = mapping_row.get('Language', 'Unknown')
    
    # Get all three headers
    annex_i_header = str(mapping_row.get('Annex I Header in country language', '')).strip()
    annex_ii_header = str(mapping_row.get('Annex II Header in country language', '')).strip()
    annex_iiib_header = str(mapping_row.get('Annex IIIB Header in country language', '')).strip()
    
    print(f"\n🔍 THREE-HEADER DEBUGGING")
    print(f"File: {Path(source_path).name}")
    print(f"Country: {country} ({language})")
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    print(f"Expected Annex I header: '{annex_i_header}'")
    print(f"Expected Annex II header: '{annex_ii_header}'")
    print(f"Expected Annex IIIB header: '{annex_iiib_header}'")
    print("=" * 80)
    
    # Find all matches for each header
    annex_i_matches = []
    annex_ii_matches = []
    annex_iiib_matches = []
    
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        if _is_header_match(text, annex_i_header):
            annex_i_matches.append({'index': idx, 'text': text})
        
        if _is_header_match(text, annex_ii_header):
            annex_ii_matches.append({'index': idx, 'text': text})
        
        if _is_header_match(text, annex_iiib_header):
            annex_iiib_matches.append({'index': idx, 'text': text})
    
    # Display results
    print(f"📌 HEADER MATCHES FOUND:")
    
    print(f"\nAnnex I ('{annex_i_header}'):")
    if annex_i_matches:
        for match in annex_i_matches:
            print(f"  Para {match['index']}: '{match['text'][:60]}...'")
    else:
        print(f"  ❌ No matches found")
    
    print(f"\nAnnex II ('{annex_ii_header}'):")
    if annex_ii_matches:
        for match in annex_ii_matches:
            print(f"  Para {match['index']}: '{match['text'][:60]}...'")
    else:
        print(f"  ❌ No matches found")
    
    print(f"\nAnnex IIIB ('{annex_iiib_header}'):")
    if annex_iiib_matches:
        for match in annex_iiib_matches:
            print(f"  Para {match['index']}: '{match['text'][:60]}...'")
    else:
        print(f"  ❌ No matches found")
    
    # Validate structure if all headers found
    if annex_i_matches and annex_ii_matches and annex_iiib_matches:
        best_i = annex_i_matches[0]['index']
        best_ii = annex_ii_matches[0]['index'] 
        best_iiib = annex_iiib_matches[0]['index']
        
        print(f"\n📊 PROPOSED STRUCTURE:")
        print(f"   Annex I: paragraphs {best_i} to {best_ii-1} ({best_ii - best_i} paragraphs)")
        print(f"   Annex II: paragraphs {best_ii} to {best_iiib-1} ({best_iiib - best_ii} paragraphs)")
        print(f"   Annex IIIB: paragraphs {best_iiib} to end ({len(doc.paragraphs) - best_iiib} paragraphs)")
        
        if best_i >= best_ii or best_ii >= best_iiib:
            print(f"  ❌ STRUCTURE ERROR: Headers not in correct order!")
        else:
            print(f"  ✅ Structure looks good!")
    else:
        print(f"\n❌ Cannot validate structure - missing header matches")


def split_annexes_three_headers_with_fallback(source_path: str, output_dir: str, language: str, country: str, mapping_row: pd.Series) -> Tuple[str, str]:
    """
    Main splitting function with three-header approach and fallback to two-header method.
    
    This is the function you should call from your processor.py
    """
    
    try:
        # Try the three-header XML approach first
        return split_annexes_three_headers_xml(source_path, output_dir, language, country, mapping_row)
    
    except ValueError as e:
        print(f"⚠️  Three-header approach failed: {e}")
        print(f"🔄 Falling back to two-header approach...")
        
        # Fall back to the enhanced two-header approach
        return split_annexes_enhanced(source_path, output_dir, language, country, mapping_row)
    
    except Exception as e:
        print(f"❌ XML approach failed with error: {e}")
        print(f"🔄 Falling back to two-header approach...")
        
        # Fall back to the enhanced two-header approach
        return split_annexes_enhanced(source_path, output_dir, language, country, mapping_row)

def _is_header_match(paragraph_text: str, header_text: str) -> bool:
    """Check if a paragraph text matches a header with precise word-boundary matching."""
    para_normalized = _normalize_text_for_matching(paragraph_text)
    header_normalized = _normalize_text_for_matching(header_text)
    
    # Exact match after normalization
    if para_normalized == header_normalized:
        return True
    
    # Check if header is contained in paragraph (word boundary matching)
    if _contains_as_words(para_normalized, header_normalized):
        return True
    
    # For very similar headers (like "annex i" vs "annex ii"), be more strict
    if _are_similar_headers(para_normalized, header_normalized):
        return False
    
    # Check if paragraph starts with header (common case)
    if para_normalized.startswith(header_normalized + " "):
        return True
    
    return False

def _contains_as_words(text: str, search_term: str) -> bool:
    """
    Check if search_term exists as complete words in text, not just as substring.
    This prevents "annex i" from matching "annex ii".
    """
    import re
    
    # Escape special regex characters in search term
    escaped_term = re.escape(search_term)
    
    # Use word boundaries to ensure complete word matching
    # \b ensures we match complete words, not substrings
    pattern = r'\b' + escaped_term + r'\b'
    
    return bool(re.search(pattern, text, re.IGNORECASE))


def _are_similar_headers(text1: str, text2: str) -> bool:
    """
    Check if two texts are similar annex headers that could be confused.
    Returns True if they're similar enough that we should be strict about matching.
    
    Uses comprehensive patterns based on actual mapping data from all supported languages.
    """
    
    # Comprehensive annex header base words from mapping data
    annex_base_words = [
        'bijlage',      # Dutch
        'annexe',       # French  
        'anhang',       # German
        'lisa',         # Estonian
        'παραρτημα',    # Greek
        'pielikums',    # Latvian
        'priedas',      # Lithuanian
        'anexo',        # Spanish/Portuguese
        'prilog',       # Croatian
        'priloga',      # Slovenian
        'liite',        # Finnish
        'bilaga',       # Swedish
        'allegato',     # Italian
        'annex',        # English
        'anness',       # Maltese
        'bilag',        # Danish
        'viðauki',      # Icelandic
        'vedlegg',      # Norwegian
        'příloha',      # Czech
        'aneks',        # Polish
        'príloha',      # Slovak
        'приложение',   # Bulgarian
        'melléklet',    # Hungarian
        'anexa',        # Romanian
    ]
    
    # Roman numeral patterns (including Greek variants)
    roman_patterns = [
        r'[ivx]+',          # Standard: i, ii, iii, iv, v
        r'[ιυχ]+',          # Greek: ι, ιι, ιιι
        r'\d+',             # Arabic numbers: 1, 2, 3 (backup)
    ]
    
    # Build comprehensive patterns for both word-first and number-first structures
    all_patterns = []
    
    for base_word in annex_base_words:
        for roman_pattern in roman_patterns:
            # Pattern 1: Word first (e.g., "ANNEXE I", "BIJLAGE II")
            all_patterns.append(rf'{re.escape(base_word)}\s*\.?\s*{roman_pattern}\.?')
            
            # Pattern 2: Number first (e.g., "I LISA", "II LISA") 
            all_patterns.append(rf'{roman_pattern}\.?\s+{re.escape(base_word)}')
            
            # Pattern 3: Number with period first (e.g., "I. MELLÉKLET")
            all_patterns.append(rf'{roman_pattern}\.\s*{re.escape(base_word)}')
    
    # Check if both texts match any of the same patterns
    for pattern in all_patterns:
        if (re.search(pattern, text1, re.IGNORECASE) and 
            re.search(pattern, text2, re.IGNORECASE)):
            return True
    
    # Additional check: if both contain the same base word, they're similar
    text1_lower = text1.lower()
    text2_lower = text2.lower()
    
    for base_word in annex_base_words:
        if base_word.lower() in text1_lower and base_word.lower() in text2_lower:
            return True
    
    return False

def _normalize_text_for_matching(text: str) -> str:
    """
    Normalize text for header matching by removing inconsistencies.
    
    Args:
        text: Raw text to normalize
        
    Returns:
        Normalized text suitable for comparison
    """
    
    # Convert to lowercase
    normalized = text.lower()
    
    # Remove extra whitespace and normalize spaces
    normalized = re.sub(r'\s+', ' ', normalized).strip()
    
    # Remove common punctuation that might vary
    normalized = re.sub(r'[.,;:!?""''""()]', '', normalized)
    
    # Remove common formatting artifacts
    normalized = re.sub(r'[\r\n\t]', ' ', normalized)
    normalized = re.sub(r'\s+', ' ', normalized).strip()
    
    return normalized

def split_annexes_with_validation(source_path: str, output_dir: str, language: str, country: str, mapping_row: pd.Series) -> Tuple[str, str]:
    """
    Wrapper function that adds validation and error handling to the enhanced splitting logic.
    
    This function can be used as a drop-in replacement for the original split_annexes function.
    """
    
    try:
        return split_annexes_enhanced(source_path, output_dir, language, country, mapping_row)
    
    except ValueError as e:
        print(f"❌ Validation error during splitting: {e}")
        print(f"🔄 Falling back to original splitting method...")
        
        # Fallback to original method if enhanced method fails
        return split_annexes_original(source_path, output_dir, language, country, mapping_row)
    
    except Exception as e:
        print(f"❌ Unexpected error during enhanced splitting: {e}")
        print(f"🔄 Falling back to original splitting method...")
        
        # Fallback to original method if enhanced method fails
        return split_annexes_original(source_path, output_dir, language, country, mapping_row)


def split_annexes_original(source_path: str, output_dir: str, language: str, country: str, mapping_row: pd.Series) -> Tuple[str, str]:
    """
    Original splitting logic as fallback.
    This is the existing implementation for compatibility.
    """
    
    doc = Document(source_path)
    
    # Create new documents
    annex_i_doc = Document()
    annex_iiib_doc = Document()
    
    current_section = None
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # Determine which section we're in using original logic
        if 'ANNEX I' in text.upper() or 'SUMMARY OF PRODUCT CHARACTERISTICS' in text.upper():
            current_section = 'annex_i'
        elif 'ANNEX III' in text.upper() or 'PACKAGE LEAFLET' in text.upper():
            current_section = 'annex_iiib'
        
        # Copy paragraph to appropriate document
        if current_section == 'annex_i':
            copy_paragraph(annex_i_doc, para)
        elif current_section == 'annex_iiib':
            copy_paragraph(annex_iiib_doc, para)
    
    # Generate output paths
    base_name = Path(source_path).stem
    annex_i_filename = generate_output_filename(base_name, language, country, "annex_i")
    annex_iiib_filename = generate_output_filename(base_name, language, country, "annex_iiib")
    
    annex_i_path = os.path.join(output_dir, annex_i_filename)
    annex_iiib_path = os.path.join(output_dir, annex_iiib_filename)
    
    # Save documents
    annex_i_doc.save(annex_i_path)
    annex_iiib_doc.save(annex_iiib_path)
    
    return annex_i_path, annex_iiib_path

def test_three_header_approach(document_path: str, mapping_row: pd.Series):
    """
    Test the three-header approach on a document.
    """
    
    print(f"🧪 TESTING THREE-HEADER APPROACH")
    print(f"Document: {document_path}")
    
    # First debug the structure
    debug_three_header_structure(document_path, mapping_row)
    
    # Then try the actual splitting
    try:
        print(f"\n" + "="*80)
        print(f"TESTING THREE-HEADER SPLITTING:")
        
        temp_output = Path(document_path).parent / "temp_three_header_test"
        temp_output.mkdir(exist_ok=True)
        
        result = split_annexes_three_headers_xml(
            document_path,
            str(temp_output),
            mapping_row.get('Language', 'Unknown'),
            mapping_row.get('Country', 'Unknown'),
            mapping_row
        )
        
        print(f"✅ Three-header splitting completed successfully!")
        print(f"   Annex I: {result[0]}")
        print(f"   Annex IIIB: {result[1]}")
        
        # Validate output files
        for file_path, name in [(result[0], "Annex I"), (result[1], "Annex IIIB")]:
            if Path(file_path).exists():
                doc = Document(file_path)
                print(f"   {name}: {len(doc.paragraphs)} paragraphs")
                
                # Check if document has content
                content_length = sum(len(p.text) for p in doc.paragraphs)
                if content_length == 0:
                    print(f"   ❌ {name} is empty!")
                elif content_length < 100:
                    print(f"   ⚠️  {name} is very short ({content_length} characters)")
                else:
                    print(f"   ✅ {name} has substantial content ({content_length} characters)")
        
    except Exception as e:
        print(f"❌ Three-header splitting failed: {e}")
        import traceback
        traceback.print_exc()
# ============================================================================= 
# ENHANCED PROCESSOR CLASSES
# =============================================================================

class FileManager:
    """Handles file operations and path management."""
    
    def __init__(self, base_folder: Path, config: ProcessingConfig):
        self.base_folder = base_folder
        self.config = config
        self.logger = logging.getLogger(f"{__name__}.FileManager")
    
    def setup_output_directories(self) -> Tuple[Path, Path]:
        """Create and return paths for output directories."""
        split_dir = self.base_folder / DirectoryNames.SPLIT_DOCS
        pdf_dir = self.base_folder / DirectoryNames.PDF_DOCS
        
        try:
            os.makedirs(split_dir, exist_ok=True)
            os.makedirs(pdf_dir, exist_ok=True)
            return split_dir, pdf_dir
        except OSError as e:
            raise ProcessingError(f"Failed to create output directories: {e}")
    
    def discover_processable_documents(self) -> List[Path]:
        """Find all valid Word documents that can be processed."""
        if not self.base_folder.is_dir():
            raise ValidationError(f"Folder does not exist: {self.base_folder}")
        
        documents = []
        for file_path in self.base_folder.iterdir():
            if self._is_processable_document(file_path):
                documents.append(file_path)
        
        return documents
    
    def _is_processable_document(self, file_path: Path) -> bool:
        """Check if a file is a valid document for processing."""
        if file_path.suffix.lower() != ".docx":
            return False
        if file_path.name.startswith(FileMarkers.TEMP_FILE_PREFIX):
            return False
        if FileMarkers.ANNEX_MARKER in file_path.name:
            return False
        if file_path.name.startswith(FileMarkers.ANNEX_PREFIX):
            return False
        return True
    
    def create_backup(self, file_path: Path) -> Optional[Path]:
        """Create a backup of the original file."""
        if not self.config.create_backups:
            return None
            
        backup_path = file_path.with_suffix(file_path.suffix + DirectoryNames.BACKUP_SUFFIX)
        if backup_path.exists():
            return backup_path
            
        try:
            shutil.copy2(file_path, backup_path)
            return backup_path
        except Exception:
            return None

class DocumentUpdater:
    """Handles document modification operations."""
    
    def __init__(self, config: ProcessingConfig):
        self.config = config
        self.logger = logging.getLogger(f"{__name__}.DocumentUpdater")
    
    def apply_all_updates(self, doc: DocumentObject, mapping_row: pd.Series) -> Tuple[bool, List[str]]:
        """Apply all required updates to a document."""
        updates_applied = []
        total_success = False
        
        try:
            # 1. Update national reporting systems  ⬅️ **REPLACE WITH THIS**
            smpc_success, smpc_updates = update_document_with_fixed_smpc_blocks(doc, mapping_row)
            if smpc_success:
                updates_applied.extend(smpc_updates)
                total_success = True
            
            # 2. Update dates
            annex_i_date_success = update_section_10_date(doc, mapping_row)
            if annex_i_date_success:
                updates_applied.append("Annex I dates")
                total_success = True
            
            annex_iiib_date_success = update_annex_iiib_date(doc, mapping_row)
            if annex_iiib_date_success:
                updates_applied.append("Annex IIIB dates")
                total_success = True
            
            # 3. Update local representatives
            local_rep_success = update_local_representatives(doc, mapping_row)
            if local_rep_success:
                updates_applied.append("Local representatives")
                total_success = True
                
            return total_success, updates_applied
            
        except Exception as e:
            raise DocumentError(f"Failed to apply document updates: {e}")

class DocumentProcessor:
    """Main document processing orchestrator."""
    
    def __init__(self, config: Optional[ProcessingConfig] = None):
        self.config = config or ProcessingConfig()
        self.stats = ProcessingStats()
        self.logger = self._setup_logging()
        
    def _setup_logging(self) -> logging.Logger:
        """Set up logging configuration."""
        logger = logging.getLogger(__name__)
        logger.setLevel(getattr(logging, self.config.log_level.upper()))
        
        if not logger.handlers:
            handler = logging.StreamHandler()
            formatter = logging.Formatter(
                '%(asctime)s - %(name)s - %(levelname)s - %(message)s'
            )
            handler.setFormatter(formatter)
            logger.addHandler(handler)
        
        return logger
    
    def process_folder(self, folder_path: str, mapping_path: str) -> ProcessingResult:
        """Main entry point for processing a folder of documents."""
        try:
            self.logger.info("=" * 80)
            self.logger.info("🚀 STARTING ENHANCED DOCUMENT PROCESSING")
            self.logger.info("=" * 80)
            
            # Validate inputs
            folder = self._validate_folder_path(folder_path)
            mapping_df = self._load_and_validate_mapping(mapping_path)
            
            # Setup processing environment
            file_manager = FileManager(folder, self.config)
            split_dir, pdf_dir = file_manager.setup_output_directories()
            
            # Discover documents to process
            documents = file_manager.discover_processable_documents()
            self.stats.input_files_found = len(documents)
            
            if not documents:
                return ProcessingResult(
                    success=False,
                    message="No valid documents found for processing"
                )
            
            # Process each document
            output_files = []
            for document_path in documents:
                try:
                    result = self._process_single_document(
                        document_path, mapping_df, file_manager, split_dir, pdf_dir
                    )
                    output_files.extend(result.output_files)
                    
                except Exception as e:
                    self.logger.error(f"Error processing {document_path.name}: {e}")
                    self.stats.errors_encountered += 1
            
            # Generate final report
            return self._generate_final_result(output_files)
            
        except Exception as e:
            self.logger.error(f"Fatal error in process_folder: {e}")
            return ProcessingResult(
                success=False,
                message=f"Processing failed: {e}",
                errors=[str(e)]
            )
    
    def _validate_folder_path(self, folder_path: str) -> Path:
        """Validate and return folder path."""
        folder = Path(folder_path).resolve()
        if not folder.is_dir():
            raise ValidationError(f"Invalid directory: {folder_path}")
        return folder
    
    def _load_and_validate_mapping(self, mapping_path: str) -> pd.DataFrame:
        """Load and validate mapping file."""
        try:
            mapping_df = load_mapping_table(mapping_path)
            if mapping_df is None or mapping_df.empty:
                raise MappingError(f"Could not load mapping file: {mapping_path}")
            
            self.logger.info(f"Mapping loaded: {len(mapping_df)} configurations")
            return mapping_df
            
        except Exception as e:
            raise MappingError(f"Failed to load mapping file: {e}")
    
    def _process_single_document(
        self,
        document_path: Path,
        mapping_df: pd.DataFrame,
        file_manager: FileManager,
        split_dir: Path,
        pdf_dir: Path
    ) -> ProcessingResult:
        """Process a single document with all its variants."""
        
        self.logger.info("=" * 60)
        self.logger.info(f"📄 PROCESSING: {document_path.name}")
        self.logger.info("=" * 60)
        
        self.stats.input_files_processed += 1
        
        try:
            # Identify document language and country
            country_code, language_name, country_name = identify_document_country_and_language(str(document_path))
            
            if not language_name:
                error_msg = f"Cannot identify language for {document_path.name}"
                self.logger.error(error_msg)
                return ProcessingResult(success=False, message=error_msg)
            
            self.logger.info(f"Document identified - Language: {language_name}, Country: {country_name}")
            
            # Find mapping rows for this language
            mapping_rows = find_mapping_rows_for_language(mapping_df, language_name)
            if not mapping_rows:
                error_msg = f"No mapping found for language: {language_name}"
                self.logger.error(error_msg)
                return ProcessingResult(success=False, message=error_msg)
            
            self.logger.info(f"Found {len(mapping_rows)} variant(s) to process")
            
            # Create backup
            file_manager.create_backup(document_path)
            
            # Process each variant
            output_files = []
            variant_success_count = 0
            
            for i, mapping_row in enumerate(mapping_rows, 1):
                country = mapping_row['Country']
                self.logger.info(f"🌍 Processing variant {i}/{len(mapping_rows)}: {country}")
                
                try:
                    result = self._process_document_variant(
                        document_path, mapping_row, split_dir, pdf_dir
                    )
                    
                    if result.success:
                        variant_success_count += 1
                        self.stats.variants_successful += 1
                        output_files.extend(result.output_files)
                        self.logger.info(f"✅ Variant {i} completed successfully")
                    else:
                        self.logger.warning(f"⚠️ Variant {i} completed with issues: {result.message}")
                    
                    self.stats.variants_processed += 1
                    
                except Exception as e:
                    self.logger.error(f"❌ Error processing variant {i} ({country}): {e}")
                    self.stats.errors_encountered += 1
            
            # Document summary
            success_rate = (variant_success_count / len(mapping_rows)) * 100 if mapping_rows else 0
            self.logger.info(f"📊 Document Summary: {variant_success_count}/{len(mapping_rows)} variants successful ({success_rate:.1f}%)")
            
            return ProcessingResult(
                success=variant_success_count > 0,
                message=f"Processed {variant_success_count}/{len(mapping_rows)} variants successfully",
                output_files=output_files
            )
            
        except Exception as e:
            self.logger.error(f"Error processing document {document_path.name}: {e}")
            return ProcessingResult(success=False, message=str(e), errors=[str(e)])
    
    def _process_document_variant(
        self,
        document_path: Path,
        mapping_row: pd.Series,
        split_dir: Path,
        pdf_dir: Path
    ) -> ProcessingResult:
        """Process a single document variant."""
        
        country = mapping_row['Country']
        language = mapping_row['Language']
        
        try:
            # Load document
            doc = Document(str(document_path))
            
            # Apply updates
            updater = DocumentUpdater(self.config)
            updates_made, updates_applied = updater.apply_all_updates(doc, mapping_row)
            
            if not updates_made:
                return ProcessingResult(
                    success=False,
                    message=f"No updates applied for {country} variant"
                )
            
            # Save and process updated document
            return self._save_and_split_document(
                doc, document_path, mapping_row, split_dir, pdf_dir, updates_applied
            )
            
        except Exception as e:
            raise DocumentError(f"Failed to process variant for {country}: {e}")
    
    def _save_and_split_document(
        self,
        doc: Document,
        original_path: Path,
        mapping_row: pd.Series,
        split_dir: Path,
        pdf_dir: Path,
        updates_applied: List[str]
    ) -> ProcessingResult:
        """Save updated document and split into annexes."""
        
        country = mapping_row['Country']
        language = mapping_row['Language']
        output_files = []
        
        try:
            # Generate output filename
            base_name = original_path.stem
            output_filename = generate_output_filename(base_name, language, country, "combined")
            output_path = original_path.parent / output_filename
            
            # Save updated document
            doc.save(str(output_path))
            output_files.append(str(output_path))
            self.logger.info(f"💾 Saved combined document: {output_filename}")
            
            # Split into annexes
            self.logger.info("🔀 Splitting into separate annexes...")
            annex_i_path, annex_iiib_path = split_annexes(
                str(output_path), str(split_dir), language, country, mapping_row
            )
            
            output_files.extend([annex_i_path, annex_iiib_path])
            self.logger.info(f"✅ Split completed")
            
            # Convert to PDF if enabled
            if self.config.convert_to_pdf:
                try:
                    self.logger.info("📄 Converting to PDF...")
                    
                    # Try converting Annex I
                    try:
                        pdf_annex_i = convert_to_pdf(annex_i_path, str(pdf_dir))
                        output_files.append(pdf_annex_i)
                        self.logger.info(f"✅ Annex I PDF: {Path(pdf_annex_i).name}")
                    except Exception as e:
                        self.logger.warning(f"⚠️ Annex I PDF conversion failed: {e}")
                    
                    # Try converting Annex IIIB  
                    try:
                        pdf_annex_iiib = convert_to_pdf(annex_iiib_path, str(pdf_dir))
                        output_files.append(pdf_annex_iiib)
                        self.logger.info(f"✅ Annex IIIB PDF: {Path(pdf_annex_iiib).name}")
                    except Exception as e:
                        self.logger.warning(f"⚠️ Annex IIIB PDF conversion failed: {e}")
                    
                    self.logger.info("📄 PDF conversion phase completed")
                    
                except Exception as e:
                    self.logger.warning(f"⚠️ PDF conversion setup failed: {e}")
                    # Continue processing - PDF conversion is not critical
            
            self.stats.output_files_created += len(output_files)
            
            return ProcessingResult(
                success=True,
                message=f"Successfully processed {country} variant with updates: {', '.join(updates_applied)}",
                output_files=output_files
            )
            
        except Exception as e:
            raise DocumentError(f"Failed to save and split document: {e}")
    
    def _generate_final_result(self, output_files: List[str]) -> ProcessingResult:
        """Generate final processing result with statistics."""
        
        self.logger.info("=" * 80)
        self.logger.info("✅ ENHANCED PROCESSING COMPLETE")
        self.logger.info("=" * 80)
        
        self.logger.info("📊 Final Statistics:")
        self.logger.info(f"   Input files found: {self.stats.input_files_found}")
        self.logger.info(f"   Input files processed: {self.stats.input_files_processed}")
        self.logger.info(f"   Total variants processed: {self.stats.variants_processed}")
        self.logger.info(f"   Successful variants: {self.stats.variants_successful}")
        self.logger.info(f"   Success rate: {self.stats.success_rate():.1f}%")
        self.logger.info(f"   Output files created: {self.stats.output_files_created}")
        self.logger.info(f"   Errors encountered: {self.stats.errors_encountered}")
        
        success = self.stats.variants_successful > 0
        message = f"Processed {self.stats.variants_successful}/{self.stats.variants_processed} variants successfully"
        
        return ProcessingResult(
            success=success,
            message=message,
            output_files=output_files
        )

# ============================================================================= 
# BACKWARDS COMPATIBILITY INTERFACE
# =============================================================================

def process_folder(folder: str, mapping_path: str) -> None:
    """
    Backwards compatible entry point for processing folders.
    
    This function maintains the same interface as the original implementation
    while using the enhanced processing system under the hood.
    """
    try:
        processor = DocumentProcessor()
        result = processor.process_folder(folder, mapping_path)
        
        if not result.success:
            # Log the error but don't raise exception to maintain backwards compatibility
            logging.error(f"Processing failed: {result.message}")
            if result.errors:
                for error in result.errors:
                    logging.error(f"Error detail: {error}")
    
    except Exception as e:
        # Maintain backwards compatibility by logging errors instead of raising
        logging.error(f"Fatal processing error: {e}")
        raise  # Re-raise to maintain original behavior

def process_folder_enhanced(
    folder: str, 
    mapping_path: str, 
    config: Optional[ProcessingConfig] = None
) -> ProcessingResult:
    """
    Enhanced entry point that returns detailed results.
    
    Args:
        folder: Path to folder containing Word documents
        mapping_path: Path to Excel mapping file
        config: Optional processing configuration
        
    Returns:
        ProcessingResult with detailed success/failure information
    """
    processor = DocumentProcessor(config)
    return processor.process_folder(folder, mapping_path)
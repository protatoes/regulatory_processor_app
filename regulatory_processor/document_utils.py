"""Document processing utilities for manipulating Word documents."""

from typing import Dict, List, Optional, Tuple
import pandas as pd
from docx import Document
from docx.document import Document as DocumentObject
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

try:
    from .date_formatter import format_date_for_country
except ImportError:
    # Fallback for when module is imported directly
    from date_formatter import format_date_for_country


def copy_paragraph(dest_doc: Document, source_para: Paragraph) -> None:
    """Copy a paragraph from one document to another, preserving runs."""
    new_para = dest_doc.add_paragraph()
    new_para.style = source_para.style
    for run in source_para.runs:
        new_run = new_para.add_run(run.text)
        new_run.bold = run.bold
        new_run.italic = run.italic
        new_run.underline = run.underline
        new_run.style = run.style


def is_hex_gray_color(hex_color: str) -> bool:
    """Check if a hex color represents a gray shade."""
    if not hex_color:
        return False
    
    hex_color = hex_color.replace('#', '').upper()
    
    gray_hex_values = [
        'BFBFBF', 'CCCCCC', 'D9D9D9', '808080', '999999', 
        '666666', 'C0C0C0', 'A0A0A0'
    ]
    
    if hex_color in gray_hex_values:
        return True
    
    # Check if R=G=B (indicates gray)
    try:
        if len(hex_color) == 6:
            r = int(hex_color[0:2], 16)
            g = int(hex_color[2:4], 16)
            b = int(hex_color[4:6], 16)
            return r == g == b
    except ValueError:
        pass
    
    return False


def is_run_gray_shaded(run: Run) -> bool:
    """Check if a run has gray shading."""
    try:
        run_pr = run._element.get_or_add_rPr()
        shading_elements = run_pr.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
        
        if shading_elements:
            for shading in shading_elements:
                fill = shading.get(qn('w:fill'))
                if fill and is_hex_gray_color(fill):
                    return True
        
        # Check font color for gray
        if run.font.color and hasattr(run.font.color, 'rgb'):
            gray_colors = [
                RGBColor(128, 128, 128), RGBColor(153, 153, 153), 
                RGBColor(102, 102, 102), RGBColor(96, 96, 96),
                RGBColor(217, 217, 217), RGBColor(191, 191, 191)
            ]
            if run.font.color.rgb in gray_colors:
                return True
                
    except Exception:
        pass
    
    return False


def is_run_hyperlink(run: Run) -> bool:
    """Check if a run is part of a hyperlink."""
    try:
        run_xml = run._r
        hyperlink_elements = run_xml.xpath('.//w:hyperlink')
        if hyperlink_elements:
            return True
            
        # Check hyperlink-style formatting
        if (run.font.color and hasattr(run.font.color, 'rgb') and 
            run.font.color.rgb == RGBColor(0, 0, 255) and run.underline):
            return True
            
    except Exception:
        pass
    
    return False


def find_target_text_runs(para: Paragraph, target_string: str) -> List[Run]:
    """Find runs that contain the target text to be replaced."""
    target_runs = []
    
    # Build text map to find position
    char_pos = 0
    run_positions = []
    
    for run in para.runs:
        run_start = char_pos
        run_end = char_pos + len(run.text)
        run_positions.append((run, run_start, run_end))
        char_pos = run_end
    
    # Find target string position in full text
    full_text = para.text
    target_start = full_text.lower().find(target_string.lower())
    
    if target_start == -1:
        return []
    
    target_end = target_start + len(target_string)
    
    # Find runs that overlap with target text
    for run, run_start, run_end in run_positions:
        # Check if run overlaps with target range
        if (run_start < target_end and run_end > target_start):
            target_runs.append(run)
    
    return target_runs


def find_target_text_range(para: Paragraph, target_string: str) -> Tuple[int, int]:
    """Find the complete target text range in paragraph."""
    full_text = para.text.lower()
    target_lower = target_string.lower()
    
    # Try exact match first
    start_pos = full_text.find(target_lower)
    if start_pos != -1:
        return start_pos, start_pos + len(target_string)
    
    # Try to find key parts
    key_phrases = ['national reporting system', 'appendix v', 'listed in appendix']
    
    earliest_start = len(full_text)
    latest_end = 0
    
    for phrase in key_phrases:
        pos = full_text.find(phrase)
        if pos != -1:
            earliest_start = min(earliest_start, pos)
            latest_end = max(latest_end, pos + len(phrase))
    
    if earliest_start < len(full_text):
        return earliest_start, latest_end
    
    return -1, -1


def find_runs_to_remove(para: Paragraph, target_string: str) -> List[Run]:
    """Find runs that should be removed (gray shaded, hyperlinks, or target text)."""
    runs_to_remove = []
    
    target_start, target_end = find_target_text_range(para, target_string)
    
    if target_start == -1:
        return runs_to_remove
    
    # Map character positions to runs
    char_pos = 0
    run_ranges = []
    
    for run in para.runs:
        run_start = char_pos
        run_end = char_pos + len(run.text)
        run_ranges.append((run, run_start, run_end))
        char_pos = run_end
    
    # Find runs to remove
    for run, run_start, run_end in run_ranges:
        should_remove = False
        
        # Check if run overlaps with target range
        if run_start < target_end and run_end > target_start:
            should_remove = True
        # Check if it's gray shaded
        elif is_run_gray_shaded(run):
            should_remove = True
        # Check if it's a hyperlink
        elif is_run_hyperlink(run):
            should_remove = True
        
        if should_remove:
            runs_to_remove.append(run)
    
    return runs_to_remove


def find_runs_to_remove_original_text(para: Paragraph, original_text: str) -> List[Run]:
    """
    Find runs that contain the specific original text to be replaced.

    Handles the realistic scenario where original text spans multiple formatted runs:
    - Gray shaded text + hyperlinks + more gray shaded text
    - Text is broken across runs due to formatting differences
    """
    runs_to_remove = []

    if not original_text or str(original_text).lower() in ['nan', '']:
        return runs_to_remove

    # Step 1: Identify all gray-shaded and hyperlink runs in the paragraph
    formatted_runs = []
    for run in para.runs:
        if is_run_gray_shaded(run) or is_run_hyperlink(run):
            formatted_runs.append(run)

    if not formatted_runs:
        return runs_to_remove

    # Step 2: Find contiguous sequences of formatted runs
    run_sequences = []
    current_sequence = []

    for i, run in enumerate(para.runs):
        if run in formatted_runs:
            current_sequence.append(run)
        else:
            # Gap in formatted runs - save current sequence if it exists
            if current_sequence:
                run_sequences.append(current_sequence.copy())
                current_sequence = []

    # Don't forget the last sequence
    if current_sequence:
        run_sequences.append(current_sequence)

    # Step 3: Smart matching - handle both contiguous sequences AND distributed runs
    original_lower = original_text.lower().strip()

    # First try: Check contiguous sequences
    for sequence in run_sequences:
        # Reconstruct text from this run sequence
        sequence_text = ''.join(run.text for run in sequence).lower().strip()

        # Check for match (allowing for some whitespace differences)
        if original_lower in sequence_text or sequence_text in original_lower:
            runs_to_remove.extend(sequence)
            return runs_to_remove  # Found our target

        # Also check if the original text is a substantial subset
        if len(original_lower) > 10:  # Only for substantial text
            # Remove extra whitespace for comparison
            import re
            clean_original = re.sub(r'\s+', ' ', original_lower)
            clean_sequence = re.sub(r'\s+', ' ', sequence_text)

            # Check if they're substantially similar (80% match)
            if clean_original in clean_sequence or clean_sequence in clean_original:
                runs_to_remove.extend(sequence)
                return runs_to_remove

    # Second try: Distributed matching across all formatted runs
    # This handles cases where original text spans formatted + unformatted runs
    if formatted_runs:
        # Get all text from formatted runs and surrounding context
        para_text = para.text.lower()

        # Check if original text exists in paragraph and overlaps with formatted runs
        if original_lower in para_text:
            # Find the position of original text in paragraph
            start_pos = para_text.find(original_lower)
            end_pos = start_pos + len(original_text)

            # Map character positions to runs to find overlapping runs
            char_pos = 0
            for run in para.runs:
                run_start = char_pos
                run_end = char_pos + len(run.text)

                # If this run overlaps with original text AND is formatted, include it
                if (run_start < end_pos and run_end > start_pos and
                    (is_run_gray_shaded(run) or is_run_hyperlink(run))):
                    runs_to_remove.append(run)

                char_pos = run_end

    return runs_to_remove


def find_run_indices_to_remove_original_text(para: Paragraph, original_text: str) -> List[int]:
    """
    Find run indices that contain the specific original text to be replaced.

    Returns list of run indices instead of run objects for reliable identification.
    Handles the realistic scenario where original text spans multiple formatted runs.
    """
    indices_to_remove = []

    if not original_text or str(original_text).lower() in ['nan', '']:
        return indices_to_remove

    # Step 1: Identify all gray-shaded and hyperlink run indices
    formatted_indices = []
    for i, run in enumerate(para.runs):
        if is_run_gray_shaded(run) or is_run_hyperlink(run):
            formatted_indices.append(i)

    if not formatted_indices:
        return indices_to_remove

    # Step 2: Check if original text exists and find overlapping formatted runs
    para_text = para.text.lower()
    original_lower = original_text.lower().strip()

    if original_lower in para_text:
        # Find the position of original text in paragraph
        start_pos = para_text.find(original_lower)
        end_pos = start_pos + len(original_text)

        # Map character positions to runs to find overlapping runs
        char_pos = 0
        for i, run in enumerate(para.runs):
            run_start = char_pos
            run_end = char_pos + len(run.text)

            # If this run overlaps with original text AND is formatted, include it
            if (run_start < end_pos and run_end > start_pos and
                i in formatted_indices):
                indices_to_remove.append(i)

            char_pos = run_end

    return indices_to_remove


def create_hyperlink_run(para: Paragraph, text: str, url: str) -> Run:
    """
    Create a proper hyperlink run in the paragraph.

    The text parameter contains the actual URL/email that should be made clickable.
    The url parameter is the same - it's the target URL/email from the mapping file.

    Handles email detection and proper mailto: prefixing.
    """
    try:
        # Determine if this is an email address
        is_email = '@' in url and not url.startswith(('http://', 'https://', 'www.'))

        # Create proper URL for the hyperlink target
        if is_email:
            target_url = f'mailto:{url}'
        else:
            target_url = url if url.startswith(('http://', 'https://')) else f'https://{url}'

        # Create hyperlink element
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('w:anchor'), target_url)

        # Create run element
        run_element = OxmlElement('w:r')
        run_props = OxmlElement('w:rPr')

        # Add hyperlink formatting
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0000FF')
        run_props.append(color)

        underline = OxmlElement('w:u')
        underline.set(qn('w:val'), 'single')
        run_props.append(underline)

        run_element.append(run_props)

        # Add text (this is the visible text - the URL/email itself)
        text_element = OxmlElement('w:t')
        text_element.text = text
        run_element.append(text_element)

        hyperlink.append(run_element)
        para._element.append(hyperlink)

        # Return the last run (newly created)
        return para.runs[-1]

    except Exception:
        # Fallback to regular run with hyperlink-like formatting
        run = para.add_run(text)
        run.font.color.rgb = RGBColor(0, 0, 255)
        run.underline = True
        return run


def apply_bold_to_specific_text(para: Paragraph, bold_text_list: List[str]) -> None:
    """
    Apply bold formatting to specific text within a paragraph.

    Searches for each text in bold_text_list within the paragraph and makes only
    that text bold, preserving other formatting.
    """
    if not bold_text_list:
        return

    for bold_text in bold_text_list:
        if not bold_text or str(bold_text).lower() in ['nan', '']:
            continue

        bold_text = bold_text.strip()
        para_text = para.text

        # Find the position of the text to bold
        start_pos = para_text.lower().find(bold_text.lower())
        if start_pos == -1:
            continue

        end_pos = start_pos + len(bold_text)

        # Map character positions to runs
        char_pos = 0
        run_ranges = []

        for run in para.runs:
            run_start = char_pos
            run_end = char_pos + len(run.text)
            run_ranges.append((run, run_start, run_end))
            char_pos = run_end

        # Apply bold to runs that overlap with the target text
        for run, run_start, run_end in run_ranges:
            if run_start < end_pos and run_end > start_pos:
                run.bold = True


def insert_country_blocks_at_position(para: Paragraph, country_blocks: List[Dict],
                                    insertion_runs: List[Run]) -> bool:
    """
    Insert formatted country blocks at the position of specified runs.

    Args:
        para: The paragraph to modify
        country_blocks: List of country block data with text, bold_texts, hyperlinks
        insertion_runs: List of runs to replace (these will be removed)

    Returns:
        bool: True if insertion was successful

    Country block format:
    {
        'lines': ['Line 1 text', 'Line 2 text', ...],
        'bold_texts': ['Country1', 'Country2', ...],
        'hyperlinks': [{'text': 'url.com', 'url': 'url.com'}, ...]
    }
    """
    if not country_blocks or not insertion_runs:
        return False

    try:
        # Find the insertion point by locating the first run to replace
        insertion_index = None
        for i, run in enumerate(para.runs):
            if run in insertion_runs:
                insertion_index = i
                break

        if insertion_index is None:
            return False

        # Remove the target runs
        for run in insertion_runs:
            if run in para.runs:
                run._element.getparent().remove(run._element)

        # Insert line break before country blocks
        para.runs[insertion_index:insertion_index] = [para.add_run('\n')]
        insertion_index += 1

        # Insert each country block
        for block_idx, block in enumerate(country_blocks):
            # Add double line break between country blocks (not before the first)
            if block_idx > 0:
                para.runs[insertion_index:insertion_index] = [para.add_run('\n\n')]
                insertion_index += 1

            # Insert each line of the country block
            for line_idx, line_text in enumerate(block.get('lines', [])):
                # Add line break between lines within a block (not before the first line)
                if line_idx > 0:
                    para.runs[insertion_index:insertion_index] = [para.add_run('\n')]
                    insertion_index += 1

                # Process the line for hyperlinks and bold text
                line_parts = [line_text]

                # Find and replace hyperlinks in the line
                for hyperlink in block.get('hyperlinks', []):
                    hyperlink_text = hyperlink.get('text', '')
                    hyperlink_url = hyperlink.get('url', '')

                    if hyperlink_text and hyperlink_url and hyperlink_text in line_text:
                        # Split the line at the hyperlink position
                        new_parts = []
                        for part in line_parts:
                            if hyperlink_text in part:
                                before, after = part.split(hyperlink_text, 1)
                                new_parts.extend([before, {'type': 'hyperlink', 'text': hyperlink_text, 'url': hyperlink_url}, after])
                            else:
                                new_parts.append(part)
                        line_parts = new_parts

                # Insert the line parts as runs
                for part in line_parts:
                    if isinstance(part, dict) and part.get('type') == 'hyperlink':
                        # Insert hyperlink run
                        hyperlink_run = create_hyperlink_run(para, part['text'], part['url'])
                        para.runs.insert(insertion_index, hyperlink_run)
                        insertion_index += 1
                    elif isinstance(part, str) and part:
                        # Insert text run
                        text_run = para.add_run(part)

                        # Apply bold formatting if this text contains bold items
                        for bold_text in block.get('bold_texts', []):
                            if bold_text and bold_text.lower() in part.lower():
                                text_run.bold = True

                        para.runs.insert(insertion_index, text_run)
                        insertion_index += 1

        # Insert line break after country blocks
        para.runs[insertion_index:insertion_index] = [para.add_run('\n')]

        return True

    except Exception as e:
        print(f"Error inserting country blocks: {e}")
        return False


def replace_original_text_with_country_blocks(para: Paragraph, original_text: str,
                                            country_blocks: List[Dict]) -> bool:
    """
    Replace original gray-shaded text with formatted country blocks.

    This is the main function that combines finding runs and inserting blocks.
    Uses run indices to avoid python-docx run object identity issues.
    """
    # Find the run indices containing the original text
    target_indices = find_run_indices_to_remove_original_text(para, original_text)

    if not target_indices:
        return False

    # Insert the country blocks at the specified indices
    return insert_country_blocks_at_indices(para, country_blocks, target_indices)


def insert_country_blocks_at_indices(para: Paragraph, country_blocks: List[Dict],
                                   target_indices: List[int]) -> bool:
    """
    Insert formatted country blocks at the specified run indices.

    Args:
        para: The paragraph to modify
        country_blocks: List of country block data
        target_indices: List of run indices to replace

    Returns:
        bool: True if insertion was successful
    """
    if not country_blocks or not target_indices:
        return False

    try:
        # --- DELETION ---
        # Sort indices in descending order to avoid index shifting during removal
        sorted_indices = sorted(target_indices, reverse=True)

        # Remove target runs from highest index to lowest
        for index in sorted_indices:
            if index < len(para.runs):
                run_element = para.runs[index]._element
                run_element.getparent().remove(run_element)

        # --- INSERTION ---
        # Determine the anchor run to insert before
        insertion_index = min(target_indices)
        anchor_run = None
        if insertion_index < len(para.runs):
            anchor_run = para.runs[insertion_index]

        def insert_run_before_anchor(new_run_element):
            """Helper to insert a new run element at the correct location."""
            if anchor_run:
                anchor_run._element.addprevious(new_run_element)
            else:
                # If there's no anchor, we are at the end of the paragraph
                para._p.append(new_run_element)

        # Helper function to create a text run element
        def create_text_run_element(text, is_bold=False):
            run_elm = OxmlElement('w:r')
            if is_bold:
                rpr = OxmlElement('w:rPr')
                b = OxmlElement('w:b')
                rpr.append(b)
                run_elm.append(rpr)
            text_elm = OxmlElement('w:t')
            text_elm.text = text
            run_elm.append(text_elm)
            return run_elm

        # Create all run elements first and then insert them
        elements_to_insert = []

        # Add line break before blocks
        elements_to_insert.append(create_text_run_element('\n'))

        for block_idx, block in enumerate(country_blocks):
            if block_idx > 0:
                elements_to_insert.append(create_text_run_element('\n\n'))

            for line_idx, line_text in enumerate(block.get('lines', [])):
                if line_idx > 0:
                    elements_to_insert.append(create_text_run_element('\n'))

                # Check if any bold text should be applied to this line
                is_bold = any(bold_text.lower() in line_text.lower()
                             for bold_text in block.get('bold_texts', [])
                             if bold_text)
                elements_to_insert.append(create_text_run_element(line_text, is_bold))

                # Add hyperlinks for this line (simplified - after the line)
                for hyperlink in block.get('hyperlinks', []):
                    if hyperlink.get('text') and hyperlink.get('text') in line_text:
                        elements_to_insert.append(create_text_run_element(f" [{hyperlink['text']}]", is_bold=False))

        # Add line break after blocks
        elements_to_insert.append(create_text_run_element('\n'))

        # Insert all created elements before the anchor
        for element in elements_to_insert:
            insert_run_before_anchor(element)

        return True

    except Exception as e:
        print(f"Error inserting country blocks: {e}")
        return False


def build_replacement_text_by_country(components: List[Dict]) -> str:
    """Build replacement text grouped by country."""
    # Group components by country
    countries = {}
    for comp in components:
        country = comp['country']
        if country not in countries:
            countries[country] = []
        countries[country].append(comp)
    
    # Build the replacement text
    lines = []
    for country, comps in countries.items():
        # Sort components by type (text first, then hyperlinks)
        text_comps = [c for c in comps if c['type'] == 'text']
        link_comps = [c for c in comps if c['type'] == 'hyperlink']
        
        country_parts = []
        for comp in text_comps + link_comps:
            if comp['type'] == 'text':
                country_parts.append(comp['content'])
            elif comp['type'] == 'hyperlink':
                country_parts.append(comp['content'])  # Just the text for now
        
        if country_parts:
            lines.append(' '.join(country_parts))
    
    return '\n'.join(lines)


def get_replacement_components(mapping_row: pd.Series, section_type: str, 
                              cached_components: Optional[List] = None, 
                              country_delimiter: str = ";") -> List:
    """Build replacement text components from mapping data."""
    if cached_components is not None:
        return cached_components
    
    components = []
    country = mapping_row['Country']
    
    if section_type.lower() == "smpc":
        # SmPC Section 4.8 components
        nrs_text = mapping_row.get('4.8 NRS Text', '')
        nrs_url = mapping_row.get('4.8 NRS URL', '')
        av_url = mapping_row.get('4.8 Appendix V URL', '')
        
        if nrs_text:
            components.append({
                'type': 'text',
                'content': str(nrs_text),
                'country': country
            })
        
        if nrs_url:
            components.append({
                'type': 'hyperlink',
                'content': 'national reporting system',
                'url': str(nrs_url),
                'country': country
            })
        
        if av_url:
            components.append({
                'type': 'hyperlink', 
                'content': 'Appendix V',
                'url': str(av_url),
                'country': country
            })
    
    elif section_type.lower() == "pl":
        # PL Section 4 components
        pl_text = mapping_row.get('Section 4 PL Text', '')
        pl_url = mapping_row.get('Section 4 PL URL', '')
        
        if pl_text:
            components.append({
                'type': 'text',
                'content': str(pl_text),
                'country': country
            })
        
        if pl_url:
            components.append({
                'type': 'hyperlink',
                'content': 'national reporting system',
                'url': str(pl_url),
                'country': country
            })
    
    return components


def update_section_10_date(doc: DocumentObject, mapping_row: pd.Series) -> bool:
    """Update the date in Section 10 of Annex I using enhanced date formatter."""
    try:
        section_10_found = False
        country = mapping_row['Country']

        # Try to use specific date header from mapping if available
        date_header = mapping_row.get('Annex I Date Header', '')

        for paragraph in doc.paragraphs:
            # Enhanced section detection using mapping file header if available
            section_match = False
            if date_header and date_header.lower() != 'nan':
                section_match = date_header.lower() in paragraph.text.lower()
            else:
                # Fallback to original detection
                section_match = ("10." in paragraph.text and
                               ("date of" in paragraph.text.lower() or
                                "revision" in paragraph.text.lower()))

            if section_match:
                section_10_found = True
                # Use enhanced date formatter
                formatted_date = format_date_for_country(country, 'annex_i')

                # Replace the date placeholder
                if "<DATE>" in paragraph.text:
                    paragraph.text = paragraph.text.replace("<DATE>", formatted_date)
                    return True

                # Look for existing date pattern and replace with enhanced pattern
                import re
                # Enhanced pattern to catch more date formats
                date_patterns = [
                    r'\d{1,2}\s+\w+\s+\d{4}',      # "12 August 2025"
                    r'\d{1,2}\.\s*\w+\s+\d{4}',    # "12. Aug 2025"
                    r'\d{1,2}/\d{1,2}/\d{4}',      # "12/08/2025"
                    r'\d{4}\.\s*\w+',              # "2025. august"
                ]

                for pattern in date_patterns:
                    if re.search(pattern, paragraph.text):
                        paragraph.text = re.sub(pattern, formatted_date, paragraph.text)
                        return True

        return section_10_found
    except Exception:
        return False


def update_annex_iiib_date(doc: Document, mapping_row: pd.Series) -> bool:
    """Update the date in Section 6 of Annex IIIB using enhanced date formatter."""
    try:
        section_6_found = False
        country = mapping_row['Country']

        # Try to use specific date text from mapping if available
        date_text = mapping_row.get('Annex IIIB Date Text', '')

        for paragraph in doc.paragraphs:
            # Enhanced section detection using mapping file text if available
            section_match = False
            if date_text and date_text.lower() != 'nan':
                section_match = date_text.lower() in paragraph.text.lower()
            else:
                # Fallback to original detection
                section_match = ("6." in paragraph.text and
                               ("detailed information" in paragraph.text.lower() or
                                "leaflet" in paragraph.text.lower()))

            if section_match:
                section_6_found = True
                # Use enhanced date formatter
                formatted_date = format_date_for_country(country, 'annex_iiib')

                # Replace the date placeholder
                if "<DATE>" in paragraph.text:
                    paragraph.text = paragraph.text.replace("<DATE>", formatted_date)
                    return True

                # Look for existing date pattern and replace with enhanced pattern
                import re
                # Enhanced pattern to catch more date formats
                date_patterns = [
                    r'\d{1,2}\s+\w+\s+\d{4}',      # "12 August 2025"
                    r'\d{1,2}\.\s*\w+\s+\d{4}',    # "12. Aug 2025"
                    r'\d{1,2}/\d{1,2}/\d{4}',      # "12/08/2025"
                    r'\d{4}\.\s*\w+',              # "2025. august"
                ]

                for pattern in date_patterns:
                    if re.search(pattern, paragraph.text):
                        paragraph.text = re.sub(pattern, formatted_date, paragraph.text)
                        return True

        return section_6_found
    except Exception:
        return False


def update_local_representatives(doc: Document, mapping_row: pd.Series) -> bool:
    """Update local representatives in Section 6 of Annex IIIB."""
    try:
        updated = False
        local_rep_text = mapping_row.get('Local Representative', '')
        
        if not local_rep_text or str(local_rep_text).lower() == 'nan':
            return False
        
        for paragraph in doc.paragraphs:
            if ("local representative" in paragraph.text.lower() or 
                "section 6" in paragraph.text.lower()):
                
                # Replace placeholder
                if "<LOCAL_REP>" in paragraph.text:
                    paragraph.text = paragraph.text.replace("<LOCAL_REP>", str(local_rep_text))
                    updated = True
        
        return updated
    except Exception:
        return False
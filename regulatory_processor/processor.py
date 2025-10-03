"""Document processing module for EU SmPC and PL documents.

This module contains helper functions to load the mapping Excel file,
update the Annex I (SmPC) and Annex IIIB (Package Leaflet) sections
of combined SmPC Word documents, split the updated documents into
separate Annex I and Annex IIIB files, and convert those files to
PDF using LibreOffice.

Complete implementation with all required functionality:
- National reporting system updates (Section 4.8 SmPC, Section 4 PL)
- Date updates (Section 10 Annex I, Section 6 Annex IIIB)
- Local representatives update (Section 6 Annex IIIB)
- Multi-file generation for languages with multiple countries
- Proper file naming conventions
"""


import os
import re
import shutil
import logging
import locale
import calendar
import asyncio
from pathlib import Path
from typing import Dict, List, Optional, Tuple, NamedTuple, Union
from dataclasses import dataclass
from datetime import datetime
import pandas as pd
import subprocess
from copy import deepcopy
from docx import Document
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.shared import RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.oxml import OxmlElement
from regulatory_processor.document_splitter import (
    clone_and_split_document,
    copy_paragraph, copy_table, _copy_paragraph_content,
    copy_document_structure, copy_headers_and_footers, copy_styles
)
from docx2pdf import convert

# Import refactored modules
from .config import (
    DirectoryNames, FileMarkers, SectionTypes,
    ProcessingConfig, ProcessingResult, ProcessingStats
)
from .exceptions import (
    ProcessingError, ValidationError, DocumentError, MappingError
)
from .date_formatter import (
    DateFormatterSystem, initialize_date_formatter, get_date_formatter,
    format_date_for_country, format_date
)
from .utils import (
    get_country_code_mapping, extract_country_code_from_filename,
    identify_document_country_and_language, find_mapping_rows_for_language,
    generate_output_filename, load_mapping_table, is_header_match
)
from .hyperlinks import (
    URLValidationResult, URLAccessibilityResult, URLValidationConfig,
    validate_url_format, test_url_accessibility, add_hyperlink_relationship,
    create_hyperlink_run_enhanced, create_hyperlink_element,
    create_styled_text_fallback_element, validate_and_test_url_complete
)

# Define utility functions that are processor-specific

class ThreadSafePDFConverter:
    """
    Singleton class to handle PDF conversion in a thread-safe manner.

    LibreOffice has issues when run from multiple threads simultaneously,
    so we use a dedicated thread with a queue to serialize all PDF conversions.
    """
    _instance = None
    _lock = None

    def __new__(cls):
        if cls._instance is None:
            import threading
            cls._lock = threading.Lock()
            with cls._lock:
                if cls._instance is None:
                    cls._instance = super().__new__(cls)
                    cls._instance._initialized = False
        return cls._instance

    def __init__(self):
        if self._initialized:
            return

        import threading
        import queue
        import os

        self._initialized = True
        self._conversion_queue = queue.Queue()
        self._worker_thread = None
        self._shutdown_event = threading.Event()
        self._start_worker()

    def _start_worker(self):
        """Start the worker thread that handles PDF conversions."""
        import threading

        if self._worker_thread is None or not self._worker_thread.is_alive():
            self._worker_thread = threading.Thread(
                target=self._worker_loop,
                daemon=True,
                name="PDFConverter"
            )
            self._worker_thread.start()

    def _worker_loop(self):
        """Main loop for the PDF conversion worker thread."""
        import subprocess
        import os
        import time
        from pathlib import Path

        while not self._shutdown_event.is_set():
            try:
                # Get conversion task from queue (timeout to check shutdown)
                task = self._conversion_queue.get(timeout=1.0)
                if task is None:  # Sentinel value for shutdown
                    break

                doc_path, output_dir, result_queue = task

                try:
                    # Perform the actual LibreOffice conversion
                    pdf_output_path = Path(output_dir) / Path(doc_path).with_suffix(".pdf").name

                    # Find LibreOffice command
                    libreoffice_cmd = _find_libreoffice_command()
                    if not libreoffice_cmd:
                        result_queue.put(("error", "LibreOffice command not found"))
                        continue

                    # Set environment variables for headless operation
                    env = os.environ.copy()
                    env.update({
                        'DISPLAY': ':0.0' if 'DISPLAY' not in env else env['DISPLAY'],
                        'LIBGL_ALWAYS_SOFTWARE': '1',  # Force software rendering
                        'QT_QPA_PLATFORM': 'offscreen',  # Qt platform for headless
                    })

                    # Run LibreOffice conversion
                    result = subprocess.run(
                        [
                            libreoffice_cmd, '--headless', '--convert-to', 'pdf',
                            '--outdir', str(output_dir), doc_path
                        ],
                        timeout=60,
                        stdout=subprocess.DEVNULL,
                        stderr=subprocess.PIPE,
                        text=True,
                        env=env  # Pass explicit environment
                    )

                    if result.returncode == 0 and pdf_output_path.exists():
                        result_queue.put(("success", str(pdf_output_path)))
                    else:
                        error_msg = result.stderr if result.stderr else f"Return code: {result.returncode}"
                        result_queue.put(("error", f"LibreOffice failed: {error_msg}"))

                except subprocess.TimeoutExpired:
                    result_queue.put(("error", "LibreOffice conversion timed out"))
                except Exception as e:
                    result_queue.put(("error", f"Conversion error: {str(e)}"))
                finally:
                    self._conversion_queue.task_done()

            except Exception:
                # Queue timeout or other error - continue loop
                continue

    def convert(self, doc_path: str, output_dir: str, timeout: float = 70.0) -> tuple[str, str]:
        """
        Convert document to PDF using the dedicated worker thread.

        Args:
            doc_path: Path to the input document
            output_dir: Directory for output PDF
            timeout: Maximum time to wait for conversion

        Returns:
            tuple: (status, result) where status is 'success' or 'error'
        """
        import queue
        import threading

        # Ensure worker is running
        self._start_worker()

        # Create result queue for this conversion
        result_queue = queue.Queue()

        # Submit conversion task
        self._conversion_queue.put((doc_path, output_dir, result_queue))

        # Wait for result with timeout
        try:
            status, result = result_queue.get(timeout=timeout)
            return status, result
        except queue.Empty:
            return "error", "Conversion timed out waiting for worker thread"

def _find_libreoffice_command():
    """Find the available LibreOffice command on the system.

    Returns:
        str or None: The LibreOffice command name if found, None otherwise.
    """
    import shutil

    # Try different possible LibreOffice command names
    commands_to_try = ['soffice', 'libreoffice', 'loffice']
    for cmd in commands_to_try:
        if shutil.which(cmd):
            return cmd
    return None

def convert_to_pdf(doc_path: str, output_dir: str) -> str:
    """Convert a Word document to PDF using thread-safe converter with multiple fallback methods."""
    import time
    import gc
    from pathlib import Path
    import sys

    # Force cleanup before conversion
    gc.collect()

    pdf_output_path = Path(output_dir) / Path(doc_path).with_suffix(".pdf").name
    print(f"   🔄 Converting: {Path(doc_path).name} → {pdf_output_path.name}")
    sys.stdout.flush()

    # Add small delay to prevent resource conflicts
    time.sleep(0.5)

    # Method 1: Try thread-safe LibreOffice converter (primary method)
    print(f"   🐧 Method 1: Using thread-safe LibreOffice conversion...")
    sys.stdout.flush()

    try:
        # Use the thread-safe PDF converter
        converter = ThreadSafePDFConverter()
        status, result = converter.convert(doc_path, output_dir, timeout=70.0)

        if status == "success":
            print(f"   ✅ Thread-safe LibreOffice conversion successful")
            return result
        else:
            print(f"   ⚠️ Thread-safe LibreOffice conversion failed: {result}")
            raise RuntimeError(f"Thread-safe LibreOffice failed: {result}")

    except Exception as e:
        print(f"   ⚠️ Thread-safe LibreOffice conversion error: {e}")

    # Method 2: Try docx2pdf with timeout protection (fallback method)
    print(f"   📝 Method 2: Attempting docx2pdf conversion...")
    sys.stdout.flush()
    try:
        # Use subprocess to run docx2pdf with timeout control
        conversion_script = f'''
import sys
from docx2pdf import convert
import time
start_time = time.time()
try:
    convert(r"{doc_path}", r"{pdf_output_path}")
    print(f"Conversion completed in {{time.time() - start_time:.2f}} seconds")
except Exception as e:
    print(f"Conversion failed: {{e}}")
    sys.exit(1)
'''
        result = subprocess.run([
            'python', '-c', conversion_script
        ], capture_output=True, text=True, timeout=15)

        if result.returncode == 0 and pdf_output_path.exists():
            print(f"   ✅ docx2pdf conversion successful")
            return str(pdf_output_path)
        else:
            print(f"   ⚠️ docx2pdf conversion failed: {result.stderr}")

    except subprocess.TimeoutExpired:
        print(f"   ⚠️ docx2pdf conversion timed out after 15 seconds")
    except Exception as e:
        print(f"   ⚠️ docx2pdf error: {e}")

    # Clean up after failed conversion attempt
    gc.collect()
    time.sleep(0.5)

    # Method 3: Try pandoc (if available)
    print(f"   📚 Method 3: Attempting pandoc conversion...")
    try:
        result = subprocess.run([
            'pandoc', doc_path, '-o', str(pdf_output_path)
        ], capture_output=True, text=True, timeout=30)

        if result.returncode == 0 and pdf_output_path.exists():
            print(f"   ✅ pandoc conversion successful")
            return str(pdf_output_path)
        else:
            print(f"   ⚠️ pandoc conversion failed: {result.stderr}")
    except (subprocess.TimeoutExpired, FileNotFoundError) as e:
        print(f"   ⚠️ pandoc not available: {e}")

    # Method 4: Create a placeholder PDF (last resort)
    print(f"   📄 Method 4: Creating placeholder file...")
    try:
        # Create a simple text file indicating conversion failed
        placeholder_path = pdf_output_path.with_suffix('.pdf.txt')
        with open(placeholder_path, 'w') as f:
            f.write(f"PDF conversion failed for: {Path(doc_path).name}\n")
            f.write(f"Original document available at: {doc_path}\n")
            f.write(f"Please convert manually or install another conversion tool.\n")

        print(f"   📝 Created placeholder file: {placeholder_path.name}")
        return str(placeholder_path)

    except Exception as e:
        print(f"   ❌ All conversion methods failed: {e}")
        raise RuntimeError(f"Failed to convert {doc_path} to PDF: All methods failed")

# copy_paragraph function moved to document_splitter.py


# copy_table function moved to document_splitter.py


# _copy_paragraph_content function moved to document_splitter.py


# copy_document_structure function moved to document_splitter.py


def copy_headers_and_footers(source_doc: Document, dest_doc: Document) -> None:
    """Copy headers and footers from source document to destination document."""

    try:
        # Ensure destination has at least one section
        if not dest_doc.sections:
            return

        source_section = source_doc.sections[0] if source_doc.sections else None
        dest_section = dest_doc.sections[0]

        if not source_section:
            return

        # Copy headers
        header_types = [
            ('first_page_header', 'first_page_header'),
            ('even_page_header', 'even_page_header'),
            ('header', 'header')  # Default header
        ]

        for source_attr, dest_attr in header_types:
            try:
                if hasattr(source_section, source_attr) and hasattr(dest_section, dest_attr):
                    source_header = getattr(source_section, source_attr)
                    dest_header = getattr(dest_section, dest_attr)

                    # Clear existing content
                    for para in dest_header.paragraphs:
                        para.clear()

                    # Copy paragraphs from source header
                    for i, source_para in enumerate(source_header.paragraphs):
                        if i < len(dest_header.paragraphs):
                            _copy_paragraph_content(dest_header.paragraphs[i], source_para)
                        else:
                            dest_header_para = dest_header.add_paragraph()
                            _copy_paragraph_content(dest_header_para, source_para)

            except Exception as e:
                print(f"⚠️ Could not copy header {source_attr}: {e}")

        # Copy footers
        footer_types = [
            ('first_page_footer', 'first_page_footer'),
            ('even_page_footer', 'even_page_footer'),
            ('footer', 'footer')  # Default footer
        ]

        for source_attr, dest_attr in footer_types:
            try:
                if hasattr(source_section, source_attr) and hasattr(dest_section, dest_attr):
                    source_footer = getattr(source_section, source_attr)
                    dest_footer = getattr(dest_section, dest_attr)

                    # Clear existing content
                    for para in dest_footer.paragraphs:
                        para.clear()

                    # Copy paragraphs from source footer
                    for i, source_para in enumerate(source_footer.paragraphs):
                        if i < len(dest_footer.paragraphs):
                            _copy_paragraph_content(dest_footer.paragraphs[i], source_para)
                        else:
                            dest_footer_para = dest_footer.add_paragraph()
                            _copy_paragraph_content(dest_footer_para, source_para)

            except Exception as e:
                print(f"⚠️ Could not copy footer {source_attr}: {e}")

    except Exception as e:
        print(f"⚠️ Could not copy headers and footers: {e}")


def copy_styles(source_doc: Document, dest_doc: Document) -> None:
    """Copy custom styles from source document to destination document."""

    try:
        # Get style collections
        source_styles = source_doc.styles
        dest_styles = dest_doc.styles

        # Copy paragraph styles
        for source_style in source_styles:
            if source_style.type == 1:  # Paragraph style
                try:
                    # Check if style already exists
                    existing_style = None
                    try:
                        existing_style = dest_styles[source_style.name]
                    except KeyError:
                        pass

                    if not existing_style:
                        # Create new style
                        new_style = dest_styles.add_style(source_style.name, 1)  # 1 = paragraph style

                        # Copy basic properties
                        if hasattr(source_style, 'font') and hasattr(new_style, 'font'):
                            if source_style.font.name:
                                new_style.font.name = source_style.font.name
                            if source_style.font.size:
                                new_style.font.size = source_style.font.size

                except Exception as e:
                    print(f"⚠️ Could not copy style {source_style.name}: {e}")

    except Exception as e:
        print(f"⚠️ Could not copy styles: {e}")

# ============================================================================= 
# CONSTANTS AND CONFIGURATION
# =============================================================================




# ============================================================================= 
# DOCUMENT UPDATE FUNCTIONS
# =============================================================================
def is_hex_gray_color(hex_color: str) -> bool:
    """Check if a hex color represents a gray shade."""
    if not hex_color:
        return False
    
    hex_color = hex_color.replace('#', '').upper()
    
    gray_hex_values = [
        'BFBFBF', 'CCCCCC', 'D9D9D9', '808080', '999999', 
        '666666', 'C0C0C0', 'A0A0A0'
    ]
    
    if hex_color in gray_hex_values:
        return True
    
    # Check if R=G=B (indicates gray)
    try:
        if len(hex_color) == 6:
            r = int(hex_color[0:2], 16)
            g = int(hex_color[2:4], 16)
            b = int(hex_color[4:6], 16)
            return r == g == b
    except ValueError:
        pass
    
    return False


def is_run_gray_shaded(run: Run) -> bool:
    """Check if a run has gray shading."""
    try:
        run_pr = run._element.get_or_add_rPr()
        shading_elements = run_pr.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
        
        if shading_elements:
            for shading in shading_elements:
                fill = shading.get(qn('w:fill'))
                if fill and is_hex_gray_color(fill):
                    return True
        
        # Check font color for gray
        if run.font.color and hasattr(run.font.color, 'rgb'):
            gray_colors = [
                RGBColor(128, 128, 128), RGBColor(153, 153, 153), 
                RGBColor(102, 102, 102), RGBColor(96, 96, 96),
                RGBColor(217, 217, 217), RGBColor(191, 191, 191)
            ]
            if run.font.color.rgb in gray_colors:
                return True
                
    except Exception:
        pass
    
    return False


def is_run_gray_shaded_enhanced(run: Run) -> bool:
    """Enhanced gray shading detection with comprehensive color matching."""
    try:
        # First use the original method
        if is_run_gray_shaded(run):
            return True

        # Check run properties for shading with more extensive color list
        run_pr = run._element.get_or_add_rPr()
        shading_elements = run_pr.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')

        if shading_elements:
            for shading in shading_elements:
                fill = shading.get(qn('w:fill'))
                if fill:
                    # Extended gray color list in hex format
                    gray_hex_colors = [
                        'd9d9d9', 'cccccc', 'c0c0c0', 'bfbfbf', 'b3b3b3', 'a0a0a0',
                        '999999', '808080', '666666', '606060', 'f5f5f5', 'e0e0e0',
                        'lightgray', 'gray', 'darkgray', 'auto'
                    ]
                    if fill.lower() in gray_hex_colors:
                        return True

        # Enhanced font color checking with more gray variations
        if run.font.color and hasattr(run.font.color, 'rgb') and run.font.color.rgb is not None:
            color = run.font.color.rgb
            # Expanded gray color list
            gray_colors = [
                # Original grays
                RGBColor(128, 128, 128), RGBColor(153, 153, 153),
                RGBColor(102, 102, 102), RGBColor(96, 96, 96),
                RGBColor(217, 217, 217), RGBColor(191, 191, 191),
                # Additional gray variations
                RGBColor(160, 160, 160), RGBColor(192, 192, 192),
                RGBColor(224, 224, 224), RGBColor(245, 245, 245),
                RGBColor(179, 179, 179), RGBColor(140, 140, 140),
                RGBColor(112, 112, 112), RGBColor(75, 75, 75)
            ]

            # Check for exact matches
            if color in gray_colors:
                return True

            # Check if color components are approximately equal (indicating gray)
            if abs(color.r - color.g) < 20 and abs(color.g - color.b) < 20 and abs(color.r - color.b) < 20:
                # It's some shade of gray
                return True

    except Exception as e:
        print(f"Warning: Gray shading detection failed: {e}")

    return False


def is_run_hyperlink(run: Run) -> bool:
    """Check if a run is part of a hyperlink."""
    try:
        run_xml = run._r
        hyperlink_elements = run_xml.xpath('.//w:hyperlink')
        if hyperlink_elements:
            return True
            
        # Check hyperlink-style formatting
        if (run.font.color and hasattr(run.font.color, 'rgb') and 
            run.font.color.rgb == RGBColor(0, 0, 255) and run.underline):
            return True
            
    except Exception:
        pass
    
    return False


def is_run_hyperlink_enhanced(run: Run) -> bool:
    """Enhanced hyperlink detection with comprehensive checks."""
    try:
        # First use the original method
        if is_run_hyperlink(run):
            return True

        # Enhanced XML-based hyperlink detection
        run_xml = run._r

        # Check multiple possible hyperlink XML patterns
        hyperlink_patterns = [
            './/w:hyperlink',
            './/*[@w:anchor]',  # Internal links
            './/*[@r:id]',      # External links with relationship ID
            './/w:instrText',   # Field codes (can contain hyperlinks)
        ]

        for pattern in hyperlink_patterns:
            try:
                elements = run_xml.xpath(pattern)
                if elements:
                    return True
            except:
                continue

        # Enhanced styling-based detection
        if run.font.color and hasattr(run.font.color, 'rgb') and run.font.color.rgb is not None:
            color = run.font.color.rgb

            # Common hyperlink colors
            hyperlink_colors = [
                RGBColor(0, 0, 255),    # Standard blue
                RGBColor(0, 0, 238),    # Slightly different blue
                RGBColor(5, 99, 193),   # Word default hyperlink blue
                RGBColor(17, 85, 204),  # Google Docs blue
                RGBColor(70, 120, 180), # Alternative blue
            ]

            # Check for hyperlink colors with or without underline
            if color in hyperlink_colors:
                return True

            # Check for blue-ish colors that might be hyperlinks
            if color.r < 100 and color.g < 150 and color.b > 150:
                return True

        # Check if run has underline (common for hyperlinks)
        if run.underline and run.font.color:
            return True

        # Check if text looks like a URL
        text = run.text.strip().lower()
        if any(url_start in text for url_start in ['http://', 'https://', 'www.', 'mailto:']):
            return True

    except Exception as e:
        print(f"Warning: Hyperlink detection failed: {e}")

    return False


def _is_old_reporting_run(run: Run, target_string: str) -> bool:
    """Check if a run matches patterns typical of old national reporting text."""
    try:
        text = run.text.strip()
        if not text:
            return False

        # Common patterns in old national reporting systems
        old_reporting_patterns = [
            'adverse',
            'reaction',
            'reporting',
            'national',
            'system',
            'side effect',
            'suspected',
            'medicine',
            'drug',
            'authority',
            'agency'
        ]

        text_lower = text.lower()
        target_lower = target_string.lower()

        # Check if text contains multiple old reporting keywords
        keyword_count = sum(1 for pattern in old_reporting_patterns if pattern in text_lower)
        if keyword_count >= 2:
            return True

        # Check if text is part of the target string
        if len(text) > 5 and text_lower in target_lower:
            return True

        # Check if text contains parts of target string
        target_words = target_lower.split()
        text_words = text_lower.split()
        matching_words = len(set(target_words) & set(text_words))
        if matching_words > 0 and len(text_words) <= 10:  # Short text with target words
            return True

    except Exception:
        pass

    return False


def find_target_text_runs(para: Paragraph, target_string: str) -> List[Run]:
    """
    Find runs that contain the target text to be replaced.
    """
    target_runs = []
    
    # Build text map to find position
    char_pos = 0
    run_positions = []
    
    for run in para.runs:
        run_start = char_pos
        run_end = char_pos + len(run.text)
        run_positions.append((run, run_start, run_end))
        char_pos = run_end
    
    # Find target string position in full text
    full_text = para.text
    target_start = full_text.lower().find(target_string.lower())
    
    if target_start == -1:
        return []
    
    target_end = target_start + len(target_string)
    
    # Find runs that overlap with target text
    for run, run_start, run_end in run_positions:
        # Check if run overlaps with target range
        if (run_start < target_end and run_end > target_start):
            target_runs.append(run)
    
    return target_runs


def find_target_text_range(para: Paragraph, target_string: str) -> Tuple[int, int]:
    """Find the complete target text range in paragraph."""
    full_text = para.text.lower()
    target_lower = target_string.lower()
    
    # Try exact match first
    start_pos = full_text.find(target_lower)
    if start_pos != -1:
        return start_pos, start_pos + len(target_string)
    
    # Try to find key parts
    key_phrases = ['national reporting system', 'appendix v', 'listed in appendix']
    
    earliest_start = len(full_text)
    latest_end = 0
    
    for phrase in key_phrases:
        pos = full_text.find(phrase)
        if pos != -1:
            earliest_start = min(earliest_start, pos)
            latest_end = max(latest_end, pos + len(phrase))
    
    if earliest_start < len(full_text):
        return earliest_start, latest_end
    
    return -1, -1


def find_runs_to_remove(para: Paragraph, target_string: str) -> List[Run]:
    """Find runs that should be removed - Enhanced with XML-based hyperlink handling.

    This function now uses XML-based removal for complex cases (like invisible hyperlinks)
    and falls back to the original run-based approach if needed.
    """
    if not target_string.strip():
        return []

    print(f"\n🎯 ENHANCED TEXT REMOVAL")
    print(f"Target: '{target_string}'")
    print(f"Paragraph text: '{para.text}'")

    # Check if we have a runs vs text mismatch (indicates invisible hyperlinks)
    para_text_len = len(para.text)
    runs_text_len = sum(len(run.text) for run in para.runs)
    has_invisible_content = para_text_len != runs_text_len

    if has_invisible_content:
        print(f"🔍 Detected invisible content (text: {para_text_len}, runs: {runs_text_len} chars)")
        print(f"🎯 Using XML-based removal for hyperlink handling...")

        # Use XML-based removal for invisible content
        success = _remove_target_text_xml_internal(para, target_string)
        if success:
            print(f"✅ XML-based removal completed")
            return []  # Return empty list since removal was done directly
        else:
            print(f"⚠️  XML removal failed, falling back to run-based approach")

    # Original run-based approach (fallback or primary for simple cases)
    print(f"🎯 Using run-based removal...")
    runs_to_remove = []

    # Find target text range
    target_start, target_end = find_target_text_range(para, target_string)

    if target_start == -1:
        print(f"❌ Target text not found")
        return runs_to_remove

    print(f"✅ Target found at position {target_start}-{target_end}")

    # Map character positions to runs
    char_pos = 0
    run_ranges = []

    for run in para.runs:
        run_start = char_pos
        run_end = char_pos + len(run.text)
        run_ranges.append((run, run_start, run_end))
        char_pos = run_end

    # Find runs that overlap with target text and are styled
    for i, (run, run_start, run_end) in enumerate(run_ranges):
        if run_start < target_end and run_end > target_start:
            is_gray = is_run_gray_shaded(run)
            is_hyperlink = is_run_hyperlink(run)

            if is_gray or is_hyperlink or run.text.strip() in target_string:
                runs_to_remove.append(run)
                reasons = []
                if is_gray: reasons.append("gray shaded")
                if is_hyperlink: reasons.append("hyperlink")
                if run.text.strip() in target_string: reasons.append("contains target text")
                print(f"  ✅ REMOVING Run {i}: '{run.text}' - {', '.join(reasons)}")
            else:
                print(f"  ⏭️  KEEPING Run {i}: '{run.text}' - not styled")
        else:
            print(f"  ⏭️  KEEPING Run {i}: '{run.text}' - outside target range")

    print(f"🗑️  Will remove {len(runs_to_remove)} runs out of {len(run_ranges)} total")
    return runs_to_remove


def _remove_target_text_xml_internal(paragraph: Paragraph, target_string: str) -> bool:
    """
    Internal XML-based text removal for invisible hyperlinks.

    This handles cases where hyperlink runs are not exposed in paragraph.runs.
    """
    p_element = paragraph._element
    w_namespace = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    try:
        # Extract text from XML
        try:
            all_text_nodes = p_element.xpath('.//w:t/text()', namespaces={'w': w_namespace})
            full_text = "".join(all_text_nodes)
        except:
            text_elements = p_element.findall(f'.//{{{w_namespace}}}t')
            full_text = "".join(t.text or '' for t in text_elements)

        # Handle text duplication (corrupted docs)
        if len(full_text) > 0 and full_text == (full_text[:len(full_text)//2] * 2):
            full_text = full_text[:len(full_text)//2]
            print(f"📝 Cleaned duplicated text")

        # Find target position
        target_start = full_text.lower().find(target_string.lower())
        if target_start == -1:
            return False

        target_end = target_start + len(target_string)

        # Get all XML runs
        try:
            all_runs = p_element.xpath('.//w:r', namespaces={'w': w_namespace})
        except:
            all_runs = p_element.findall(f'.//{{{w_namespace}}}r')

        print(f"🔍 Processing {len(all_runs)} XML runs...")

        # Process runs and mark for deletion/modification
        current_pos = 0
        runs_to_delete = []
        runs_to_modify = []

        for run_element in all_runs:
            try:
                text_elements = run_element.xpath('./w:t', namespaces={'w': w_namespace})
            except:
                text_elements = run_element.findall(f'{{{w_namespace}}}t')

            run_text = "".join(t.text or '' for t in text_elements)
            if len(run_text) > 0 and run_text == (run_text[:len(run_text)//2] * 2):
                run_text = run_text[:len(run_text)//2]

            run_start = current_pos
            run_end = current_pos + len(run_text)

            # Check overlap with target
            if run_start < target_end and run_end > target_start:
                if run_start >= target_start and run_end <= target_end:
                    runs_to_delete.append(run_element)
                else:
                    runs_to_modify.append((run_element, text_elements))

            current_pos = run_end

        # Execute deletions
        for run_element in runs_to_delete:
            # Check if run is inside a hyperlink
            hyperlink_parent = run_element.getparent()
            while hyperlink_parent is not None:
                if hyperlink_parent.tag.endswith('hyperlink'):
                    hyperlink_parent.getparent().remove(hyperlink_parent)
                    break
                hyperlink_parent = hyperlink_parent.getparent()
            else:
                parent = run_element.getparent()
                if parent is not None:
                    parent.remove(run_element)

        # Execute modifications
        for run_element, text_elements in runs_to_modify:
            for t in text_elements:
                t.text = ""

        return True

    except Exception as e:
        print(f"❌ XML removal error: {e}")
        return False


def find_gray_and_hyperlink_runs(para: Paragraph, target_string: str) -> List[Run]:
    """
    Find all gray shaded runs and hyperlink runs that should be removed.
    """
    runs_to_remove = []
    
    # First find runs containing target text
    target_runs = find_target_text_runs(para, target_string)
    
    # Then find additional gray/hyperlink runs in vicinity
    for run in para.runs:
        should_remove = False
        
        # Remove if it's a target run
        if run in target_runs:
            should_remove = True
        # Remove if it's gray shaded
        elif is_run_gray_shaded(run):
            should_remove = True
        # Remove if it's a hyperlink
        elif is_run_hyperlink(run):
            should_remove = True
        
        if should_remove:
            runs_to_remove.append(run)
    
    return runs_to_remove


# OLD BROKEN IMPLEMENTATION - REMOVED IN STEP 3.4
# The create_hyperlink_run function has been moved to Step 3.4 section with proper implementation
# This broken version used w:anchor instead of proper document relationships











def build_replacement_text_by_country(components: List[Dict]) -> str:
    """Build replacement text grouped by country."""
    # Group components by country
    countries = {}
    for comp in components:
        country = comp['country']
        if country not in countries:
            countries[country] = []
        countries[country].append(comp)
    
    # Build text for each country
    country_blocks = []
    
    for country_name, country_components in countries.items():
        # Sort components by line number
        sorted_components = sorted(country_components, key=lambda x: x['line'])
        
        # Build country block
        country_lines = []
        
        for comp in sorted_components:
            line_text = comp['text']
            hyperlink = comp.get('hyperlink')
            email = comp.get('email')
            
            # Add hyperlinks if present and not already in text
            if hyperlink and hyperlink not in line_text:
                line_text += f" {hyperlink}"
            if email and email not in line_text:
                line_text += f" {email}"
            
            country_lines.append(line_text)
        
        # Join lines for this country
        country_block = '\n'.join(country_lines)
        country_blocks.append(country_block)
    
    # Join country blocks with double line breaks
    return '\n\n'.join(country_blocks)

def get_replacement_components(mapping_row: pd.Series, section_type: str,
                              cached_components: Optional[List] = None,
                              country_delimiter: str = ";") -> List:
    """Build replacement text components from mapping data.

    Now supports multi-country block separation by grouping components by country.
    Each country gets its own complete block that will be separated by double line breaks.
    """
    if cached_components is not None:
        return cached_components

    components = []

    # Get line columns for this section type
    line_columns = [col for col in mapping_row.index
                   if col.startswith('Line ') and section_type in col]

    if not line_columns:
        return components

    # Get hyperlinks and email links
    hyperlinks_col = f'Hyperlinks {section_type}'
    email_col = f'Link for email - {section_type}'

    hyperlinks_str = str(mapping_row.get(hyperlinks_col, '')).strip()
    email_str = str(mapping_row.get(email_col, '')).strip()

    # Parse hyperlinks and emails (semicolon separated)
    hyperlinks = [h.strip() for h in hyperlinks_str.split(country_delimiter)
                 if h.strip() and h.strip().lower() != 'nan']
    emails = [e.strip() for e in email_str.split(country_delimiter)
             if e.strip() and e.strip().lower() != 'nan']

    # Sort line columns by number
    def extract_line_number(col_name):
        match = re.search(r'Line (\d+)', col_name)
        return int(match.group(1)) if match else 999

    sorted_columns = sorted(line_columns, key=extract_line_number)

    # Find Line 1 to get countries
    line_1_col = None
    for col in sorted_columns:
        if extract_line_number(col) == 1:
            line_1_col = col
            break

    if not line_1_col:
        return components

    line_1_text = str(mapping_row.get(line_1_col, '')).strip()
    if not line_1_text or line_1_text.lower() == 'nan':
        return components

    # Get countries from dedicated bold country column
    bold_countries_col = f'Line 1 - Country names to be bolded - {section_type}'
    bold_countries_str = str(mapping_row.get(bold_countries_col, '')).strip()

    # Parse countries using comma/semicolon delimiter
    if bold_countries_str and bold_countries_str.lower() != 'nan':
        # Try comma first (as seen in mapping file), then semicolon as fallback
        if ',' in bold_countries_str:
            countries = [c.strip() for c in bold_countries_str.split(',') if c.strip()]
        else:
            countries = [c.strip() for c in bold_countries_str.split(country_delimiter) if c.strip()]
    else:
        # Fallback: extract from line text (backwards compatibility)
        countries = [c.strip() for c in line_1_text.split(country_delimiter) if c.strip()]

    if not countries:
        return components

    # Process each line
    for col in sorted_columns:
        line_num = extract_line_number(col)
        content = str(mapping_row.get(col, '')).strip()

        if not content or content.lower() == 'nan':
            continue

        # Split content by countries using semicolon delimiter
        parts = [p.strip() for p in content.split(country_delimiter)]

        for i, country in enumerate(countries):
            if i < len(parts) and parts[i]:
                text = parts[i]

                # Determine links for this country position
                hyperlink = hyperlinks[i] if i < len(hyperlinks) else None
                email = emails[i] if i < len(emails) else None

                components.append({
                    'line': line_num,
                    'country': country,
                    'country_index': i,  # NEW: Add country index for grouping
                    'text': text,
                    'hyperlink': hyperlink,
                    'email': email
                })

    return components


def insert_formatted_replacement_surgically(para: Paragraph, insertion_point: int, 
                                          components: List[Dict], country_delimiter: str = ";",
                                          document: Document = None):
    """
    Insert properly formatted replacement text at a specific position in the paragraph.
    
    NEW: Only hyperlinks the specific text that matches the hyperlink/email value.
    """
    # Group components by country first, then by line within each country
    countries = {}
    for comp in components:
        country_key = (comp['country'], comp.get('country_index', 0))
        line_num = comp['line']

        if country_key not in countries:
            countries[country_key] = {'country': comp['country'], 'lines': {}}

        if line_num not in countries[country_key]['lines']:
            countries[country_key]['lines'][line_num] = []
        countries[country_key]['lines'][line_num].append(comp)

    sorted_country_keys = sorted(countries.keys(), key=lambda x: x[1])

    current_element = None
    if insertion_point < len(para.runs):
        current_element = para.runs[insertion_point]._element

    # Add a single line break BEFORE the first country block
    first_break_run_xml = OxmlElement('w:r')
    first_break_run_xml.append(OxmlElement('w:br'))
    
    if current_element is not None:
        current_element.addnext(first_break_run_xml)
        current_element = first_break_run_xml
    else:
        para._p.append(first_break_run_xml)
        current_element = first_break_run_xml

    # Build replacement text country by country
    for country_idx, country_key in enumerate(sorted_country_keys):
        country_info = countries[country_key]
        lines = country_info['lines']

        if country_idx > 0:
            double_break_run_xml = OxmlElement('w:r')
            double_break_run_xml.append(OxmlElement('w:br'))
            double_break_run_xml.append(OxmlElement('w:br'))
            current_element.addnext(double_break_run_xml)
            current_element = double_break_run_xml

        # Build lines within this country
        for line_idx, line_num in enumerate(sorted(lines.keys())):
            line_components = lines[line_num]

            if line_idx > 0:
                line_break_run_xml = OxmlElement('w:r')
                line_break_run_xml.append(OxmlElement('w:br'))
                current_element.addnext(line_break_run_xml)
                current_element = line_break_run_xml

            for comp_idx, comp in enumerate(line_components):
                text = comp['text']
                country = comp['country']
                hyperlink = comp.get('hyperlink')
                email = comp.get('email')

                # ==========================================================
                # START: MODIFIED LOGIC
                # ==========================================================
                
                is_email_link = email and email in text
                is_hyperlink = not is_email_link and hyperlink and hyperlink in text

                # --- 1. RENDER AS EMAIL LINK (SPLIT) ---
                if is_email_link:
                    parts = re.split(f'({re.escape(email)})', text, 1)
                    email_url = f'mailto:{email}' if not email.startswith('mailto:') else email

                    if parts[0]:
                        text_run = para.add_run(parts[0])
                        current_element.addnext(text_run._element)
                        current_element = text_run._element
                    
                    link_element = create_hyperlink_element(para, email, email_url, document)
                    current_element.addnext(link_element)
                    current_element = link_element

                    if len(parts) > 2 and parts[2]:
                        text_run = para.add_run(parts[2])
                        current_element.addnext(text_run._element)
                        current_element = text_run._element

                # --- 2. RENDER AS HYPERLINK (SPLIT) ---
                elif is_hyperlink:
                    parts = re.split(f'({re.escape(hyperlink)})', text, 1)
                    
                    if parts[0]:
                        text_run = para.add_run(parts[0])
                        current_element.addnext(text_run._element)
                        current_element = text_run._element
                    
                    link_element = create_hyperlink_element(para, hyperlink, hyperlink, document)
                    current_element.addnext(link_element)
                    current_element = link_element

                    if len(parts) > 2 and parts[2]:
                        text_run = para.add_run(parts[2])
                        current_element.addnext(text_run._element)
                        current_element = text_run._element

                # --- 3. RENDER AS PLAIN TEXT (with potential bolding) ---
                else:
                    if country and country in text:
                        parts = text.split(country, 1)
                        
                        if parts[0]:
                            text_run = para.add_run(parts[0])
                            current_element.addnext(text_run._element)
                            current_element = text_run._element
                        
                        country_run = para.add_run(country)
                        country_run.bold = True
                        current_element.addnext(country_run._element)
                        current_element = country_run._element
                        
                        if len(parts) > 1 and parts[1]:
                            remaining_run = para.add_run(parts[1])
                            current_element.addnext(remaining_run._element)
                            current_element = remaining_run._element
                    else:
                        text_run = para.add_run(text)
                        current_element.addnext(text_run._element)
                        current_element = text_run._element
                
                # ==========================================================
                # END: MODIFIED LOGIC
                # ==========================================================


def debug_paragraph_structure(para: Paragraph, target_string: str):
    """
    Debug function to understand paragraph structure and identify issues.
    """
    print(f"\n🔍 DEBUGGING PARAGRAPH STRUCTURE")
    print(f"Full paragraph text: '{para.text}'")
    print(f"Target string: '{target_string}'")
    print(f"Target found: {target_string.lower() in para.text.lower()}")
    print(f"Number of runs: {len(para.runs)}")
    
    for i, run in enumerate(para.runs):
        print(f"\nRun {i}:")
        print(f"  Text: '{run.text}'")
        print(f"  Bold: {run.bold}")
        print(f"  Underline: {run.underline}")
        print(f"  Font color: {run.font.color.rgb if run.font.color else 'None'}")
        
        # Check for shading
        is_shaded = is_run_gray_shaded_debug(run)
        is_hyperlink = is_run_hyperlink_debug(run)
        
        print(f"  Is gray shaded: {is_shaded}")
        print(f"  Is hyperlink: {is_hyperlink}")
        print(f"  Should remove: {is_shaded or is_hyperlink}")


def is_run_gray_shaded_debug(run: Run) -> bool:
    """
    Debug version of gray shading detection with detailed output.
    """
    try:
        # Check run properties for shading
        run_pr = run._element.get_or_add_rPr()
        shading_elements = run_pr.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
        
        if shading_elements:
            print(f"    Found shading elements: {len(shading_elements)}")
            for shading in shading_elements:
                fill = shading.get(qn('w:fill'))
                print(f"    Shading fill: {fill}")
                if fill and fill.lower() in ['d9d9d9', 'cccccc', 'gray', 'lightgray', 'auto']:
                    return True
        
        # Check font color for gray
        if run.font.color and hasattr(run.font.color, 'rgb'):
            color = run.font.color.rgb
            print(f"    Font color RGB: {color}")
            gray_colors = [
                RGBColor(128, 128, 128),  # Standard gray
                RGBColor(153, 153, 153),  # Light gray
                RGBColor(102, 102, 102),  # Dark gray
                RGBColor(96, 96, 96),     # Another common gray
                RGBColor(217, 217, 217),  # Very light gray
            ]
            if color in gray_colors:
                return True
                
    except Exception as e:
        print(f"    Error checking shading: {e}")
    
    return False


def is_run_hyperlink_debug(run: Run) -> bool:
    """
    Debug version of hyperlink detection.
    """
    try:
        # Check if run is within a hyperlink element
        run_xml = run._r
        hyperlink_elements = run_xml.xpath('.//w:hyperlink', 
                                         namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
        if hyperlink_elements:
            print(f"    Found hyperlink elements: {len(hyperlink_elements)}")
            return True
            
        # Check hyperlink-style formatting
        if (run.font.color and hasattr(run.font.color, 'rgb') and 
            run.font.color.rgb == RGBColor(0, 0, 255) and run.underline):
            print(f"    Has hyperlink-style formatting (blue + underline)")
            return True
            
    except Exception as e:
        print(f"    Error checking hyperlink: {e}")
    
    return False


def find_runs_to_remove_aggressive(para: Paragraph, target_string: str) -> List[Run]:
    """
    More aggressive approach to find runs that should be removed.
    """
    runs_to_remove = []
    
    # First, find runs containing target text
    target_start = para.text.lower().find(target_string.lower())
    if target_start == -1:
        return runs_to_remove
    
    target_end = target_start + len(target_string)
    
    # Map character positions to runs
    char_pos = 0
    run_ranges = []
    
    for run in para.runs:
        run_start = char_pos
        run_end = char_pos + len(run.text)
        run_ranges.append((run, run_start, run_end))
        char_pos = run_end
    
    print(f"\n🎯 TARGET RANGE: {target_start} to {target_end}")
    
    # Find runs that overlap with target or are adjacent problematic runs
    for run, run_start, run_end in run_ranges:
        should_remove = False
        reason = ""
        
        # Check if run overlaps with target range
        if run_start < target_end and run_end > target_start:
            should_remove = True
            reason = "overlaps with target text"
        
        # Check if it's gray shaded
        elif is_run_gray_shaded_debug(run):
            should_remove = True
            reason = "is gray shaded"
        
        # Check if it's a hyperlink
        elif is_run_hyperlink_debug(run):
            should_remove = True
            reason = "is hyperlink"
        
        # Check if it's a small connector (like period, comma) adjacent to target
        elif (len(run.text.strip()) <= 2 and 
              run.text.strip() in '.,;:' and
              abs(run_start - target_end) <= 5):  # Within 5 chars of target end
            should_remove = True
            reason = "is small connector near target"
        
        print(f"Run {run_start}-{run_end}: '{run.text}' -> Remove: {should_remove} ({reason})")
        
        if should_remove:
            runs_to_remove.append(run)
    
    return runs_to_remove


def build_replacement_components_simple(mapping_row: pd.Series, section_type: str, 
                                       country_delimiter: str = ";") -> List[Dict]:
    """
    Simplified version that focuses on getting the components right.
    """
    print(f"\n🔨 Building replacement components for {section_type}")
    
    components = []
    
    # Get line columns for this section type
    line_columns = [col for col in mapping_row.index 
                   if col.startswith('Line ') and section_type in col]
    
    print(f"Found line columns: {line_columns}")
    
    if not line_columns:
        print("No line columns found")
        return components
    
    # Get hyperlinks and email links
    hyperlinks_col = f'Hyperlinks {section_type}'
    email_col = f'Link for email - {section_type}'
    
    hyperlinks_str = str(mapping_row.get(hyperlinks_col, '')).strip()
    email_str = str(mapping_row.get(email_col, '')).strip()
    
    print(f"Hyperlinks: '{hyperlinks_str}'")
    print(f"Emails: '{email_str}'")
    
    # Parse hyperlinks and emails (semicolon separated)
    hyperlinks = [h.strip() for h in hyperlinks_str.split(country_delimiter) 
                 if h.strip() and h.strip().lower() != 'nan']
    emails = [e.strip() for e in email_str.split(country_delimiter) 
             if e.strip() and e.strip().lower() != 'nan']
    
    print(f"Parsed hyperlinks: {hyperlinks}")
    print(f"Parsed emails: {emails}")
    
    # Sort line columns by number
    def extract_line_number(col_name):
        match = re.search(r'Line (\d+)', col_name)
        return int(match.group(1)) if match else 999
    
    sorted_columns = sorted(line_columns, key=extract_line_number)
    
    # Find Line 1 to get countries
    line_1_col = None
    for col in sorted_columns:
        if extract_line_number(col) == 1:
            line_1_col = col
            break
    
    if not line_1_col:
        print("No Line 1 column found")
        return components
    
    line_1_text = str(mapping_row.get(line_1_col, '')).strip()
    print(f"Line 1 text: '{line_1_text}'")
    
    if not line_1_text or line_1_text.lower() == 'nan':
        print("Line 1 text is empty")
        return components
    
    # Get countries from dedicated bold country column
    bold_countries_col = f'Line 1 - Country names to be bolded - {section_type}'
    bold_countries_str = str(mapping_row.get(bold_countries_col, '')).strip()
    print(f"Bold countries column: '{bold_countries_col}' = '{bold_countries_str}'")
    
    # Parse countries using comma/semicolon delimiter
    if bold_countries_str and bold_countries_str.lower() != 'nan':
        # Try comma first (as seen in mapping file), then semicolon as fallback
        if ',' in bold_countries_str:
            countries = [c.strip() for c in bold_countries_str.split(',') if c.strip()]
        else:
            countries = [c.strip() for c in bold_countries_str.split(country_delimiter) if c.strip()]
        print(f"Countries from bold column: {countries}")
    else:
        # Fallback: extract from line text (backwards compatibility)
        countries = [c.strip() for c in line_1_text.split(country_delimiter) if c.strip()]
        print(f"Countries from fallback (line text): {countries}")
    
    if not countries:
        print("No countries found")
        return components
    
    # Process each line
    for col in sorted_columns:
        line_num = extract_line_number(col)
        content = str(mapping_row.get(col, '')).strip()
        
        print(f"Processing Line {line_num}: '{content}'")
        
        if not content or content.lower() == 'nan':
            continue
        
        # Split content by countries using semicolon delimiter
        parts = [p.strip() for p in content.split(country_delimiter)]
        print(f"  Split into parts: {parts}")
        
        for i, country in enumerate(countries):
            if i < len(parts) and parts[i]:
                text = parts[i]
                
                # Determine links for this country position
                hyperlink = hyperlinks[i] if i < len(hyperlinks) else None
                email = emails[i] if i < len(emails) else None
                
                component = {
                    'line': line_num,
                    'country': country,
                    'text': text,
                    'hyperlink': hyperlink,
                    'email': email
                }
                
                components.append(component)
                print(f"  Added component: {component}")
    
    print(f"Total components built: {len(components)}")
    return components


def insert_replacement_simple(para: Paragraph, insertion_point: int, components: List[Dict], 
                            section_type: str, mapping_row: pd.Series, country_delimiter: str = ";"):
    """
    Simplified insertion that adds text at the insertion point.
    """
    print(f"\n📝 INSERTING REPLACEMENT at position {insertion_point}")
    print(f"Components to insert: {len(components)}")
    
    # Group components by line
    lines = {}
    for comp in components:
        line_num = comp['line']
        if line_num not in lines:
            lines[line_num] = []
        lines[line_num].append(comp)
    
    # Build replacement text
    replacement_text = ""
    
    for line_idx, line_num in enumerate(sorted(lines.keys())):
        line_components = lines[line_num]
        
        # Add line break before non-first lines
        if line_idx > 0:
            replacement_text += "\n"
        
        # Add components for this line
        line_texts = []
        for comp in line_components:
            text = comp['text']
            country = comp['country']
            hyperlink = comp.get('hyperlink')
            email = comp.get('email')
            
            # For now, just add the text (we'll enhance formatting later)
            component_text = text
            if hyperlink:
                component_text += f" {hyperlink}"
            if email:
                component_text += f" {email}"
            
            line_texts.append(component_text)
        
        # Join with semicolon delimiter
        replacement_text += f"{country_delimiter} ".join(line_texts)
    
    # For PL sections, append the additional text
    if section_type == "PL":
        additional_text = str(mapping_row.get('Text to be appended after National reporting system PL', '')).strip()
        if additional_text and additional_text.lower() != 'nan':
            replacement_text += f"\n\n{additional_text}"
    
    print(f"Replacement text: '{replacement_text}'")
    
    # Simple insertion - add a new run with the replacement text
    new_run = para.add_run(replacement_text)
    print("✅ Replacement text inserted")
    
    return True


def run_annex_update_v2(doc: Document, mapping_row: pd.Series, section_type: str, 
                       cached_components: Optional[List] = None, 
                       country_delimiter: str = ";") -> Tuple[bool, Optional[List]]:
    """Update national reporting systems in SmPC or PL sections."""
    # Get the target text to find and replace
    target_col = f'Original text national reporting - {section_type}'
    target_string = str(mapping_row.get(target_col, '')).strip()
    
    if ":" in target_string:
        target_string = target_string.split(':', 1)[-1].strip()

    if not target_string or target_string.lower() == 'nan':
        return False, None

    # Get replacement components
    components = get_replacement_components(mapping_row, section_type, cached_components, country_delimiter)
    
    if not components:
        return False, None
    
    # Find and update the target text
    found = False
    for para in doc.paragraphs:
        if target_string.lower() in para.text.lower():
            
            # Find runs to remove - enhanced with XML-based hyperlink handling
            runs_to_remove = find_runs_to_remove(para, target_string)

            if runs_to_remove:
                # Remove only the identified runs (traditional approach)
                print(f"Removing {len(runs_to_remove)} specific runs...")
                for run in runs_to_remove:
                    try:
                        run._element.getparent().remove(run._element)
                        print(f"  Removed: '{run.text[:30]}...'")
                    except Exception as e:
                        print(f"  Error removing run: {e}")

                # Check remaining text
                remaining_text = para.text.strip()
                print(f"Text after removal: '{remaining_text}'")
            else:
                # Empty list could mean XML removal was already done, or no runs to remove
                remaining_text = para.text.strip()
                if target_string.lower() in remaining_text.lower():
                    print(f"Target still present - XML removal may have failed")
                else:
                    print(f"XML-based removal completed - proceeding with insertion")

            # Insert formatted replacement at the end of the paragraph (ALWAYS after removal)
            try:
                insertion_point = len(para.runs)

                insert_formatted_replacement_surgically(
                    para, insertion_point, components,
                    country_delimiter=country_delimiter, document=doc
                )

                # Insertion successful

                # For PL sections, append additional text
                if section_type == "PL":
                    additional_text = str(mapping_row.get('Text to be appended after National reporting system PL', '')).strip()
                    if additional_text and additional_text.lower() != 'nan':
                        para.add_run(f"\n\n{additional_text}")

            except Exception as e:
                print(f"Error during insertion: {e}")
                import traceback
                traceback.print_exc()
                # Return False but still return components (not the error message)
                return False, components

            found = True
            break
    
    return found, components


def update_document_with_fixed_smpc_blocks(doc: Document, mapping_row: pd.Series) -> Tuple[bool, List[str]]:
    """
    Main function to update document with fixed SmPC block handling.
    
    This function coordinates the updates for both SmPC and PL sections
    using the fixed functions that properly handle semicolon delimiters
    and selective text replacement.
    """
    updates_applied = []
    total_success = False
    
    try:
        # Get the correct country delimiter from config (default to semicolon)
        country_delimiter = ";"  # This should come from ProcessingConfig
        
        # 1. Update SmPC national reporting systems
        smpc_success, smpc_components = run_annex_update_v2(
            doc, mapping_row, "SmPC", None, country_delimiter=country_delimiter
        )
        if smpc_success:
            updates_applied.append("SmPC national reporting")
            total_success = True

        # 2. Update PL national reporting systems
        pl_success, _ = run_annex_update_v2(
            doc, mapping_row, "PL", smpc_components if smpc_success else None, country_delimiter=country_delimiter
        )
        if pl_success:
            updates_applied.append("PL national reporting")
            total_success = True
        
        return total_success, updates_applied
        
    except Exception as e:
        raise Exception(f"Failed to apply SmPC block updates: {e}")


def handle_pl_additional_text(para: Paragraph, mapping_row: pd.Series) -> bool:
    """
    Handle the additional text that needs to be appended after PL national reporting system.
    
    This text comes from the "Text to be appended after National reporting system PL" column
    and provides additional safety reporting information.
    """
    additional_text = str(mapping_row.get('Text to be appended after National reporting system PL', '')).strip()
    
    if not additional_text or additional_text.lower() == 'nan':
        return False
    
    # Add spacing and the additional text
    para.add_run('\n\n')
    
    # Add the additional text with appropriate formatting
    additional_run = para.add_run(additional_text)
    # You can customize formatting here if needed (e.g., italic, different color, etc.)
    
    return True


def create_pl_replacement_block(mapping_row: pd.Series, country_delimiter: str = ";") -> str:
    """
    Create the complete PL replacement block including the main content and additional text.
    
    This handles the case where PL uses block format rather than line-by-line format.
    """
    # Get main PL content
    main_content = str(mapping_row.get('National reporting system PL', '')).strip()
    
    # Get additional text
    additional_text = str(mapping_row.get('Text to be appended after National reporting system PL', '')).strip()
    
    # Combine them
    full_content = []
    
    if main_content and main_content.lower() != 'nan':
        full_content.append(main_content)
    
    if additional_text and additional_text.lower() != 'nan':
        full_content.append(additional_text)
    
    return '\n\n'.join(full_content) if full_content else ''


def _find_section_10_header(doc: Document, date_header: str) -> Optional[int]:
    """
    Find the paragraph index containing the Section 10 header.

    Uses the language-specific date_header from the mapping file to identify
    the correct Section 10 header paragraph.

    Args:
        doc: Document object
        date_header: Language-specific header text from mapping file
                    (e.g., 'Annex I Date Text' column)

    Returns:
        int: Paragraph index if found, None otherwise
    """
    date_header_lower = date_header.lower().strip()

    # Store potential matches with scores for better selection
    potential_matches = []

    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        text_lower = text.lower()

        # Primary approach: Look for paragraph that contains the exact
        # language-specific date header from the mapping file
        if date_header_lower in text_lower:
            # Additional validation: should also contain "10" to confirm it's Section 10
            if '10' in text_lower:
                print(f"✅ Found Section 10 header (exact match) at paragraph {idx}: '{text[:80]}...'")
                return idx  # Exact match - return immediately

        # Secondary approach: Look for section number pattern + length validation
        # This handles cases where the header might be slightly different
        if '10.' in text_lower:
            # Must be at start of text or after whitespace (proper section header)
            if text_lower.startswith('10.') or ' 10.' in text_lower:
                # Must be substantial content (not just "10." alone)
                if len(text_lower) > 15:  # More restrictive length requirement
                    # Additional validation: should contain date-related keywords
                    date_keywords = ['date', 'first', 'authorisation', 'authorization', 'renewal', 'fecha', 'première', 'première']
                    if any(keyword in text_lower for keyword in date_keywords):
                        score = len(text_lower)  # Longer headers get higher priority
                        potential_matches.append((idx, score, text))

    # If we have potential matches, return the one with highest score (longest text)
    if potential_matches:
        best_match = max(potential_matches, key=lambda x: x[1])
        idx, score, text = best_match
        print(f"✅ Found Section 10 header (pattern match) at paragraph {idx}: '{text[:80]}...'")
        return idx

    print(f"❌ Could not find Section 10 header with text: '{date_header}'")
    return None


def _insert_date_after_header(doc: Document, header_index: int, formatted_date: str) -> bool:
    """
    Insert date content in the paragraph immediately following the header.

    Args:
        doc: Document object
        header_index: Index of the Section 10 header paragraph
        formatted_date: Formatted date string to insert

    Returns:
        bool: True if successful
    """
    try:
        # Check if next paragraph exists
        if header_index + 1 < len(doc.paragraphs):
            next_para = doc.paragraphs[header_index + 1]

            # Clear existing content and insert date
            next_para.clear()
            run = next_para.add_run(formatted_date)
            run.bold = False
            print(f"✅ Date inserted in existing paragraph {header_index + 1}")
            return True

        else:
            # Create new paragraph after header
            # We need to use a different approach since python-docx doesn't have direct insertion
            # Let's add a paragraph at the end and then move content around
            new_para = doc.add_paragraph()
            run = new_para.add_run(formatted_date)
            run.bold = False

            # Move the new paragraph to the correct position
            # This is a workaround - we add at end then reorder
            paragraphs = doc.paragraphs
            if len(paragraphs) > header_index + 2:
                # Move the last paragraph (our new one) to position header_index + 1
                last_para_element = paragraphs[-1]._element
                header_para_element = paragraphs[header_index]._element
                header_para_element.addnext(last_para_element)

            print(f"✅ Created new paragraph after header at position {header_index + 1}")
            return True

    except Exception as e:
        print(f"❌ Error inserting date after header: {e}")
        return False


def update_section_10_date(doc: Document, mapping_row: pd.Series, mapping_file_path: Optional[str] = None) -> bool:
    """
    Update date in Annex I Section 10 - ENHANCED VERSION.

    This function now:
    1. Finds the Section 10 header using language-specific text from mapping file
    2. Preserves the header paragraph unchanged
    3. Inserts the formatted date in the next paragraph after the header
    """
    country = mapping_row.get('Country', '')
    date_header = mapping_row.get('Annex I Date Text', 'Date of first authorisation/renewal of the authorisation')

    print(f"🔧 DEBUG: update_section_10_date called for {country}")
    print(f"   Date header text: '{date_header}'")

    if not country:
        print("❌ DEBUG: No country found, returning False")
        return False

    # Get formatted date (existing logic works fine)
    try:
        # Ensure date formatter is initialized
        if mapping_file_path:
            try:
                get_date_formatter()
            except RuntimeError:
                initialize_date_formatter(mapping_file_path)

        formatted_date = format_date_for_country(country, 'annex_i')
        print(f"✅ Formatted date: '{formatted_date}'")
    except Exception as e:
        print(f"⚠️ Date formatting failed, using fallback: {e}")
        date_format = mapping_row.get('Annex I Date Format', '')
        formatted_date = datetime.now().strftime("%d %B %Y")

    # NEW: Find Section 10 header specifically using mapping file data
    header_index = _find_section_10_header(doc, date_header)
    if header_index is None:
        print(f"❌ Could not find Section 10 header with text: '{date_header}'")
        return False

    # NEW: Insert date in next paragraph, preserving header
    success = _insert_date_after_header(doc, header_index, formatted_date)
    if success:
        print(f"✅ Section 10 date inserted successfully for {country}")
    else:
        print(f"❌ Failed to insert Section 10 date for {country}")

    return success

def update_annex_iiib_date(doc: Document, mapping_row: pd.Series, mapping_file_path: Optional[str] = None) -> bool:
    """Update date in Annex IIIB Section 6."""
    country = mapping_row.get('Country', '')
    date_text = mapping_row.get('Annex IIIB Date Text', 'This leaflet was last revised in')
    
    if not country:
        return False
    
    try:
        # Ensure date formatter is initialized
        if mapping_file_path:
            try:
                get_date_formatter()
            except RuntimeError:
                initialize_date_formatter(mapping_file_path)

        formatted_date = format_date_for_country(country, 'annex_iiib')
    except Exception:
        formatted_date = datetime.now().strftime("%d %B %Y")
    
    found = False
    for para in doc.paragraphs:
        text_lower = para.text.lower()
        
        if (date_text.lower() in text_lower or
            'leaflet was last revised' in text_lower or
            'dernière approbation' in text_lower or
            'última revisión' in text_lower):
            
            para.clear()
            run = para.add_run(f"{date_text} {formatted_date}")
            run.bold = False
            found = True
            break
    
    return found

def filter_local_representatives(doc: Document, mapping_row: pd.Series) -> bool:
    """
    Filter local representatives in Section 6 of Annex IIIB to keep only applicable ones.

    Enhanced to support both table-based and paragraph-based local rep processing.
    First attempts table processing (modern approach), then falls back to paragraph
    processing (legacy compatibility).

    Args:
        doc: Document object to modify
        mapping_row: Row from mapping file containing local rep filtering data

    Returns:
        bool: True if local representatives were successfully filtered
    """
    country = str(mapping_row.get('Country', '')).strip()

    print(f"🔧 DEBUG: filter_local_representatives called")
    print(f"   Country extracted: '{country}'")

    if not country:
        print("❌ DEBUG: No country found, returning False")
        return False

    # Try table-based processing first (new capability)
    print("🔧 DEBUG: Attempting table-based processing...")
    try:
        from .local_rep_table_processor import LocalRepTableProcessor

        table_processor = LocalRepTableProcessor()
        table_result = table_processor.process_local_rep_table(doc, mapping_row)
        print(f"🔧 DEBUG: Table processing result: {table_result}")

        if table_result:
            print("✅ Local representatives processed using table-based approach")
            return True
        else:
            print("⚠️  Table processing returned False, trying paragraph fallback...")

    except Exception as e:
        print(f"⚠️  Table processing failed with exception, falling back to paragraph processing: {e}")
        import traceback
        traceback.print_exc()

    # Fallback to existing paragraph processing (preserved for compatibility)
    print("🔧 DEBUG: Attempting paragraph-based processing...")
    paragraph_result = _filter_local_representatives_paragraphs(doc, mapping_row)
    print(f"🔧 DEBUG: Paragraph processing result: {paragraph_result}")
    return paragraph_result


def _filter_local_representatives_paragraphs(doc: Document, mapping_row: pd.Series) -> bool:
    """
    Legacy paragraph-based local representative filtering.

    Preserved for backward compatibility and documents that don't use table format.
    """
    # Get applicable local representatives for this language/country
    applicable_reps = str(mapping_row.get('Local Representative', '')).strip()
    country = str(mapping_row.get('Country', '')).strip()
    language = str(mapping_row.get('Language', '')).strip()

    if not applicable_reps or applicable_reps.lower() == 'nan':
        return False
    # Parse countries that should be bold formatted
    bold_countries_str = str(mapping_row.get('Country names to be bolded - Local Reps', '')).strip()
    bold_countries = [c.strip() for c in bold_countries_str.split(',')
                     if c.strip() and c.strip().lower() != 'nan']

    found = False
    in_section_6 = False
    in_local_rep_section = False
    paragraphs_to_remove = []

    # Phase 1: Identify Section 6 and locate local representative paragraphs
    for idx, para in enumerate(doc.paragraphs):
        text_lower = para.text.lower()

        # Check if we're entering Section 6
        if ('6.' in text_lower and 'contents of the pack' in text_lower) or \
           ('section 6' in text_lower) or \
           ('contenu de l\'emballage' in text_lower):
            in_section_6 = True
            continue

        # Check if we've left Section 6 (entering next section)
        if in_section_6 and _is_section_header(para.text):
            break

        # Look for local representative section header
        if in_section_6 and ('local representative' in text_lower or 'représentant local' in text_lower):
            in_local_rep_section = True
            continue

        # Collect local rep entries to potentially remove
        if in_local_rep_section:
            # Stop if we hit marketing auth holder or other major section
            if ('marketing authorisation holder' in text_lower or
                'manufacturing authorisation holder' in text_lower or
                'this leaflet was last revised' in text_lower or
                _is_section_header(para.text)):
                break

            # Check if this paragraph contains a local rep entry
            if _contains_country_local_rep_entry(para.text):
                # Determine if this local rep should be kept or removed
                if not _should_keep_local_rep_entry(para.text, country, applicable_reps):
                    paragraphs_to_remove.append(idx)
                else:
                    # This is the applicable local rep - apply bold formatting
                    _apply_bold_formatting_to_paragraph(para, bold_countries)
                    found = True

    # Phase 2: Remove non-applicable local representative paragraphs
    # Remove in reverse order to maintain correct indices
    for idx in reversed(paragraphs_to_remove):
        # Get the paragraph to remove
        para_to_remove = doc.paragraphs[idx]
        # Remove the paragraph's content
        para_to_remove.clear()

    return found


def _contains_country_local_rep_entry(text: str) -> bool:
    """
    Check if paragraph contains a country-specific local representative entry.
    These typically start with a country name followed by a colon.
    """
    text_stripped = text.strip()
    if not text_stripped:
        return False
        
    # Look for patterns like "Germany:", "France:", "Ireland:", etc.
    import re
    # Match country name at start of line followed by colon
    return bool(re.match(r'^[A-Za-z\s]+:', text_stripped))


def _should_keep_local_rep_entry(para_text: str, target_country: str, applicable_reps: str) -> bool:
    """
    Determine if a local representative entry should be kept based on the target country.
    """
    # Check if the paragraph contains the target country
    return target_country.lower() in para_text.lower()


def _apply_bold_formatting_to_paragraph(para: Paragraph, bold_countries: list) -> None:
    """
    Apply bold formatting to country names within an existing paragraph.
    """
    if not bold_countries:
        return
        
    # Get the current text
    current_text = para.text
    
    # Clear and rebuild the paragraph with proper formatting
    para.clear()
    
    remaining_text = current_text
    for country in bold_countries:
        if country.lower() in remaining_text.lower():
            # Find the country name (case-insensitive)
            import re
            match = re.search(re.escape(country), remaining_text, re.IGNORECASE)
            if match:
                # Add text before country name
                before_text = remaining_text[:match.start()]
                if before_text:
                    para.add_run(before_text)
                
                # Add country name with bold formatting
                bold_run = para.add_run(match.group())
                bold_run.bold = True
                
                # Continue with remaining text
                remaining_text = remaining_text[match.end():]
    
    # Add any remaining text
    if remaining_text:
        para.add_run(remaining_text)


def _is_section_header(text: str) -> bool:
    """Check if text appears to be a section header (like "7.", "8.", etc.)"""
    text_lower = text.strip().lower()
    # Look for patterns like "7.", "section 7", etc.
    import re
    return bool(re.match(r'^\s*\d+\.', text) or re.match(r'^\s*section\s+\d+', text_lower))


def update_local_representatives(doc: Document, mapping_row: pd.Series) -> bool:
    """
    Update local representatives - wrapper function with debug logging.

    This function is called by the main processing workflow and delegates to
    the enhanced filter_local_representatives function.
    """
    country = mapping_row.get('Country', 'MISSING')
    language = mapping_row.get('Language', 'MISSING')
    local_rep_data = mapping_row.get('Local Representative', 'MISSING')

    print(f"🔧 DEBUG: Starting local rep processing")
    print(f"   Country: {country}")
    print(f"   Language: {language}")
    print(f"   Local Rep Data: {local_rep_data[:100] if isinstance(local_rep_data, str) else local_rep_data}")
    print(f"   Document has {len(doc.tables)} tables")

    try:
        result = filter_local_representatives(doc, mapping_row)
        print(f"🔧 DEBUG: Local rep processing result: {result}")
        if result:
            print("✅ Local rep processing succeeded!")
        else:
            print("❌ Local rep processing returned False")
        return result
    except Exception as e:
        print(f"💥 DEBUG: Local rep processing failed with exception: {e}")
        import traceback
        traceback.print_exc()
        return False




# Legacy function for backwards compatibility - now calls the new filtering function
def update_local_representatives(doc: Document, mapping_row: pd.Series) -> bool:
    """
    Legacy function for backwards compatibility.
    Now calls the new filter_local_representatives function.
    """
    return filter_local_representatives(doc, mapping_row)

# ============================================================================= 
# Split Annexes Workflow
# =============================================================================

def split_annexes(source_path: str, output_dir: str, language: str, country: str, mapping_row: pd.Series) -> Tuple[str, str]:
    """
    Split a combined SmPC document into Annex I and Annex IIIB documents.

    ENHANCED VERSION: Uses clone-and-prune approach for perfect document preservation.
    This preserves ALL formatting, hyperlinks, headers, footers, and scaffolding.
    """
    print("🚀 Using enhanced clone-and-prune document splitting")

    try:
        # Get actual headers from mapping file (what's really in the document)
        annex_i_header = str(mapping_row.get('Annex I Header in country language', 'ANNEX I')).strip()
        annex_iiib_header = str(mapping_row.get('Annex IIIB Header in country language', 'ANNEX III')).strip()

        print(f"📋 Using headers from mapping file:")
        print(f"   Annex I: '{annex_i_header}'")
        print(f"   Annex IIIB: '{annex_iiib_header}'")

        # Use clone-and-prune approach with actual document headers
        result_paths = clone_and_split_document(
            source_path=source_path,
            output_dir=output_dir,
            country_code=country,
            target_annexes=[annex_i_header, annex_iiib_header],  # Use actual headers from mapping
            language=language,
            mapping_row=mapping_row
        )

        # Extract paths for return (maintain backward compatibility)
        # Map back to expected keys
        annex_i_path = result_paths.get(annex_i_header)
        annex_iiib_path = result_paths.get(annex_iiib_header)

        if not annex_i_path or not annex_iiib_path:
            raise ValueError(f"Failed to split document - could not find required annexes. Found: {list(result_paths.keys())}")

        print(f"✅ Successfully split documents using clone-and-prune:")
        print(f"   ANNEX I: {annex_i_path}")
        print(f"   ANNEX IIIB: {annex_iiib_path}")

        return annex_i_path, annex_iiib_path

    except Exception as e:
        print(f"❌ Clone-and-prune error: {e}")
        raise ProcessingError(f"Document splitting failed: {e}") from e






def debug_three_header_structure(source_path: str, mapping_row: pd.Series) -> None:
    """
    Debug the three-header approach to validate header detection.
    
    Args:
        source_path: Path to document to analyze
        mapping_row: Mapping row with header information
    """
    
    doc = Document(source_path)
    country = mapping_row.get('Country', 'Unknown')
    language = mapping_row.get('Language', 'Unknown')
    
    # Get all three headers
    annex_i_header = str(mapping_row.get('Annex I Header in country language', '')).strip()
    annex_ii_header = str(mapping_row.get('Annex II Header in country language', '')).strip()
    annex_iiib_header = str(mapping_row.get('Annex IIIB Header in country language', '')).strip()
    
    print(f"\n🔍 THREE-HEADER DEBUGGING")
    print(f"File: {Path(source_path).name}")
    print(f"Country: {country} ({language})")
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    print(f"Expected Annex I header: '{annex_i_header}'")
    print(f"Expected Annex II header: '{annex_ii_header}'")
    print(f"Expected Annex IIIB header: '{annex_iiib_header}'")
    print("=" * 80)
    
    # Find all matches for each header
    annex_i_matches = []
    annex_ii_matches = []
    annex_iiib_matches = []
    
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        if _is_header_match(text, annex_i_header):
            annex_i_matches.append({'index': idx, 'text': text})
        
        if _is_header_match(text, annex_ii_header):
            annex_ii_matches.append({'index': idx, 'text': text})
        
        if _is_header_match(text, annex_iiib_header):
            annex_iiib_matches.append({'index': idx, 'text': text})
    
    # Display results
    print(f"📌 HEADER MATCHES FOUND:")
    
    print(f"\nAnnex I ('{annex_i_header}'):")
    if annex_i_matches:
        for match in annex_i_matches:
            print(f"  Para {match['index']}: '{match['text'][:60]}...'")
    else:
        print(f"  ❌ No matches found")
    
    print(f"\nAnnex II ('{annex_ii_header}'):")
    if annex_ii_matches:
        for match in annex_ii_matches:
            print(f"  Para {match['index']}: '{match['text'][:60]}...'")
    else:
        print(f"  ❌ No matches found")
    
    print(f"\nAnnex IIIB ('{annex_iiib_header}'):")
    if annex_iiib_matches:
        for match in annex_iiib_matches:
            print(f"  Para {match['index']}: '{match['text'][:60]}...'")
    else:
        print(f"  ❌ No matches found")
    
    # Validate structure if all headers found
    if annex_i_matches and annex_ii_matches and annex_iiib_matches:
        best_i = annex_i_matches[0]['index']
        best_ii = annex_ii_matches[0]['index'] 
        best_iiib = annex_iiib_matches[0]['index']
        
        print(f"\n📊 PROPOSED STRUCTURE:")
        print(f"   Annex I: paragraphs {best_i} to {best_ii-1} ({best_ii - best_i} paragraphs)")
        print(f"   Annex II: paragraphs {best_ii} to {best_iiib-1} ({best_iiib - best_ii} paragraphs)")
        print(f"   Annex IIIB: paragraphs {best_iiib} to end ({len(doc.paragraphs) - best_iiib} paragraphs)")
        
        if best_i >= best_ii or best_ii >= best_iiib:
            print(f"  ❌ STRUCTURE ERROR: Headers not in correct order!")
        else:
            print(f"  ✅ Structure looks good!")
    else:
        print(f"\n❌ Cannot validate structure - missing header matches")



def _is_header_match(paragraph_text: str, header_text: str) -> bool:
    """Check if a paragraph text matches a header with precise word-boundary matching."""
    para_normalized = _normalize_text_for_matching(paragraph_text)
    header_normalized = _normalize_text_for_matching(header_text)
    
    # Exact match after normalization
    if para_normalized == header_normalized:
        return True
    
    # Check if header is contained in paragraph (word boundary matching)
    if _contains_as_words(para_normalized, header_normalized):
        return True
    
    # For very similar headers (like "annex i" vs "annex ii"), be more strict
    if _are_similar_headers(para_normalized, header_normalized):
        return False
    
    # Check if paragraph starts with header (common case)
    if para_normalized.startswith(header_normalized + " "):
        return True
    
    return False

def _contains_as_words(text: str, search_term: str) -> bool:
    """
    Check if search_term exists as complete words in text, not just as substring.
    This prevents "annex i" from matching "annex ii".
    """
    import re
    
    # Escape special regex characters in search term
    escaped_term = re.escape(search_term)
    
    # Use word boundaries to ensure complete word matching
    # \b ensures we match complete words, not substrings
    pattern = r'\b' + escaped_term + r'\b'
    
    return bool(re.search(pattern, text, re.IGNORECASE))


def _are_similar_headers(text1: str, text2: str) -> bool:
    """
    Check if two texts are similar annex headers that could be confused.
    Returns True if they're similar enough that we should be strict about matching.
    
    Uses comprehensive patterns based on actual mapping data from all supported languages.
    """
    
    # Comprehensive annex header base words from mapping data
    annex_base_words = [
        'bijlage',      # Dutch
        'annexe',       # French  
        'anhang',       # German
        'lisa',         # Estonian
        'παραρτημα',    # Greek
        'pielikums',    # Latvian
        'priedas',      # Lithuanian
        'anexo',        # Spanish/Portuguese
        'prilog',       # Croatian
        'priloga',      # Slovenian
        'liite',        # Finnish
        'bilaga',       # Swedish
        'allegato',     # Italian
        'annex',        # English
        'anness',       # Maltese
        'bilag',        # Danish
        'viðauki',      # Icelandic
        'vedlegg',      # Norwegian
        'příloha',      # Czech
        'aneks',        # Polish
        'príloha',      # Slovak
        'приложение',   # Bulgarian
        'melléklet',    # Hungarian
        'anexa',        # Romanian
    ]
    
    # Roman numeral patterns (including Greek variants)
    roman_patterns = [
        r'[ivx]+',          # Standard: i, ii, iii, iv, v
        r'[ιυχ]+',          # Greek: ι, ιι, ιιι
        r'\d+',             # Arabic numbers: 1, 2, 3 (backup)
    ]
    
    # Build comprehensive patterns for both word-first and number-first structures
    all_patterns = []
    
    for base_word in annex_base_words:
        for roman_pattern in roman_patterns:
            # Pattern 1: Word first (e.g., "ANNEXE I", "BIJLAGE II")
            all_patterns.append(rf'{re.escape(base_word)}\s*\.?\s*{roman_pattern}\.?')
            
            # Pattern 2: Number first (e.g., "I LISA", "II LISA") 
            all_patterns.append(rf'{roman_pattern}\.?\s+{re.escape(base_word)}')
            
            # Pattern 3: Number with period first (e.g., "I. MELLÉKLET")
            all_patterns.append(rf'{roman_pattern}\.\s*{re.escape(base_word)}')
    
    # Check if both texts match any of the same patterns
    for pattern in all_patterns:
        if (re.search(pattern, text1, re.IGNORECASE) and 
            re.search(pattern, text2, re.IGNORECASE)):
            return True
    
    # Additional check: if both contain the same base word, they're similar
    text1_lower = text1.lower()
    text2_lower = text2.lower()
    
    for base_word in annex_base_words:
        if base_word.lower() in text1_lower and base_word.lower() in text2_lower:
            return True
    
    return False

def _normalize_text_for_matching(text: str) -> str:
    """
    Normalize text for header matching by removing inconsistencies.
    
    Args:
        text: Raw text to normalize
        
    Returns:
        Normalized text suitable for comparison
    """
    
    # Convert to lowercase
    normalized = text.lower()
    
    # Remove extra whitespace and normalize spaces
    normalized = re.sub(r'\s+', ' ', normalized).strip()
    
    # Remove common punctuation that might vary
    normalized = re.sub(r'[.,;:!?""''""()]', '', normalized)
    
    # Remove common formatting artifacts
    normalized = re.sub(r'[\r\n\t]', ' ', normalized)
    normalized = re.sub(r'\s+', ' ', normalized).strip()
    
    return normalized



def split_annexes_original(source_path: str, output_dir: str, language: str, country: str, mapping_row: pd.Series) -> Tuple[str, str]:
    """
    Original splitting logic as fallback.
    This is the existing implementation for compatibility.
    """
    
    doc = Document(source_path)
    
    # Create new documents
    annex_i_doc = Document()
    annex_iiib_doc = Document()
    
    current_section = None
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # Determine which section we're in using original logic
        if 'ANNEX I' in text.upper() or 'SUMMARY OF PRODUCT CHARACTERISTICS' in text.upper():
            current_section = 'annex_i'
        elif 'ANNEX III' in text.upper() or 'PACKAGE LEAFLET' in text.upper():
            current_section = 'annex_iiib'
        
        # Copy paragraph to appropriate document
        if current_section == 'annex_i':
            copy_paragraph(annex_i_doc, para)
        elif current_section == 'annex_iiib':
            copy_paragraph(annex_iiib_doc, para)
    
    # Generate output paths
    base_name = Path(source_path).stem
    annex_i_filename = generate_output_filename(base_name, language, country, "annex_i")
    annex_iiib_filename = generate_output_filename(base_name, language, country, "annex_iiib")
    
    annex_i_path = os.path.join(output_dir, annex_i_filename)
    annex_iiib_path = os.path.join(output_dir, annex_iiib_filename)
    
    # Save documents
    annex_i_doc.save(annex_i_path)
    annex_iiib_doc.save(annex_iiib_path)
    
    return annex_i_path, annex_iiib_path

# ============================================================================= 
# ENHANCED PROCESSOR CLASSES
# =============================================================================

class FileManager:
    """Handles file operations and path management."""
    
    def __init__(self, base_folder: Path, config: ProcessingConfig):
        self.base_folder = base_folder
        self.config = config
        self.logger = logging.getLogger(f"{__name__}.FileManager")
    
    def setup_output_directories(self) -> Tuple[Path, Path]:
        """Create and return paths for output directories."""
        split_dir = self.base_folder / DirectoryNames.SPLIT_DOCS
        pdf_dir = self.base_folder / DirectoryNames.PDF_DOCS
        
        try:
            os.makedirs(split_dir, exist_ok=True)
            os.makedirs(pdf_dir, exist_ok=True)
            return split_dir, pdf_dir
        except OSError as e:
            raise ProcessingError(f"Failed to create output directories: {e}")
    
    def discover_processable_documents(self) -> List[Path]:
        """Find all valid Word documents that can be processed."""
        if not self.base_folder.is_dir():
            raise ValidationError(f"Folder does not exist: {self.base_folder}")
        
        documents = []
        for file_path in self.base_folder.iterdir():
            if self._is_processable_document(file_path):
                documents.append(file_path)
        
        return documents
    
    def _is_processable_document(self, file_path: Path) -> bool:
        """Check if a file is a valid document for processing."""
        if file_path.suffix.lower() != ".docx":
            return False
        if file_path.name.startswith(FileMarkers.TEMP_FILE_PREFIX):
            return False
        if FileMarkers.ANNEX_MARKER in file_path.name:
            return False
        if file_path.name.startswith(FileMarkers.ANNEX_PREFIX):
            return False
        return True
    
    def create_backup(self, file_path: Path) -> Optional[Path]:
        """Create a backup of the original file."""
        if not self.config.create_backups:
            return None
            
        backup_path = file_path.with_suffix(file_path.suffix + DirectoryNames.BACKUP_SUFFIX)
        if backup_path.exists():
            return backup_path
            
        try:
            shutil.copy2(file_path, backup_path)
            return backup_path
        except Exception:
            return None

class DocumentUpdater:
    """Handles document modification operations."""
    
    def __init__(self, config: ProcessingConfig):
        self.config = config
        self.logger = logging.getLogger(f"{__name__}.DocumentUpdater")
    
    def apply_all_updates(self, doc: Document, mapping_row: pd.Series, mapping_file_path: Optional[str] = None) -> Tuple[bool, List[str]]:
        """Apply all required updates to a document."""
        updates_applied = []
        total_success = False
        
        try:
            # 1. Update national reporting systems  ⬅️ **REPLACE WITH THIS**
            smpc_success, smpc_updates = update_document_with_fixed_smpc_blocks(doc, mapping_row)
            if smpc_success:
                updates_applied.extend(smpc_updates)
                total_success = True
            
            # 2. Update dates
            annex_i_date_success = update_section_10_date(doc, mapping_row, mapping_file_path)
            if annex_i_date_success:
                updates_applied.append("Annex I dates")
                total_success = True

            annex_iiib_date_success = update_annex_iiib_date(doc, mapping_row, mapping_file_path)
            if annex_iiib_date_success:
                updates_applied.append("Annex IIIB dates")
                total_success = True
            
            # 3. Update local representatives
            local_rep_success = update_local_representatives(doc, mapping_row)
            if local_rep_success:
                updates_applied.append("Local representatives")
                total_success = True
                
            return total_success, updates_applied
            
        except Exception as e:
            raise DocumentError(f"Failed to apply document updates: {e}")

class DocumentProcessor:
    """Main document processing orchestrator."""
    
    def __init__(self, config: Optional[ProcessingConfig] = None):
        self.config = config or ProcessingConfig()
        self.stats = ProcessingStats()
        self.logger = self._setup_logging()
        
    def _setup_logging(self) -> logging.Logger:
        """Set up logging configuration."""
        logger = logging.getLogger(__name__)
        logger.setLevel(getattr(logging, self.config.log_level.upper()))
        
        if not logger.handlers:
            handler = logging.StreamHandler()
            formatter = logging.Formatter(
                '%(asctime)s - %(name)s - %(levelname)s - %(message)s'
            )
            handler.setFormatter(formatter)
            logger.addHandler(handler)
        
        return logger
    
    def process_folder(self, folder_path: str, mapping_path: str) -> ProcessingResult:
        """Main entry point for processing a folder of documents."""
        try:
            self.logger.info("=" * 80)
            self.logger.info("🚀 STARTING ENHANCED DOCUMENT PROCESSING")
            self.logger.info("=" * 80)
            
            # Validate inputs
            folder = self._validate_folder_path(folder_path)
            mapping_df = self._load_and_validate_mapping(mapping_path)

            # Initialize date formatter with mapping file
            initialize_date_formatter(mapping_path)
            self.logger.info("✅ Date formatter initialized")
            
            # Setup processing environment
            file_manager = FileManager(folder, self.config)
            split_dir, pdf_dir = file_manager.setup_output_directories()
            
            # Discover documents to process
            documents = file_manager.discover_processable_documents()
            self.stats.input_files_found = len(documents)
            
            if not documents:
                return ProcessingResult(
                    success=False,
                    message="No valid documents found for processing"
                )
            
            # Process each document
            output_files = []
            for document_path in documents:
                try:
                    result = self._process_single_document(
                        document_path, mapping_df, file_manager, split_dir, pdf_dir, mapping_path
                    )
                    output_files.extend(result.output_files)

                except Exception as e:
                    self.logger.error(f"Error processing {document_path.name}: {e}")
                    self.stats.errors_encountered += 1

            # NEW: Batch convert PDFs after all document processing
            if self.config.convert_to_pdf and not self.config.skip_pdf_in_background:
                pdf_files = self._batch_convert_pdfs(pdf_dir)
                output_files.extend(pdf_files)
                self.stats.output_files_created += len(pdf_files)
            elif self.config.convert_to_pdf and self.config.skip_pdf_in_background:
                self.logger.info("📄 PDF conversion skipped (running in background context)")
                self.logger.info(f"📄 {len(getattr(self, '_pending_pdf_conversions', []))} documents queued for manual PDF conversion")

            # Generate final report
            return self._generate_final_result(output_files)
            
        except Exception as e:
            self.logger.error(f"Fatal error in process_folder: {e}")
            return ProcessingResult(
                success=False,
                message=f"Processing failed: {e}",
                errors=[str(e)]
            )
    
    def _validate_folder_path(self, folder_path: str) -> Path:
        """Validate and return folder path."""
        folder = Path(folder_path).resolve()
        if not folder.is_dir():
            raise ValidationError(f"Invalid directory: {folder_path}")
        return folder
    
    def _load_and_validate_mapping(self, mapping_path: str) -> pd.DataFrame:
        """Load and validate mapping file."""
        try:
            mapping_df = load_mapping_table(mapping_path)
            if mapping_df is None or mapping_df.empty:
                raise MappingError(f"Could not load mapping file: {mapping_path}")
            
            self.logger.info(f"Mapping loaded: {len(mapping_df)} configurations")
            return mapping_df
            
        except Exception as e:
            raise MappingError(f"Failed to load mapping file: {e}")
    
    def _process_single_document(
        self,
        document_path: Path,
        mapping_df: pd.DataFrame,
        file_manager: FileManager,
        split_dir: Path,
        pdf_dir: Path,
        mapping_path: str
    ) -> ProcessingResult:
        """Process a single document with all its variants."""
        
        self.logger.info("=" * 60)
        self.logger.info(f"📄 PROCESSING: {document_path.name}")
        self.logger.info("=" * 60)
        
        self.stats.input_files_processed += 1
        
        try:
            # Identify document language and country
            country_code, language_name, country_name = identify_document_country_and_language(str(document_path))
            
            if not language_name:
                error_msg = f"Cannot identify language for {document_path.name}"
                self.logger.error(error_msg)
                return ProcessingResult(success=False, message=error_msg)
            
            self.logger.info(f"Document identified - Language: {language_name}, Country: {country_name}")
            
            # Find mapping rows for this language
            mapping_rows = find_mapping_rows_for_language(mapping_df, language_name)
            if not mapping_rows:
                error_msg = f"No mapping found for language: {language_name}"
                self.logger.error(error_msg)
                return ProcessingResult(success=False, message=error_msg)
            
            self.logger.info(f"Found {len(mapping_rows)} variant(s) to process")
            
            # Create backup
            file_manager.create_backup(document_path)
            
            # Process each variant
            output_files = []
            variant_success_count = 0
            
            for i, mapping_row in enumerate(mapping_rows, 1):
                country = mapping_row['Country']
                self.logger.info(f"🌍 Processing variant {i}/{len(mapping_rows)}: {country}")
                
                try:
                    result = self._process_document_variant(
                        document_path, mapping_row, split_dir, pdf_dir, mapping_path
                    )
                    
                    if result.success:
                        variant_success_count += 1
                        self.stats.variants_successful += 1
                        output_files.extend(result.output_files)
                        self.logger.info(f"✅ Variant {i} completed successfully")
                    else:
                        self.logger.warning(f"⚠️ Variant {i} completed with issues: {result.message}")
                    
                    self.stats.variants_processed += 1
                    
                except Exception as e:
                    self.logger.error(f"❌ Error processing variant {i} ({country}): {e}")
                    self.stats.errors_encountered += 1
            
            # Document summary
            success_rate = (variant_success_count / len(mapping_rows)) * 100 if mapping_rows else 0
            self.logger.info(f"📊 Document Summary: {variant_success_count}/{len(mapping_rows)} variants successful ({success_rate:.1f}%)")
            
            return ProcessingResult(
                success=variant_success_count > 0,
                message=f"Processed {variant_success_count}/{len(mapping_rows)} variants successfully",
                output_files=output_files
            )
            
        except Exception as e:
            self.logger.error(f"Error processing document {document_path.name}: {e}")
            return ProcessingResult(success=False, message=str(e), errors=[str(e)])
    
    def _process_document_variant(
        self,
        document_path: Path,
        mapping_row: pd.Series,
        split_dir: Path,
        pdf_dir: Path,
        mapping_path: str
    ) -> ProcessingResult:
        """Process a single document variant."""
        
        country = mapping_row['Country']
        language = mapping_row['Language']
        
        try:
            # Load document
            doc = Document(str(document_path))
            
            # Apply updates
            updater = DocumentUpdater(self.config)
            updates_made, updates_applied = updater.apply_all_updates(doc, mapping_row, mapping_path)
            
            if not updates_made:
                return ProcessingResult(
                    success=False,
                    message=f"No updates applied for {country} variant"
                )
            
            # Save and process updated document - use synchronous wrapper
            return self._save_and_split_document_sync(
                doc, document_path, mapping_row, split_dir, pdf_dir, updates_applied
            )
            
        except Exception as e:
            raise DocumentError(f"Failed to process variant for {country}: {e}")

    def _save_and_split_document_sync(
        self,
        doc: Document,
        original_path: Path,
        mapping_row: pd.Series,
        split_dir: Path,
        pdf_dir: Path,
        updates_applied: List[str]
    ) -> ProcessingResult:
        """Synchronous wrapper that calls async version."""
        import asyncio

        # Run the async version in a new event loop
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        try:
            return loop.run_until_complete(
                self._save_and_split_document(
                    doc, original_path, mapping_row, split_dir, pdf_dir, updates_applied
                )
            )
        finally:
            loop.close()

    async def _save_and_split_document(
        self,
        doc: Document,
        original_path: Path,
        mapping_row: pd.Series,
        split_dir: Path,
        pdf_dir: Path,
        updates_applied: List[str]
    ) -> ProcessingResult:
        """Save updated document and split into annexes."""
        
        country = mapping_row['Country']
        language = mapping_row['Language']
        output_files = []
        
        try:
            # Generate output filename
            base_name = original_path.stem
            output_filename = generate_output_filename(base_name, language, country, "combined")
            output_path = original_path.parent / output_filename
            
            # Save updated document
            doc.save(str(output_path))
            output_files.append(str(output_path))
            self.logger.info(f"💾 Saved combined document: {output_filename}")
            
            # Split into annexes
            self.logger.info("🔀 Splitting into separate annexes...")
            annex_i_path, annex_iiib_path = split_annexes(
                str(output_path), str(split_dir), language, country, mapping_row
            )
            
            output_files.extend([annex_i_path, annex_iiib_path])
            self.logger.info(f"✅ Split completed")
            
            # Store paths for later PDF conversion (don't convert yet)
            if self.config.convert_to_pdf:
                if not hasattr(self, '_pending_pdf_conversions'):
                    self._pending_pdf_conversions = []
                self._pending_pdf_conversions.append((annex_i_path, str(pdf_dir)))
                self._pending_pdf_conversions.append((annex_iiib_path, str(pdf_dir)))
                self.logger.info(f"📄 Queued 2 documents for batch PDF conversion")
            
            self.stats.output_files_created += len(output_files)
            
            return ProcessingResult(
                success=True,
                message=f"Successfully processed {country} variant with updates: {', '.join(updates_applied)}",
                output_files=output_files
            )
            
        except Exception as e:
            raise DocumentError(f"Failed to save and split document: {e}")

    def _batch_convert_pdfs(self, pdf_dir: Path) -> List[str]:
        """Convert all pending Word documents to PDF after main processing."""
        if not hasattr(self, '_pending_pdf_conversions') or not self._pending_pdf_conversions:
            return []

        self.logger.info("=" * 80)
        self.logger.info(f"📄 Starting batch PDF conversion for {len(self._pending_pdf_conversions)} documents...")
        self.logger.info("=" * 80)

        pdf_files = []
        successful = 0
        failed = 0

        for idx, (doc_path, output_dir) in enumerate(self._pending_pdf_conversions, 1):
            self.logger.info(f"🔄 Converting {idx}/{len(self._pending_pdf_conversions)}: {Path(doc_path).name}")
            try:
                pdf_path = convert_to_pdf(doc_path, output_dir)
                pdf_files.append(pdf_path)
                successful += 1
                self.logger.info(f"✅ Success: {Path(pdf_path).name}")
            except Exception as e:
                failed += 1
                self.logger.warning(f"❌ Failed: {Path(doc_path).name} - {e}")

        self.logger.info("=" * 80)
        self.logger.info(f"📄 Batch PDF conversion complete: {successful} successful, {failed} failed")
        self.logger.info("=" * 80)

        return pdf_files

    def _generate_final_result(self, output_files: List[str]) -> ProcessingResult:
        """Generate final processing result with statistics."""
        
        self.logger.info("=" * 80)
        self.logger.info("✅ ENHANCED PROCESSING COMPLETE")
        self.logger.info("=" * 80)
        
        self.logger.info("📊 Final Statistics:")
        self.logger.info(f"   Input files found: {self.stats.input_files_found}")
        self.logger.info(f"   Input files processed: {self.stats.input_files_processed}")
        self.logger.info(f"   Total variants processed: {self.stats.variants_processed}")
        self.logger.info(f"   Successful variants: {self.stats.variants_successful}")
        self.logger.info(f"   Success rate: {self.stats.success_rate():.1f}%")
        self.logger.info(f"   Output files created: {self.stats.output_files_created}")
        self.logger.info(f"   Errors encountered: {self.stats.errors_encountered}")
        
        success = self.stats.variants_successful > 0
        message = f"Processed {self.stats.variants_successful}/{self.stats.variants_processed} variants successfully"
        
        # Get pending PDF conversions if they exist
        pending_conversions = getattr(self, '_pending_pdf_conversions', [])

        return ProcessingResult(
            success=success,
            message=message,
            output_files=output_files,
            pending_pdf_conversions=pending_conversions
        )

# ============================================================================= 
# BACKWARDS COMPATIBILITY INTERFACE
# =============================================================================

def process_folder(folder: str, mapping_path: str) -> None:
    """
    Backwards compatible entry point for processing folders.
    
    This function maintains the same interface as the original implementation
    while using the enhanced processing system under the hood.
    """
    try:
        processor = DocumentProcessor()
        result = processor.process_folder(folder, mapping_path)
        
        if not result.success:
            # Log the error but don't raise exception to maintain backwards compatibility
            logging.error(f"Processing failed: {result.message}")
            if result.errors:
                for error in result.errors:
                    logging.error(f"Error detail: {error}")
    
    except Exception as e:
        # Maintain backwards compatibility by logging errors instead of raising
        logging.error(f"Fatal processing error: {e}")
        raise  # Re-raise to maintain original behavior

def process_folder_enhanced(
    folder: str, 
    mapping_path: str, 
    config: Optional[ProcessingConfig] = None
) -> ProcessingResult:
    """
    Enhanced entry point that returns detailed results.
    
    Args:
        folder: Path to folder containing Word documents
        mapping_path: Path to Excel mapping file
        config: Optional processing configuration
        
    Returns:
        ProcessingResult with detailed success/failure information
    """
    processor = DocumentProcessor(config)
    return processor.process_folder(folder, mapping_path)
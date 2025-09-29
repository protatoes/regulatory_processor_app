"""
Local Representative Table Processor

Handles table-based local representative filtering in Annex IIIB documents.
This module provides clean, efficient table processing using direct table access
and cell merging for professional formatting.

Key Features:
- Direct table location via doc.tables[-1]
- Country-based validation
- Cell clearing and merging for clean formatting
- Integration with existing workflow
"""

from docx import Document
from docx.table import Table, _Row
import pandas as pd
from typing import Optional


class LocalRepTableProcessor:
    """Handles table-based local representative filtering in Annex IIIB."""

    def process_local_rep_table(self, doc: Document, mapping_row: pd.Series) -> bool:
        """
        Main entry point for table-based local rep processing.

        Args:
            doc: Document object to modify
            mapping_row: Row from mapping file containing processing data

        Returns:
            bool: True if table was successfully processed, False otherwise
        """
        target_country = mapping_row.get('Country', '').strip()

        # Always convert to country list (handles both single and combined countries)
        country_list = [c.strip() for c in target_country.split('/')]

        print(f"🔧 DEBUG: LocalRepTableProcessor.process_local_rep_table called")
        print(f"   Target countries: {country_list}")
        print(f"   Document has {len(doc.tables)} tables")

        if not target_country:
            print("❌ DEBUG: No target country provided")
            return False

        # Step 1: Locate the local rep table
        print("🔧 DEBUG: Step 1 - Locating local rep table...")
        local_rep_table = self._locate_local_rep_table(doc, country_list)
        if not local_rep_table:
            print("❌ DEBUG: No local rep table found")
            return False

        print(f"✅ DEBUG: Found local rep table with {len(local_rep_table.rows)} rows, {len(local_rep_table.columns)} columns")

        # Step 2: Filter table content for target countries
        print("🔧 DEBUG: Step 2 - Filtering table content...")
        result = self._filter_table_content(local_rep_table, country_list)
        print(f"🔧 DEBUG: Table filtering result: {result}")
        return result

    def _locate_local_rep_table(self, doc: Document, country_list: list) -> Optional[Table]:
        """
        Locate local rep table - it's the last table in the document.

        Args:
            doc: Document to search
            country_list: List of countries to validate table content against

        Returns:
            Optional[Table]: Located table if found and validated, None otherwise
        """
        print(f"🔧 DEBUG: _locate_local_rep_table called for countries: {country_list}")

        if not doc.tables:
            print("❌ DEBUG: No tables found in document")
            return None

        print(f"🔧 DEBUG: Document has {len(doc.tables)} tables")

        # Direct access to last table (O(1) operation)
        last_table = doc.tables[-1]
        print(f"🔧 DEBUG: Last table has {len(last_table.rows)} rows, {len(last_table.columns)} columns")

        # Print first few cells of table for debugging
        print("🔧 DEBUG: Table content preview:")
        for i, row in enumerate(last_table.rows[:3]):  # Show first 3 rows
            row_content = []
            for j, cell in enumerate(row.cells[:3]):  # Show first 3 cells
                cell_text = cell.text.strip()[:50]  # First 50 chars
                row_content.append(f"[{j}]: '{cell_text}'")
            if i < 3:  # Only show first 3 rows
                print(f"   Row {i}: {' | '.join(row_content)}")

        # Validate by checking if any target country is present
        print(f"🔧 DEBUG: Checking if table contains any of countries: {country_list}")
        contains_country = self._table_contains_country(last_table, country_list)
        print(f"🔧 DEBUG: Table contains target countries: {contains_country}")

        if contains_country:
            print("✅ DEBUG: Local rep table located and validated")
            return last_table
        else:
            print("❌ DEBUG: Table found but doesn't contain any target countries")

        return None

    def _table_contains_country(self, table: Table, country_list: list) -> bool:
        """
        Confirm table contains any target country in any cell.

        Args:
            table: Table to validate
            country_list: List of countries to search for (e.g., ['Ireland', 'Malta'])

        Returns:
            bool: True if any of the countries found in table, False otherwise
        """
        print(f"🔧 DEBUG: Searching for countries: {country_list}")

        countries_found = []
        found_countries = []

        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                if cell_text:  # Only check non-empty cells
                    countries_found.append(f"Row {i}, Col {j}: '{cell_text[:30]}'")

                    # Check if cell starts with any of our target countries
                    for country in country_list:
                        if self._cell_starts_with_country(cell_text, country.lower()):
                            print(f"✅ DEBUG: Found '{country}' at start of cell [{i},{j}]: '{cell_text[:50]}'")
                            found_countries.append(country)

        print(f"🔧 DEBUG: Countries/text found in table:")
        for country in countries_found[:10]:  # Show first 10 entries
            print(f"   {country}")

        if found_countries:
            print(f"✅ DEBUG: Found target countries: {found_countries}")
            return True
        else:
            print(f"❌ DEBUG: None of the target countries {country_list} found in table")
            return False

    def _filter_table_content(self, table: Table, country_list: list) -> bool:
        """
        Filter table using cell clearing and merging for clean formatting.

        This method processes each row independently:
        1. Identifies cells containing any target country
        2. Clears cells that don't contain target countries
        3. Merges kept cells across cleared neighbors for professional formatting

        Args:
            table: Table to filter
            country_list: List of countries to keep (e.g., ['Ireland', 'Malta'])

        Returns:
            bool: True if any rows were processed successfully
        """
        # Convert to lowercase for matching
        country_list_lower = [c.lower() for c in country_list]
        print(f"🔧 DEBUG: Filtering table for countries: {country_list_lower}")

        rows_processed = 0

        # Process each row independently
        for i, row in enumerate(table.rows):
            if self._process_table_row(row, country_list_lower, i):
                rows_processed += 1

        print(f"🔧 DEBUG: Table filtering completed. Rows processed: {rows_processed}")

        # Step 3: Clean up empty rows (keep first empty row for formatting)
        print("🔧 DEBUG: Step 3 - Cleaning up empty rows...")
        rows_removed = self._cleanup_empty_rows(table)
        print(f"🔧 DEBUG: Empty row cleanup completed. Rows removed: {rows_removed}")

        return rows_processed > 0

    def _process_table_row(self, row: _Row, country_list: list, row_index: int) -> bool:
        """
        Clear and merge cells in a single row to keep only target country data.

        Process:
        1. Identify cells to keep (contain any target country)
        2. Clear cells that don't contain target countries
        3. Merge kept cells over cleared cells for clean formatting

        Args:
            row: Table row to process
            country_list: List of lowercase target countries to keep
            row_index: Row index for debugging

        Returns:
            bool: True if row had cells to keep, False if row was fully cleared
        """
        cells_to_keep_indices = []

        print(f"🔧 DEBUG: Processing row {row_index}")

        # 1. Identify cells to keep and clear the rest
        for i, cell in enumerate(row.cells):
            cell_text = cell.text.strip()
            should_keep = False

            # Check if cell starts with any of the target countries (handles "Ireland: Regeneron Ireland DAC" correctly)
            for country in country_list:
                if self._cell_starts_with_country(cell_text, country):
                    should_keep = True
                    print(f"✅ DEBUG: Keeping cell [{row_index},{i}] - starts with '{country}': '{cell_text[:50]}'")
                    break

            if should_keep:
                cells_to_keep_indices.append(i)
            else:
                if cell_text:  # Only log if clearing non-empty cell
                    print(f"🗑️  DEBUG: Clearing cell [{row_index},{i}]: '{cell_text[:50]}'")
                cell.text = ''  # Clear content of unwanted cells

        # 2. If no cells to keep, row is fully cleared
        if not cells_to_keep_indices:
            print(f"❌ DEBUG: Row {row_index} - no cells to keep, fully cleared")
            return False

        print(f"✅ DEBUG: Row {row_index} - keeping {len(cells_to_keep_indices)} cells")

        # 3. Merge kept cells over cleared ones for clean formatting
        first_kept_cell_index = cells_to_keep_indices[0]
        merge_target_cell = row.cells[first_kept_cell_index]

        for i, cell in enumerate(row.cells):
            if i != first_kept_cell_index and cell.text == '':
                try:
                    merge_target_cell.merge(cell)
                    print(f"🔗 DEBUG: Merged cell [{row_index},{i}] into [{row_index},{first_kept_cell_index}]")
                except Exception as e:
                    # Log merge failure but continue processing
                    print(f"⚠️  Warning: Could not merge cell [{row_index},{i}] - {e}")

        return True

    def _cell_starts_with_country(self, cell_text: str, target_country: str) -> bool:
        """
        Check if cell text starts with the target country name.

        Handles patterns like:
        - "Ireland: Regeneron Ireland DAC" ✅ (starts with Ireland)
        - "France: Regeneron Ireland DAC" ❌ (starts with France, not Ireland)
        - "Ireland" ✅ (exact match)
        - "Regeneron Ireland DAC" ❌ (doesn't start with Ireland)

        Args:
            cell_text: Text content of the cell
            target_country: Country name to check for (lowercase)

        Returns:
            bool: True if cell starts with target country
        """
        if not cell_text:
            return False

        cell_lower = cell_text.lower().strip()

        # Check if starts with "Country:" pattern (most common)
        if cell_lower.startswith(f"{target_country}:"):
            return True

        # Check if starts with "Country " pattern (with space)
        if cell_lower.startswith(f"{target_country} "):
            return True

        # Check if exact match (just the country name)
        if cell_lower == target_country:
            return True

        # Extract first word and compare
        first_word = cell_lower.split()[0] if cell_lower.split() else ""
        # Remove common punctuation from first word
        first_word_clean = first_word.rstrip(":.,-")

        return first_word_clean == target_country

    def _cleanup_empty_rows(self, table: Table) -> int:
        """
        Remove empty rows after filtering to clean up table formatting.

        Strategy:
        1. Iterate through table rows in reverse order (to maintain indices)
        2. Check if all cells in row are empty (no text content)
        3. Remove empty rows, but keep the first empty row for spacing
        4. Log cleanup operations for debugging

        Args:
            table: Table to clean up

        Returns:
            int: Number of rows removed
        """
        rows_to_remove = []
        first_empty_row_found = False

        print(f"🔧 DEBUG: Starting empty row cleanup on table with {len(table.rows)} rows")

        # Identify empty rows (iterate forward to find first empty row)
        for i, row in enumerate(table.rows):
            all_cells_empty = True

            for cell in row.cells:
                if cell.text.strip():  # If any cell has content
                    all_cells_empty = False
                    break

            if all_cells_empty:
                if first_empty_row_found:
                    # This is not the first empty row, mark for removal
                    rows_to_remove.append(i)
                    print(f"🗑️  DEBUG: Marking empty row {i} for removal")
                else:
                    # This is the first empty row, keep it for formatting
                    first_empty_row_found = True
                    print(f"✅ DEBUG: Keeping first empty row {i} for formatting")

        # Remove rows in reverse order to maintain indices
        for row_index in reversed(rows_to_remove):
            try:
                row_element = table.rows[row_index]._element
                row_element.getparent().remove(row_element)
                print(f"🔗 DEBUG: Removed empty row {row_index}")
            except Exception as e:
                print(f"⚠️  Warning: Could not remove row {row_index} - {e}")

        print(f"🔧 DEBUG: Empty row cleanup completed. {len(rows_to_remove)} rows removed")
        return len(rows_to_remove)


def process_local_rep_table_standalone(doc: Document, target_country: str) -> bool:
    """
    Standalone function for processing local rep table without mapping file.

    Args:
        doc: Document to process
        target_country: Country to filter for

    Returns:
        bool: True if processing succeeded
    """
    processor = LocalRepTableProcessor()

    # Create minimal mapping row for compatibility
    mapping_row = pd.Series({'Country': target_country})

    return processor.process_local_rep_table(doc, mapping_row)
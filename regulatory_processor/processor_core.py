"""Main document processor classes and orchestration logic."""

import logging
import os
import shutil
from pathlib import Path
from typing import List, Optional, Tuple
from docx import Document
from docx.document import Document as DocumentObject

try:
    from .config import ProcessingConfig, ProcessingResult, ProcessingStats, DirectoryNames, FileMarkers, ValidationError, ProcessingError, DocumentError, MappingError
    from .file_manager import (
        load_mapping_table,
        identify_document_country_and_language,
        generate_output_filename,
        convert_to_pdf,
    )
    from .mapping_table import MappingRow, MappingTable
    from .document_utils import update_section_10_date, update_annex_iiib_date, update_local_representatives
    from .document_splitter import split_annexes
except ImportError:
    # Fallback for when module is imported directly
    from config import ProcessingConfig, ProcessingResult, ProcessingStats, DirectoryNames, FileMarkers, ValidationError, ProcessingError, DocumentError, MappingError
    from file_manager import (
        load_mapping_table,
        identify_document_country_and_language,
        generate_output_filename,
        convert_to_pdf,
    )
    from mapping_table import MappingRow, MappingTable
    from document_utils import update_section_10_date, update_annex_iiib_date, update_local_representatives
    from document_splitter import split_annexes


class FileManager:
    """Handles file operations and path management."""
    
    def __init__(self, base_folder: Path, config: ProcessingConfig):
        self.base_folder = base_folder
        self.config = config
        self.logger = logging.getLogger(f"{__name__}.FileManager")
    
    def setup_output_directories(self) -> Tuple[Path, Path]:
        """Create and return paths for output directories."""
        split_dir = self.base_folder / DirectoryNames.SPLIT_DOCS
        pdf_dir = self.base_folder / DirectoryNames.PDF_DOCS
        
        try:
            os.makedirs(split_dir, exist_ok=True)
            os.makedirs(pdf_dir, exist_ok=True)
            return split_dir, pdf_dir
        except OSError as e:
            raise ProcessingError(f"Failed to create output directories: {e}")
    
    def discover_processable_documents(self) -> List[Path]:
        """Find all valid Word documents that can be processed."""
        if not self.base_folder.is_dir():
            raise ValidationError(f"Folder does not exist: {self.base_folder}")
        
        documents = []
        for file_path in self.base_folder.iterdir():
            if self._is_processable_document(file_path):
                documents.append(file_path)
        
        return documents
    
    def _is_processable_document(self, file_path: Path) -> bool:
        """Check if a file is a valid document for processing."""
        if file_path.suffix.lower() != ".docx":
            return False
        if file_path.name.startswith(FileMarkers.TEMP_FILE_PREFIX):
            return False
        if FileMarkers.ANNEX_MARKER in file_path.name:
            return False
        if file_path.name.startswith(FileMarkers.ANNEX_PREFIX):
            return False
        return True
    
    def create_backup(self, file_path: Path) -> Optional[Path]:
        """Create a backup of the original file."""
        if not self.config.create_backups:
            return None
            
        backup_path = file_path.with_suffix(file_path.suffix + DirectoryNames.BACKUP_SUFFIX)
        if backup_path.exists():
            return backup_path
            
        try:
            shutil.copy2(file_path, backup_path)
            return backup_path
        except Exception as e:
            self.logger.warning(f"Failed to create backup: {e}")
            return None


class DocumentUpdater:
    """Handles document modification operations."""
    
    def __init__(self, config: ProcessingConfig):
        self.config = config
        self.logger = logging.getLogger(f"{__name__}.DocumentUpdater")
    
    def apply_all_updates(self, doc: DocumentObject, mapping_row: MappingRow) -> Tuple[bool, List[str]]:
        """Apply all required updates to a document."""
        updates_applied = []
        total_success = False
        
        try:
            # 1. Update national reporting systems - SmPC Section 4.8
            smpc_success, smpc_updates = self._update_document_with_fixed_smpc_blocks(doc, mapping_row)
            if smpc_success:
                updates_applied.extend(smpc_updates)
                total_success = True

            # 2. Update national reporting systems - PL Section 4
            pl_success, pl_updates = self._update_document_with_fixed_pl_blocks(doc, mapping_row)
            if pl_success:
                updates_applied.extend(pl_updates)
                total_success = True

            # 3. Update dates
            annex_i_date_success = update_section_10_date(doc, mapping_row)
            if annex_i_date_success:
                updates_applied.append("Annex I dates")
                total_success = True

            annex_iiib_date_success = update_annex_iiib_date(doc, mapping_row)
            if annex_iiib_date_success:
                updates_applied.append("Annex IIIB dates")
                total_success = True

            # 4. Update local representatives
            local_rep_success = update_local_representatives(doc, mapping_row)
            if local_rep_success:
                updates_applied.append("Local representatives")
                total_success = True
                
            return total_success, updates_applied
            
        except Exception as e:
            raise DocumentError(f"Failed to apply document updates: {e}")
    
    def _update_document_with_fixed_smpc_blocks(self, doc: Document, mapping_row: MappingRow) -> Tuple[bool, List[str]]:
        """
        Update SmPC Section 4.8 with multi-country formatted text blocks.

        Builds complete text blocks for each country using Lines 1-10, applies bold
        formatting to country names, creates hyperlinks, and replaces original text.
        """
        updates = []
        success = False

        try:
            # Get original text to replace
            original_text = mapping_row.get('Original text national reporting - SmPC', '')
            if not original_text or str(original_text).lower() in ['nan', '']:
                return False, updates

            # Build country blocks from mapping data
            country_blocks = self._build_smpc_country_blocks(mapping_row)
            if not country_blocks:
                return False, updates

            # Find SmPC Section 4.8 paragraphs
            smpc_paragraphs = self._find_smpc_section_48_paragraphs(doc)
            if not smpc_paragraphs:
                return False, updates

            # Replace original text with country blocks in each relevant paragraph
            blocks_inserted = False
            for para in smpc_paragraphs:
                try:
                    from .document_utils import replace_original_text_with_country_blocks
                except ImportError:
                    from document_utils import replace_original_text_with_country_blocks
                if replace_original_text_with_country_blocks(para, original_text, country_blocks):
                    blocks_inserted = True
                    break  # Usually only one paragraph contains the target text

            if blocks_inserted:
                success = True
                updates.append("SmPC Section 4.8")

        except Exception as e:
            print(f"Error updating SmPC blocks: {e}")

        return success, updates

    def _build_smpc_country_blocks(self, mapping_row: MappingRow) -> List[Dict]:
        """
        Build formatted country blocks from mapping data.

        Returns list of country blocks with lines, bold_texts, and hyperlinks.
        """
        country_blocks = []

        try:
            # Parse countries from Line 1 - Country names to be bolded - SmPC
            country_bold_text = mapping_row.get('Line 1 - Country names to be bolded - SmPC', '')
            if not country_bold_text or str(country_bold_text).lower() in ['nan', '']:
                return country_blocks

            # Split countries by comma/semicolon
            import re
            countries = [c.strip() for c in re.split('[,;]', country_bold_text) if c.strip()]

            # Get all line texts (Lines 1-10)
            line_texts = []
            for i in range(1, 11):
                line_key = f'Line {i} - SmPC'
                line_text = mapping_row.get(line_key, '')
                if line_text and str(line_text).lower() not in ['nan', '']:
                    line_texts.append(str(line_text))

            if not line_texts:
                return country_blocks

            # Get hyperlink data
            hyperlinks_text = mapping_row.get('Hyperlinks SmPC', '')
            hyperlinks_urls = mapping_row.get('Link for email - SmPC', '')

            # Parse hyperlinks
            hyperlink_data = self._parse_hyperlink_data(hyperlinks_text, hyperlinks_urls)

            # Create one block per country (or single block if multiple countries share same text)
            # For now, create single block with all countries - can be refined later
            country_block = {
                'lines': line_texts,
                'bold_texts': countries,
                'hyperlinks': hyperlink_data
            }
            country_blocks.append(country_block)

        except Exception as e:
            print(f"Error building SmPC country blocks: {e}")

        return country_blocks

    def _parse_hyperlink_data(self, hyperlinks_text: str, hyperlinks_urls: str) -> List[Dict]:
        """
        Parse hyperlink text and URLs from mapping columns.

        Args:
            hyperlinks_text: e.g., "national reporting system, Appendix V"
            hyperlinks_urls: e.g., "https://url1.com, adr@example.com"

        Returns:
            List of {'text': 'url.com', 'url': 'url.com'} dicts
        """
        hyperlink_data = []

        try:
            if not hyperlinks_urls or str(hyperlinks_urls).lower() in ['nan', '', '*n/a*']:
                return hyperlink_data

            # Split URLs by comma
            import re
            urls = [u.strip() for u in re.split('[,;]', str(hyperlinks_urls)) if u.strip()]

            # For each URL, the hyperlink text is the URL itself (as per your clarification)
            for url in urls:
                if url and url.lower() not in ['*n/a*', 'nan']:
                    hyperlink_data.append({
                        'text': url,  # The visible text is the URL/email itself
                        'url': url    # The target is the same URL/email
                    })

        except Exception as e:
            print(f"Error parsing hyperlink data: {e}")

        return hyperlink_data

    def _find_smpc_section_48_paragraphs(self, doc: Document) -> List:
        """
        Find paragraphs in SmPC Section 4.8 that contain reporting information.

        Looks for sections containing "4.8" and "adverse" or "undesirable" effects.
        """
        section_paragraphs = []

        try:
            for para in doc.paragraphs:
                para_text = para.text.lower()

                # Look for Section 4.8 indicators
                if ("4.8" in para_text and
                    ("adverse" in para_text or "undesirable" in para_text or
                     "reporting" in para_text or "suspected" in para_text)):
                    section_paragraphs.append(para)

        except Exception as e:
            print(f"Error finding SmPC Section 4.8 paragraphs: {e}")

        return section_paragraphs

    def _update_document_with_fixed_pl_blocks(self, doc: Document, mapping_row: MappingRow) -> Tuple[bool, List[str]]:
        """
        Update PL Section 4 with multi-country formatted text blocks.

        Reuses SmPC text block construction logic and appends PL-specific text.
        """
        updates = []
        success = False

        try:
            # SmPC/PL Consistency Validation
            smpc_text = mapping_row.get('National reporting system SmPC', '')
            pl_text = mapping_row.get('National reporting system PL', '')

            if smpc_text != pl_text and both_exist(smpc_text, pl_text):
                print(f"⚠️  Warning: SmPC text differs from PL text - using SmPC base")
                print(f"    SmPC: {smpc_text[:50]}...")
                print(f"    PL: {pl_text[:50]}...")

            # Get original text to replace
            original_text = mapping_row.get('Original text national reporting - PL', '')
            if not original_text or str(original_text).lower() in ['nan', '']:
                return False, updates

            # Build PL country blocks (reuse SmPC logic + append PL text)
            country_blocks = self._build_pl_country_blocks(mapping_row)
            if not country_blocks:
                return False, updates

            # Find PL Section 4 paragraphs
            pl_paragraphs = self._find_pl_section_4_paragraphs(doc)
            if not pl_paragraphs:
                return False, updates

            # Replace original text with country blocks in each relevant paragraph
            blocks_inserted = False
            for para in pl_paragraphs:
                try:
                    from .document_utils import replace_original_text_with_country_blocks
                except ImportError:
                    from document_utils import replace_original_text_with_country_blocks
                if replace_original_text_with_country_blocks(para, original_text, country_blocks):
                    blocks_inserted = True
                    break  # Usually only one paragraph contains the target text

            if blocks_inserted:
                success = True
                updates.append("PL Section 4")

        except Exception as e:
            print(f"Error updating PL blocks: {e}")

        return success, updates

    def _build_pl_country_blocks(self, mapping_row: MappingRow) -> List[Dict]:
        """
        Build PL country blocks by reusing SmPC logic and appending PL-specific text.
        """
        try:
            # Start with SmPC country blocks
            country_blocks = self._build_smpc_country_blocks(mapping_row)

            # Get PL-specific text to append
            pl_append_text = mapping_row.get('Text to be appended after National reporting system PL', '')

            # Append PL text to each country block
            if pl_append_text and str(pl_append_text).lower() not in ['nan', '']:
                for block in country_blocks:
                    # Add the append text as a new line
                    block['lines'].append(str(pl_append_text))

            return country_blocks

        except Exception as e:
            print(f"Error building PL country blocks: {e}")
            return []

    def _find_pl_section_4_paragraphs(self, doc: Document) -> List:
        """
        Find paragraphs in PL Section 4 that contain side effects information.

        Looks for sections containing "4" or "section 4" and "side effects" or "adverse".
        """
        section_paragraphs = []

        try:
            for para in doc.paragraphs:
                para_text = para.text.lower()

                # Look for PL Section 4 indicators
                if (("4." in para_text or "section 4" in para_text) and
                    ("side effects" in para_text or "adverse" in para_text or
                     "reporting" in para_text or "suspected" in para_text)):
                    section_paragraphs.append(para)

        except Exception as e:
            print(f"Error finding PL Section 4 paragraphs: {e}")

        return section_paragraphs


def both_exist(text1: str, text2: str) -> bool:
    """Helper function to check if both texts exist and are not nan/empty."""
    return (text1 and str(text1).lower() not in ['nan', ''] and
            text2 and str(text2).lower() not in ['nan', ''])


class DocumentProcessor:
    """Main document processing orchestrator."""
    
    def __init__(self, config: Optional[ProcessingConfig] = None):
        self.config = config or ProcessingConfig()
        self.stats = ProcessingStats()
        self.logger = self._setup_logging()
        
    def _setup_logging(self) -> logging.Logger:
        """Set up logging configuration."""
        logger = logging.getLogger(__name__)
        logger.setLevel(getattr(logging, self.config.log_level.upper()))
        
        if not logger.handlers:
            handler = logging.StreamHandler()
            formatter = logging.Formatter(
                '%(asctime)s - %(name)s - %(levelname)s - %(message)s'
            )
            handler.setFormatter(formatter)
            logger.addHandler(handler)
        
        return logger
    
    def process_folder(
        self, folder_path: str, mapping_path: Optional[str] = None
    ) -> ProcessingResult:
        """Main entry point for processing a folder of documents."""
        try:
            self.logger.info("=" * 80)
            self.logger.info("🚀 STARTING ENHANCED DOCUMENT PROCESSING")
            self.logger.info("=" * 80)
            
            # Validate inputs
            folder = self._validate_folder_path(folder_path)
            mapping_table = self._load_and_validate_mapping(mapping_path)
            
            # Setup processing environment
            file_manager = FileManager(folder, self.config)
            split_dir, pdf_dir = file_manager.setup_output_directories()
            
            # Discover documents to process
            documents = file_manager.discover_processable_documents()
            self.stats.input_files_found = len(documents)
            
            if not documents:
                return ProcessingResult(
                    success=False,
                    message="No valid documents found for processing"
                )
            
            # Process each document
            output_files = []
            for document_path in documents:
                try:
                    result = self._process_single_document(
                        document_path, mapping_table, file_manager, split_dir, pdf_dir
                    )
                    output_files.extend(result.output_files)
                    
                except Exception as e:
                    self.logger.error(f"Error processing {document_path.name}: {e}")
                    self.stats.errors_encountered += 1
            
            # Generate final report
            return self._generate_final_result(output_files)
            
        except Exception as e:
            self.logger.error(f"Fatal error in process_folder: {e}")
            return ProcessingResult(
                success=False,
                message=f"Processing failed: {e}",
                errors=[str(e)]
            )
    
    def _validate_folder_path(self, folder_path: str) -> Path:
        """Validate and return folder path."""
        folder = Path(folder_path).resolve()
        if not folder.is_dir():
            raise ValidationError(f"Invalid directory: {folder_path}")
        return folder
    
    def _load_and_validate_mapping(self, mapping_path: Optional[str]) -> MappingTable:
        """Load and validate mapping file."""
        try:
            mapping_table = load_mapping_table(mapping_path, self.config)
            if mapping_table is None or len(mapping_table) == 0:
                if mapping_path:
                    raise MappingError(f"Could not load mapping file: {mapping_path}")
                raise MappingError("Default mapping table could not be loaded")

            source = mapping_path if mapping_path else "built-in default"
            self.logger.info(
                f"Mapping loaded from {source}: {len(mapping_table)} configurations"
            )
            return mapping_table

        except Exception as e:
            raise MappingError(f"Failed to load mapping file: {e}")
    
    def _process_single_document(
        self,
        document_path: Path,
        mapping_table: MappingTable,
        file_manager: FileManager,
        split_dir: Path,
        pdf_dir: Path
    ) -> ProcessingResult:
        """Process a single document with all its variants."""
        
        self.logger.info("=" * 60)
        self.logger.info(f"📄 PROCESSING: {document_path.name}")
        self.logger.info("=" * 60)
        
        self.stats.input_files_processed += 1
        
        try:
            # Identify document language and country
            country_code, language_name, country_name = identify_document_country_and_language(str(document_path))
            
            if not language_name:
                error_msg = f"Cannot identify language for {document_path.name}"
                self.logger.error(error_msg)
                return ProcessingResult(success=False, message=error_msg)
            
            self.logger.info(f"Document identified - Language: {language_name}, Country: {country_name}")
            
            # Find mapping rows for this language
            mapping_rows = mapping_table.for_language(language_name)
            if not mapping_rows:
                error_msg = f"No mapping found for language: {language_name}"
                self.logger.error(error_msg)
                return ProcessingResult(success=False, message=error_msg)
            
            self.logger.info(f"Found {len(mapping_rows)} variant(s) to process")
            
            # Create backup
            backup_path = file_manager.create_backup(document_path)
            if backup_path:
                self.stats.record_backup()
            
            # Process each variant
            output_files = []
            variant_success_count = 0
            
            for i, mapping_row in enumerate(mapping_rows, 1):
                country = mapping_row.country_display or str(mapping_row.get('Country', '')).strip()
                self.logger.info(f"🌍 Processing variant {i}/{len(mapping_rows)}: {country}")
                
                try:
                    result = self._process_document_variant(
                        document_path, mapping_row, split_dir, pdf_dir
                    )
                    
                    if result.success:
                        variant_success_count += 1
                        output_files.extend(result.output_files)
                        self.logger.info(f"✅ Variant {i} completed successfully")
                    else:
                        self.logger.warning(f"⚠️ Variant {i} completed with issues: {result.message}")

                    self.stats.record_variant(result.success)

                except Exception as e:
                    self.logger.error(f"❌ Error processing variant {i} ({country}): {e}")
                    self.stats.record_variant(False)
            
            # Document summary
            success_rate = (variant_success_count / len(mapping_rows)) * 100 if mapping_rows else 0
            self.logger.info(f"📊 Document Summary: {variant_success_count}/{len(mapping_rows)} variants successful ({success_rate:.1f}%)")
            
            return ProcessingResult(
                success=variant_success_count > 0,
                message=f"Processed {variant_success_count}/{len(mapping_rows)} variants successfully",
                output_files=output_files
            )
            
        except Exception as e:
            self.logger.error(f"Error processing document {document_path.name}: {e}")
            return ProcessingResult(success=False, message=str(e), errors=[str(e)])
    
    def _process_document_variant(
        self,
        document_path: Path,
        mapping_row: MappingRow,
        split_dir: Path,
        pdf_dir: Path
    ) -> ProcessingResult:
        """Process a single document variant."""
        
        country = mapping_row.country_display or str(mapping_row.get('Country', '')).strip()
        language = mapping_row.language or str(mapping_row.get('Language', '')).strip()
        
        try:
            # Load document
            doc = Document(str(document_path))
            
            # Apply updates
            updater = DocumentUpdater(self.config)
            updates_made, updates_applied = updater.apply_all_updates(doc, mapping_row)
            
            if not updates_made:
                return ProcessingResult(
                    success=False,
                    message=f"No updates applied for {country} variant"
                )
            
            # Save and process updated document
            return self._save_and_split_document(
                doc, document_path, mapping_row, split_dir, pdf_dir, updates_applied
            )
            
        except Exception as e:
            raise DocumentError(f"Failed to process variant for {country}: {e}")
    
    def _save_and_split_document(
        self,
        doc: Document,
        original_path: Path,
        mapping_row: MappingRow,
        split_dir: Path,
        pdf_dir: Path,
        updates_applied: List[str]
    ) -> ProcessingResult:
        """Save updated document and split into annexes."""

        country = mapping_row.country_display or str(mapping_row.get('Country', '')).strip()
        language = mapping_row.language or str(mapping_row.get('Language', '')).strip()
        output_files: List[str] = []

        try:
            # Generate output filename
            base_name = original_path.stem
            output_filename = generate_output_filename(base_name, language, country, "combined")
            output_path = original_path.parent / output_filename

            # Save updated document
            doc.save(str(output_path))
            output_files.append(str(output_path))
            self.logger.info(f"💾 Saved combined document: {output_filename}")

            # Split into annexes
            self.logger.info("🔀 Splitting into separate annexes...")
            annex_i_path, annex_iiib_path = split_annexes(
                str(output_path), str(split_dir), language, country, mapping_row
            )

            if annex_i_path:
                output_files.append(annex_i_path)
                self.stats.annex_i_created += 1
            if annex_iiib_path:
                output_files.append(annex_iiib_path)
                self.stats.annex_iiib_created += 1
            self.logger.info("✅ Split completed")

            # Convert to PDF if enabled
            if self.config.convert_to_pdf:
                try:
                    self.logger.info("📄 Converting to PDF...")

                    pdf_annex_i = convert_to_pdf(annex_i_path, str(pdf_dir))
                    if pdf_annex_i:
                        output_files.append(pdf_annex_i)
                        self.stats.record_pdf_result(True)
                    else:
                        self.stats.record_pdf_result(False)

                    pdf_annex_iiib = convert_to_pdf(annex_iiib_path, str(pdf_dir))
                    if pdf_annex_iiib:
                        output_files.append(pdf_annex_iiib)
                        self.stats.record_pdf_result(True)
                    else:
                        self.stats.record_pdf_result(False)

                    self.logger.info("✅ PDF conversion completed")

                except Exception as exc:
                    self.logger.warning(f"PDF conversion failed: {exc}")
                    self.stats.record_pdf_result(False)

            self.stats.record_outputs(len(output_files))

            return ProcessingResult(
                success=True,
                message=f"Successfully processed variant: {country}",
                output_files=output_files
            )

        except Exception as e:
            raise DocumentError(f"Failed to save and split document: {e}")


    def _generate_final_result(self, output_files: List[str]) -> ProcessingResult:
        """Generate final processing result."""
        success_rate = self.stats.success_rate()
        
        if self.stats.variants_successful == 0:
            return ProcessingResult(
                success=False,
                message="No documents were processed successfully",
                errors=["Processing failed for all variants"]
            )
        
        message = (
            f"Processing completed: {self.stats.variants_successful}/{self.stats.variants_processed} "
            f"variants successful ({success_rate:.1f}%)"
        )
        
        return ProcessingResult(
            success=True,
            message=message,
            output_files=output_files
        )


def process_folder(folder: str, mapping_path: Optional[str] = None) -> None:
    """Legacy wrapper function for backward compatibility."""
    try:
        processor = DocumentProcessor()
        result = processor.process_folder(folder, mapping_path)
        
        if not result.success:
            # Log the error but don't raise exception to maintain backwards compatibility
            logging.error(f"Processing failed: {result.message}")
            if result.errors:
                for error in result.errors:
                    logging.error(f"Error detail: {error}")
    
    except Exception as e:
        # Maintain backwards compatibility by logging errors instead of raising
        logging.error(f"Fatal processing error: {e}")
        raise  # Re-raise to maintain original behavior


def process_folder_enhanced(
    folder: str,
    mapping_path: Optional[str] = None,
    config: Optional[ProcessingConfig] = None
) -> ProcessingResult:
    """Enhanced entry point that returns detailed results."""
    processor = DocumentProcessor(config)
    return processor.process_folder(folder, mapping_path)
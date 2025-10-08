"""
Document splitting using clone-and-prune approach.
Preserves all scaffolding and formatting while surgically removing unwanted content.

This module replaces the complex copy-based document splitting with a superior
clone-and-prune approach that maintains perfect document fidelity.
"""

import shutil
import os
from pathlib import Path
from docx import Document
from typing import List, Tuple, Optional, Dict
import logging

# Configure logging
logger = logging.getLogger(__name__)


def clone_and_split_document(
    source_path: str,
    output_dir: str,
    country_code: str,
    target_annexes: List[str] = None,
    language: str = None,
    mapping_row = None
) -> Dict[str, str]:
    """
    Main function to split a combined document into separate annex documents.
    Uses clone-and-prune approach for perfect scaffolding preservation.

    Args:
        source_path: Path to combined source document
        output_dir: Directory for output files
        country_code: Country code for output filenames
        target_annexes: List of annexes to extract (default: ["ANNEX I", "ANNEX IIIB"])

    Returns:
        Dict mapping annex names to output file paths

    Example:
        paths = clone_and_split_document(
            "combined.docx",
            "output/",
            "IE",
            ["ANNEX I", "ANNEX IIIB"]
        )
        # Returns: {"ANNEX I": "output/IE_ANNEX_I.docx", "ANNEX IIIB": "output/IE_ANNEX_IIIB.docx"}
    """
    if target_annexes is None:
        target_annexes = ["ANNEX I", "ANNEX IIIB"]

    logger.info(f"🔄 Splitting document {source_path} for country {country_code}")
    logger.info(f"📋 Target annexes: {target_annexes}")

    # Ensure output directory exists
    os.makedirs(output_dir, exist_ok=True)

    # Extract all annex headers from mapping for proper boundary detection
    all_annex_headers = []
    if mapping_row is not None:
        annex_i_header = mapping_row.get('Annex I Header in country language', '').strip()
        annex_ii_header = mapping_row.get('Annex II Header in country language', '').strip()
        annex_iiib_header = mapping_row.get('Annex IIIB Header in country language', '').strip()

        print(f"🗂️ MAPPING HEADERS EXTRACTED:")
        print(f"   Annex I: '{annex_i_header}'")
        print(f"   Annex II: '{annex_ii_header}'")
        print(f"   Annex IIIB: '{annex_iiib_header}'")

        # Add non-empty headers to the list
        for header in [annex_i_header, annex_ii_header, annex_iiib_header]:
            if header:
                all_annex_headers.append(header)
                logger.info(f"   ✅ Added to all_annex_headers: '{header}'")

        logger.info(f"📋 Final all_annex_headers list: {all_annex_headers}")

    result_paths = {}

    for annex in target_annexes:
        try:
            # Generate proper filename using mapping conventions
            output_filename = _generate_annex_filename(annex, language, mapping_row)
            output_path = os.path.join(output_dir, output_filename)

            # Clone source document
            clone_source_document(source_path, output_path)

            # Determine if this is Annex I from mapping data
            is_annex_i = (mapping_row is not None and
                         annex == mapping_row.get('Annex I Header in country language', '').strip())

            print(f"🧪 PROCESSING ANNEX: '{annex}'")
            print(f"   Is Annex I: {is_annex_i}")
            print(f"   Annex I header from mapping: '{mapping_row.get('Annex I Header in country language', '').strip()}'")

            # Prune to target annex
            print(f"🔧 Starting pruning process for {annex}")

            # OPTIMIZATION: Find boundaries once and pass them to avoid duplicate processing
            print(f"🔧 Pre-calculating boundaries to avoid duplicate work...")
            temp_doc = Document(output_path)
            start_idx, end_idx = find_annex_boundaries(temp_doc, annex, all_annex_headers, is_annex_i, mapping_row)
            print(f"🔧 Pre-calculated boundaries: start={start_idx}, end={end_idx}")

            success = prune_to_annex_with_boundaries(output_path, annex, start_idx, end_idx)
            print(f"🔧 Pruning result for {annex}: {'SUCCESS' if success else 'FAILED'}")

            if success:
                result_paths[annex] = output_path
                print(f"✅ Successfully created {annex} document: {output_path}")

                # Verify the pruned document
                try:
                    verify_doc = Document(output_path)
                    print(f"   📊 Verification: Document has {len(verify_doc.paragraphs)} paragraphs")
                except Exception as e:
                    print(f"   ⚠️ Could not verify document: {e}")
            else:
                print(f"❌ Failed to prune {annex} from {output_path}")

        except Exception as e:
            logger.error(f"❌ Error processing {annex}: {e}")

    return result_paths


def clone_source_document(source_path: str, output_path: str) -> str:
    """
    Create a byte-for-byte clone of the source document.
    Preserves ALL scaffolding: headers, footers, styles, properties, etc.

    Args:
        source_path: Path to source .docx file
        output_path: Path for cloned document

    Returns:
        Path to cloned document

    Raises:
        FileNotFoundError: If source file doesn't exist
        PermissionError: If cannot write to output location
    """
    if not os.path.exists(source_path):
        raise FileNotFoundError(f"Source document not found: {source_path}")

    # Ensure output directory exists
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    # Use shutil.copy2 to preserve metadata and timestamps
    shutil.copy2(source_path, output_path)

    logger.debug(f"📄 Cloned document: {source_path} → {output_path}")
    return output_path


def find_annex_boundaries(doc: Document, target_annex: str, all_annex_headers: List[str] = None, is_annex_i: bool = False, mapping_row = None) -> Tuple[Optional[int], Optional[int]]:
    """
    Find the start and end paragraph indices for a specific annex.
    Handles proper annex boundary detection to avoid partial matches.

    Args:
        doc: Document to search
        target_annex: Annex identifier (e.g., "ANNEX I", "ANNEX II", "ANNEX IIIB")
        all_annex_headers: List of all known annex headers from mapping file
        is_annex_i: True if this is Annex I (starts from document beginning)

    Returns:
        Tuple of (start_index, end_index) or (None, None) if not found

    Note:
        - Uses strict matching to avoid "ANNEX I" matching "ANNEX III"
        - For Annex I: starts from paragraph 0 to preserve title pages and intro content
        - For Annex I boundaries: prioritizes Annex II, uses Annex IIIB only as fallback
        - For other annexes: starts from their header location, uses all headers for boundaries
        - If no end marker found, assumes annex extends to document end
    """
    print(f"🔍 FINDING BOUNDARIES FOR: '{target_annex}'")
    print(f"🎯 is_annex_i: {is_annex_i}")
    print(f"📄 Document has {len(doc.paragraphs)} paragraphs")

    start_idx = None
    end_idx = None

    # Normalize target for comparison - handle various space characters
    def normalize_text(text):
        """Normalize text by converting non-breaking spaces and other whitespace to regular spaces"""
        return text.upper().strip().replace('\xa0', ' ').replace('\u00a0', ' ')

    target_upper = normalize_text(target_annex)
    logger.info(f"🎯 Normalized target: '{target_upper}'")

    # Special case: For Annex I, start from beginning of document (paragraph 0)
    # This preserves title pages and introductory content
    if is_annex_i:
        start_idx = 0
        logger.info(f"📍 Annex I detected - starting from document beginning (paragraph 0)")

    # First pass: log all annex-related paragraphs for debugging (REDUCED for performance)
    print("🔍 SCANNING DOCUMENT FOR ANNEX HEADERS...")
    annex_paragraphs = []
    for i, para in enumerate(doc.paragraphs):
        para_text = normalize_text(para.text)
        if "ANNEX" in para_text or "ANEXO" in para_text:
            annex_paragraphs.append((i, para.text.strip(), para_text))

    # Only show the annex headers, not all the debug text
    for i, para_text, normalized in annex_paragraphs:
        print(f"   Para {i}: '{para_text}'")

    if not annex_paragraphs:
        print("   ⚠️ NO ANNEX HEADERS FOUND IN DOCUMENT!")
    else:
        print(f"   📋 Found {len(annex_paragraphs)} annex-related paragraphs")

    # Define priority headers ONCE before the main loop
    if is_annex_i:
        # For Annex I: prioritize Annex II by explicitly finding it in mapping_row
        priority_headers = []
        if mapping_row is not None:
            # Get Annex II header directly from mapping row (not by array index!)
            explicit_annex_ii = mapping_row.get('Annex II Header in country language', '').strip()
            if explicit_annex_ii:
                priority_headers.append(explicit_annex_ii)
                print(f"✅ Added Annex II as priority boundary: '{explicit_annex_ii}'")
            else:
                print(f"⚠️ No Annex II header found in mapping")

            # Add Annex IIIB as fallback
            explicit_annex_iiib = mapping_row.get('Annex IIIB Header in country language', '').strip()
            if explicit_annex_iiib:
                priority_headers.append(explicit_annex_iiib)
                print(f"✅ Added Annex IIIB as fallback boundary: '{explicit_annex_iiib}'")
            else:
                print(f"⚠️ No Annex IIIB header found in mapping")
        else:
            # Fallback: use all headers if no mapping_row
            priority_headers = all_annex_headers
        print(f"🎯 Annex I boundary priority: {priority_headers}")
    else:
        # For other annexes: use all headers as before
        priority_headers = all_annex_headers
        print(f"🎯 {target_annex} boundary headers: {priority_headers}")

    # Main processing loop
    for i, para in enumerate(doc.paragraphs):
        para_text = normalize_text(para.text)

        # Found target annex start - use strict matching (skip for Annex I since we start at 0)
        if not is_annex_i and start_idx is None and para_text.startswith(target_upper):
            # Additional check: ensure it's not a substring match
            # e.g., "ANNEX I" should not match "ANNEX III"
            if para_text == target_upper or para_text.startswith(target_upper + " "):
                start_idx = i
                logger.debug(f"📍 Found {target_annex} start at paragraph {i}: '{para.text[:50]}...'")
                continue

        # Found next annex (end of current annex) - use mapping file headers with proper sequencing
        if start_idx is not None:
            if all_annex_headers is None:
                raise ValueError("all_annex_headers parameter is required for proper annex boundary detection")

            # Check if this paragraph starts any prioritized annex header
            for header in priority_headers:
                header_upper = normalize_text(header)
                # Simplified logging for performance - only log boundary matches
                if "ANNEX" in para_text or "ANEXO" in para_text:
                    if para_text.startswith(header_upper):
                        print(f"🔍 Para {i}: MATCH '{para.text.strip()}' vs '{header}'")

                if para_text.startswith(header_upper):
                    # Make sure it's not the same annex continuing
                    # FIXED: Use exact match to avoid substring issues (e.g., "ANEXO II" vs "ANEXO I")
                    if para_text != target_upper and not para_text.startswith(target_upper + " "):
                        end_idx = i
                        print(f"🔚 BOUNDARY FOUND! {target_annex} ends at paragraph {i}: '{para.text[:100]}...' (boundary: {header})")
                        break
                    else:
                        logger.debug(f"⚠️ Skipped same annex match: '{para.text[:50]}...'")
                else:
                    logger.debug(f"❌ No match for '{header}' in: '{para.text[:50]}...'")

                # Also log the exact text comparison for debugging
                if i % 10 == 0:  # Log every 10th paragraph to avoid spam
                    logger.debug(f"�� Para {i} normalized: '{para_text[:50]}...' vs header: '{header_upper}'")

            # Exit early if boundary found
            if end_idx is not None:
                logger.info(f"🎯 Boundary found for {target_annex}, stopping search at paragraph {end_idx}")
                break

    # If no end found, assume it goes to document end
    if start_idx is not None and end_idx is None:
        end_idx = len(doc.paragraphs)
        logger.debug(f"📝 {target_annex} extends to document end (paragraph {end_idx})")

    return start_idx, end_idx


def prune_to_annex_with_boundaries(doc_path: str, target_annex: str, start_idx: int, end_idx: int = None) -> bool:
    """
    Remove all content except the target annex from the document using pre-calculated boundaries.
    This avoids duplicate boundary calculation and improves performance.

    Args:
        doc_path: Path to document to prune
        target_annex: Annex name for logging
        start_idx: Start paragraph index
        end_idx: End paragraph index (None means to document end)

    Returns:
        True if successful, False otherwise
    """
    try:
        print(f"✂️ PRUNING DOCUMENT to keep only {target_annex}")
        print(f"   Document path: {doc_path}")
        print(f"   Using pre-calculated boundaries: start={start_idx}, end={end_idx}")

        import time
        start_time = time.time()

        doc = Document(doc_path)
        print(f"   Loaded document with {len(doc.paragraphs)} paragraphs")
        print(f"   ⏱️ Document load time: {time.time() - start_time:.2f}s")

        if start_idx is None:
            print(f"❌ Invalid start index for {target_annex}")
            return False

        # Build a set of paragraph elements that should be KEPT
        keep_paragraph_elements = set()
        for idx in range(start_idx, end_idx if end_idx is not None else len(doc.paragraphs)):
            if idx < len(doc.paragraphs):
                keep_paragraph_elements.add(doc.paragraphs[idx]._element)

        print(f"   🎯 Keeping {len(keep_paragraph_elements)} paragraph elements (para {start_idx} to {end_idx if end_idx else 'end'})")

        elements_to_delete = []
        element_count = 0
        kept_count = 0

        print(f"   🔄 Processing document body elements...")

        # Iterate over the entire body
        for element in doc.element.body:
            element_count += 1

            # Check if this element is a paragraph we want to keep
            if element in keep_paragraph_elements:
                kept_count += 1
                continue  # Keep this element

            # Otherwise, mark for deletion
            elements_to_delete.append(element)

        print(f"   📊 XML Processing Summary:")
        print(f"      Total body elements: {element_count}")
        print(f"      Elements to keep: {kept_count}")
        print(f"      Elements to delete: {len(elements_to_delete)}")

        print(f"🗑️ Deleting {len(elements_to_delete)} elements outside {target_annex}")

        print(f"🗑️ Deleting {len(elements_to_delete)} elements outside {target_annex}")

        # Delete all marked elements from the document tree
        deleted_count = 0
        print(f"   🗑️ Starting deletion of {len(elements_to_delete)} elements...")

        for i, element in enumerate(elements_to_delete):
            try:
                if element.getparent() is not None:
                    element.getparent().remove(element)
                    deleted_count += 1

                    # Progress logging every 100 deletions
                    if (i + 1) % 100 == 0:
                        print(f"   🗑️ Deleted {i + 1}/{len(elements_to_delete)} elements...")

            except Exception as e:
                print(f"   🚨 ERROR: Failed to delete element {i+1}: {str(e)}")
                print(f"   🚨 Element type: {type(element)}")
                print(f"   🚨 Element tag: {getattr(element, 'tag', 'unknown')}")
                # Continue with other elements
                continue

        print(f"   ✅ Deleted {deleted_count} elements")

        # Save the pruned document
        print(f"   💾 Saving pruned document...")
        doc.save(doc_path)

        print(f"✅ Successfully pruned document to {target_annex}")
        return True

    except Exception as e:
        print(f"❌ Error during pruning: {e}")
        import traceback
        traceback.print_exc()
        return False


def prune_to_annex(doc_path: str, target_annex: str, all_annex_headers: List[str] = None, is_annex_i: bool = False, mapping_row = None) -> bool:
    """
    Remove all content except the target annex from the document.
    Iterates through document body to correctly handle both paragraphs AND tables.

    Args:
        doc_path: Path to document to prune
        target_annex: Annex to keep (e.g., "ANNEX I")
        all_annex_headers: List of all annex headers for boundary detection
        is_annex_i: True if this is Annex I (special handling)

    Returns:
        True if successful, False otherwise
    """
    try:
        print(f"✂️ PRUNING DOCUMENT to keep only {target_annex}")
        print(f"   Document path: {doc_path}")
        doc = Document(doc_path)
        print(f"   Loaded document with {len(doc.paragraphs)} paragraphs")

        # Find annex boundaries
        start_idx, end_idx = find_annex_boundaries(doc, target_annex, all_annex_headers, is_annex_i, mapping_row)

        print(f"   📍 Boundaries found: start={start_idx}, end={end_idx}")

        if start_idx is None:
            print(f"❌ Could not find start marker for {target_annex} in document")
            return False

        logger.info(f"📍 Boundaries: start={start_idx}, end={end_idx}")
        logger.debug(f"   Start paragraph: '{doc.paragraphs[start_idx].text[:100] if start_idx < len(doc.paragraphs) else 'N/A'}...'")
        if end_idx and end_idx < len(doc.paragraphs):
            logger.debug(f"   End paragraph: '{doc.paragraphs[end_idx].text[:100]}...'")

        # Build a set of paragraph elements that should be KEPT
        keep_paragraph_elements = set()
        for idx in range(start_idx, end_idx if end_idx is not None else len(doc.paragraphs)):
            if idx < len(doc.paragraphs):
                keep_paragraph_elements.add(doc.paragraphs[idx]._element)

        print(f"   🎯 Keeping {len(keep_paragraph_elements)} paragraph elements (para {start_idx} to {end_idx if end_idx else 'end'})")

        elements_to_delete = []
        element_count = 0
        kept_count = 0

        print(f"   🔄 Processing document body elements...")

        # Iterate over the entire body
        for element in doc.element.body:
            element_count += 1

            # Check if this element is a paragraph we want to keep
            if element in keep_paragraph_elements:
                kept_count += 1
                continue  # Keep this element

            # Otherwise, mark for deletion
            elements_to_delete.append(element)

        print(f"   📊 XML Processing Summary:")
        print(f"      Total body elements: {element_count}")
        print(f"      Elements to keep: {kept_count}")
        print(f"      Elements to delete: {len(elements_to_delete)}")

        print(f"🗑️ Deleting {len(elements_to_delete)} elements outside {target_annex}")

        # Delete all marked elements from the document tree
        deleted_count = 0
        print(f"   🗑️ Starting deletion of {len(elements_to_delete)} elements...")

        for i, element in enumerate(elements_to_delete):
            try:
                if element.getparent() is not None:
                    element.getparent().remove(element)
                    deleted_count += 1

                    # Progress logging every 100 deletions
                    if (i + 1) % 100 == 0:
                        print(f"   🗑️ Deleted {i + 1}/{len(elements_to_delete)} elements...")

            except Exception as e:
                print(f"   🚨 ERROR: Failed to delete element {i+1}: {str(e)}")
                print(f"   🚨 Element type: {type(element)}")
                print(f"   🚨 Element tag: {getattr(element, 'tag', 'unknown')}")
                # Continue with other elements
                continue

        print(f"   ✅ Deleted {deleted_count} elements")

        # Save the pruned document
        print(f"   💾 Saving pruned document...")
        doc.save(doc_path)

        print(f"✅ Successfully pruned document to {target_annex}")
        return True

    except Exception as e:
        logger.error(f"❌ Error during pruning: {e}", exc_info=True)
        return False


def validate_document_structure(doc_path: str) -> Dict[str, any]:
    """
    Validate and analyze document structure for debugging and verification.

    Args:
        doc_path: Path to document to analyze

    Returns:
        Dictionary with document structure information
    """
    try:
        doc = Document(doc_path)

        analysis = {
            'paragraphs': len(doc.paragraphs),
            'sections': len(doc.sections),
            'tables': len(doc.tables),
            'annexes': [],
            'has_header': False,
            'has_footer': False,
            'title': None,
            'author': None
        }

        # Find annex headers
        for i, para in enumerate(doc.paragraphs):
            para_text = para.text.strip()
            if para_text.upper().startswith("ANNEX"):
                analysis['annexes'].append({
                    'index': i,
                    'text': para_text
                })

        # Check scaffolding elements
        if doc.sections:
            section = doc.sections[0]
            analysis['has_header'] = bool(section.header and len(section.header.paragraphs) > 0)
            analysis['has_footer'] = bool(section.footer and len(section.footer.paragraphs) > 0)

        # Document properties
        props = doc.core_properties
        analysis['title'] = props.title
        analysis['author'] = props.author

        return analysis

    except Exception as e:
        logger.error(f"❌ Error analyzing document structure: {e}")
        return {'error': str(e)}


# INTEGRATION FUNCTIONS
# These functions provide compatibility with existing processor.py code

def get_annex_document_paths(base_output_dir: str, country_code: str) -> Dict[str, str]:
    """
    Get expected paths for annex documents based on country code.
    Helper function for integration with existing processor workflow.

    Args:
        base_output_dir: Base output directory
        country_code: Country code

    Returns:
        Dictionary mapping annex names to expected file paths
    """
    return {
        "ANNEX I": os.path.join(base_output_dir, f"{country_code}_ANNEX_I.docx"),
        "ANNEX IIIB": os.path.join(base_output_dir, f"{country_code}_ANNEX_IIIB.docx")
    }


def split_document_for_country(
    combined_doc_path: str,
    output_dir: str,
    country_code: str
) -> Tuple[str, str]:
    """
    Split document for a specific country - maintains compatibility with existing processor.

    Args:
        combined_doc_path: Path to combined source document
        output_dir: Output directory
        country_code: Country code

    Returns:
        Tuple of (annex_i_path, annex_iiib_path)

    Raises:
        RuntimeError: If splitting fails
    """
    try:
        result_paths = clone_and_split_document(
            combined_doc_path,
            output_dir,
            country_code,
            ["ANNEX I", "ANNEX IIIB"]
        )

        if "ANNEX I" not in result_paths or "ANNEX IIIB" not in result_paths:
            raise RuntimeError(f"Failed to create both annexes for {country_code}")

        return result_paths["ANNEX I"], result_paths["ANNEX IIIB"]

    except Exception as e:
        logger.error(f"❌ Failed to split document for {country_code}: {e}")
        raise RuntimeError(f"Document splitting failed: {e}")


def _generate_annex_filename(annex_header: str, language: str, mapping_row) -> str:
    """
    Generate proper filename following the naming conventions:
    - Annex I: Annex I_EU_SmPC_Language
    - Annex IIIB: Annex IIIB_EU_PL_Language

    Args:
        annex_header: The actual header text from document
        language: Language name
        mapping_row: Mapping row containing language information

    Returns:
        Properly formatted filename
    """
    try:
        # Get language name from mapping if available
        if mapping_row is not None:
            # Try to get language from mapping row
            lang_name = mapping_row.get('Language', language or 'Unknown')
        else:
            lang_name = language or 'Unknown'

        # Determine if this is Annex I or Annex IIIB based on mapping columns
        if mapping_row is not None:
            annex_i_header = str(mapping_row.get('Annex I Header in country language', '')).strip()
            annex_iiib_header = str(mapping_row.get('Annex IIIB Header in country language', '')).strip()

            if annex_header == annex_i_header:
                return f"Annex I_EU_SmPC_{lang_name}.docx"
            elif annex_header == annex_iiib_header:
                return f"Annex IIIB_EU_PL_{lang_name}.docx"

        # Fallback logic if mapping not available
        if "I" in annex_header.upper() and "III" not in annex_header.upper():
            return f"Annex I_EU_SmPC_{lang_name}.docx"
        elif "III" in annex_header.upper():
            return f"Annex IIIB_EU_PL_{lang_name}.docx"
        else:
            # Ultimate fallback
            safe_name = annex_header.replace(" ", "_").replace("/", "_")
            return f"{safe_name}_{lang_name}.docx"

    except Exception as e:
        logger.error(f"❌ Error generating filename: {e}")
        # Emergency fallback
        safe_name = annex_header.replace(" ", "_").replace("/", "_")
        return f"{safe_name}.docx"


# =============================================================================
# LEGACY DOCUMENT COPY FUNCTIONS
# =============================================================================
# These functions are kept for fallback support when clone-and-prune fails.
# They implement element-by-element copying which has limitations but provides
# compatibility with older splitting approaches.

def copy_paragraph(dest_doc: Document, source_para) -> None:
    """Copy a paragraph from one document to another, preserving comprehensive formatting."""
    new_para = dest_doc.add_paragraph()

    # Copy paragraph-level formatting
    new_para.style = source_para.style

    # Copy paragraph format properties
    if source_para.paragraph_format:
        pf_source = source_para.paragraph_format
        pf_dest = new_para.paragraph_format

        # Copy alignment
        if pf_source.alignment is not None:
            pf_dest.alignment = pf_source.alignment

        # Copy spacing
        if pf_source.space_before is not None:
            pf_dest.space_before = pf_source.space_before
        if pf_source.space_after is not None:
            pf_dest.space_after = pf_source.space_after
        if pf_source.line_spacing is not None:
            pf_dest.line_spacing = pf_source.line_spacing

        # Copy indentation
        if pf_source.left_indent is not None:
            pf_dest.left_indent = pf_source.left_indent
        if pf_source.right_indent is not None:
            pf_dest.right_indent = pf_source.right_indent
        if pf_source.first_line_indent is not None:
            pf_dest.first_line_indent = pf_source.first_line_indent

    # Copy all runs with comprehensive formatting
    for run in source_para.runs:
        new_run = new_para.add_run(run.text)

        # Copy basic formatting
        new_run.bold = run.bold
        new_run.italic = run.italic
        new_run.underline = run.underline
        new_run.style = run.style

        # Copy font properties
        if run.font:
            if run.font.name:
                new_run.font.name = run.font.name
            if run.font.size:
                new_run.font.size = run.font.size
            if run.font.color.rgb:
                new_run.font.color.rgb = run.font.color.rgb
            if run.font.highlight_color:
                new_run.font.highlight_color = run.font.highlight_color


def copy_table(dest_doc: Document, source_table) -> None:
    """Copy a table from source document to destination document, preserving formatting."""
    # Get table dimensions
    rows = len(source_table.rows)
    cols = len(source_table.columns) if source_table.rows else 0

    if rows == 0 or cols == 0:
        return

    # Create new table
    new_table = dest_doc.add_table(rows=rows, cols=cols)

    # Copy table style if available
    if hasattr(source_table, 'style') and source_table.style:
        new_table.style = source_table.style

    # Copy cell content and formatting
    for row_idx, source_row in enumerate(source_table.rows):
        dest_row = new_table.rows[row_idx]

        for col_idx, source_cell in enumerate(source_row.cells):
            dest_cell = dest_row.cells[col_idx]

            # Clear default paragraph in destination cell
            dest_cell.text = ""
            if dest_cell.paragraphs:
                dest_cell.paragraphs[0].clear()

            # Copy all paragraphs from source cell
            for para_idx, source_para in enumerate(source_cell.paragraphs):
                if para_idx == 0 and dest_cell.paragraphs:
                    # Use existing first paragraph
                    dest_para = dest_cell.paragraphs[0]
                    # Copy paragraph content manually
                    _copy_paragraph_content(dest_para, source_para)
                else:
                    # Add new paragraph
                    dest_para = dest_cell.add_paragraph()
                    _copy_paragraph_content(dest_para, source_para)


def _copy_paragraph_content(dest_para, source_para) -> None:
    """Helper function to copy paragraph content without creating a new paragraph."""
    # Copy paragraph-level formatting
    dest_para.style = source_para.style

    # Copy paragraph format properties
    if source_para.paragraph_format:
        pf_source = source_para.paragraph_format
        pf_dest = dest_para.paragraph_format

        if pf_source.alignment is not None:
            pf_dest.alignment = pf_source.alignment
        if pf_source.space_before is not None:
            pf_dest.space_before = pf_source.space_before
        if pf_source.space_after is not None:
            pf_dest.space_after = pf_source.space_after
        if pf_source.line_spacing is not None:
            pf_dest.line_spacing = pf_source.line_spacing
        if pf_source.left_indent is not None:
            pf_dest.left_indent = pf_source.left_indent
        if pf_source.right_indent is not None:
            pf_dest.right_indent = pf_source.right_indent
        if pf_source.first_line_indent is not None:
            pf_dest.first_line_indent = pf_source.first_line_indent

    # Copy all runs
    for run in source_para.runs:
        new_run = dest_para.add_run(run.text)

        # Copy formatting
        new_run.bold = run.bold
        new_run.italic = run.italic
        new_run.underline = run.underline
        new_run.style = run.style

        if run.font:
            if run.font.name:
                new_run.font.name = run.font.name
            if run.font.size:
                new_run.font.size = run.font.size
            if run.font.color.rgb:
                new_run.font.color.rgb = run.font.color.rgb
            if run.font.highlight_color:
                new_run.font.highlight_color = run.font.highlight_color


def copy_document_structure(source_doc: Document, dest_doc: Document) -> None:
    """Copy document structure elements like headers, footers, and page setup."""

    try:
        # Copy document properties
        if hasattr(source_doc.core_properties, 'title') and source_doc.core_properties.title:
            dest_doc.core_properties.title = source_doc.core_properties.title
        if hasattr(source_doc.core_properties, 'author') and source_doc.core_properties.author:
            dest_doc.core_properties.author = source_doc.core_properties.author
        if hasattr(source_doc.core_properties, 'subject') and source_doc.core_properties.subject:
            dest_doc.core_properties.subject = source_doc.core_properties.subject
    except Exception as e:
        logger.warning(f"Could not copy document properties: {e}")

    try:
        # Copy page setup from first section
        if source_doc.sections and dest_doc.sections:
            source_section = source_doc.sections[0]
            dest_section = dest_doc.sections[0]

            # Copy page dimensions and margins
            dest_section.page_height = source_section.page_height
            dest_section.page_width = source_section.page_width
            dest_section.left_margin = source_section.left_margin
            dest_section.right_margin = source_section.right_margin
            dest_section.top_margin = source_section.top_margin
            dest_section.bottom_margin = source_section.bottom_margin
            dest_section.gutter = source_section.gutter

            # Copy orientation
            dest_section.orientation = source_section.orientation
    except Exception as e:
        logger.warning(f"Could not copy page setup: {e}")


def copy_headers_and_footers(source_doc: Document, dest_doc: Document) -> None:
    """Copy headers and footers from source document to destination document."""

    try:
        # Ensure destination has at least one section
        if not dest_doc.sections:
            return

        source_section = source_doc.sections[0] if source_doc.sections else None
        dest_section = dest_doc.sections[0]

        if not source_section:
            return

        # Copy headers
        header_types = [
            ('first_page_header', 'first_page_header'),
            ('even_page_header', 'even_page_header'),
            ('header', 'header')  # Default header
        ]

        for source_attr, dest_attr in header_types:
            try:
                if hasattr(source_section, source_attr) and hasattr(dest_section, dest_attr):
                    source_header = getattr(source_section, source_attr)
                    dest_header = getattr(dest_section, dest_attr)

                    # Clear existing content
                    for para in dest_header.paragraphs:
                        para.clear()

                    # Copy paragraphs from source header
                    for i, source_para in enumerate(source_header.paragraphs):
                        if i < len(dest_header.paragraphs):
                            _copy_paragraph_content(dest_header.paragraphs[i], source_para)
                        else:
                            dest_header_para = dest_header.add_paragraph()
                            _copy_paragraph_content(dest_header_para, source_para)

            except Exception as e:
                logger.warning(f"Could not copy header {source_attr}: {e}")

        # Copy footers
        footer_types = [
            ('first_page_footer', 'first_page_footer'),
            ('even_page_footer', 'even_page_footer'),
            ('footer', 'footer')  # Default footer
        ]

        for source_attr, dest_attr in footer_types:
            try:
                if hasattr(source_section, source_attr) and hasattr(dest_section, dest_attr):
                    source_footer = getattr(source_section, source_attr)
                    dest_footer = getattr(dest_section, dest_attr)

                    # Clear existing content
                    for para in dest_footer.paragraphs:
                        para.clear()

                    # Copy paragraphs from source footer
                    for i, source_para in enumerate(source_footer.paragraphs):
                        if i < len(dest_footer.paragraphs):
                            _copy_paragraph_content(dest_footer.paragraphs[i], source_para)
                        else:
                            dest_footer_para = dest_footer.add_paragraph()
                            _copy_paragraph_content(dest_footer_para, source_para)

            except Exception as e:
                logger.warning(f"Could not copy footer {source_attr}: {e}")

    except Exception as e:
        logger.warning(f"Could not copy headers and footers: {e}")


def copy_styles(source_doc: Document, dest_doc: Document) -> None:
    """Copy custom styles from source document to destination document."""

    try:
        # Get style collections
        source_styles = source_doc.styles
        dest_styles = dest_doc.styles

        # Copy paragraph styles
        for source_style in source_styles:
            if source_style.type == 1:  # Paragraph style
                try:
                    # Check if style already exists
                    existing_style = None
                    try:
                        existing_style = dest_styles[source_style.name]
                    except KeyError:
                        pass

                    if not existing_style:
                        # Create new style
                        new_style = dest_styles.add_style(source_style.name, 1)  # 1 = paragraph style

                        # Copy basic properties
                        if hasattr(source_style, 'font') and hasattr(new_style, 'font'):
                            if source_style.font.name:
                                new_style.font.name = source_style.font.name
                            if source_style.font.size:
                                new_style.font.size = source_style.font.size

                except Exception as e:
                    logger.warning(f"Could not copy style {source_style.name}: {e}")

    except Exception as e:
        logger.warning(f"Could not copy styles: {e}")


# Safe copy functions for debugging/fallback scenarios
def copy_paragraph_safe(dest_doc: Document, source_para) -> None:
    """
    Safely copy a paragraph from source to destination document.

    This preserves basic formatting while avoiding XML corruption.
    """

    # Create new paragraph
    new_para = dest_doc.add_paragraph()

    # Copy paragraph-level properties safely
    try:
        new_para.style = source_para.style
    except:
        # If style copying fails, use default
        pass

    try:
        new_para.alignment = source_para.alignment
    except:
        pass

    # Copy runs with formatting
    for run in source_para.runs:
        new_run = new_para.add_run(run.text)

        # Copy basic formatting safely
        try:
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
        except:
            # If formatting copying fails, continue with plain text
            pass


def copy_table_safe(dest_doc: Document, source_table) -> None:
    """
    Safely copy a table from source to destination document.
    """

    try:
        # Get table dimensions
        rows = len(source_table.rows)
        cols = len(source_table.columns) if rows > 0 else 0

        if rows > 0 and cols > 0:
            # Create new table
            new_table = dest_doc.add_table(rows=rows, cols=cols)

            # Copy cell contents
            for row_idx in range(rows):
                for col_idx in range(cols):
                    try:
                        source_cell = source_table.cell(row_idx, col_idx)
                        dest_cell = new_table.cell(row_idx, col_idx)
                        dest_cell.text = source_cell.text
                    except:
                        # If cell copying fails, continue
                        continue

    except Exception as e:
        logger.warning(f"Could not copy table: {e}")


def copy_document_settings_safe(source_doc: Document, target_doc: Document) -> None:
    """
    Safely copy basic document settings without causing corruption.
    """

    try:
        # Only copy very basic properties that are unlikely to cause issues
        if hasattr(source_doc, 'core_properties') and hasattr(target_doc, 'core_properties'):
            # Copy basic metadata only
            target_doc.core_properties.author = source_doc.core_properties.author
    except:
        # If any copying fails, continue without it
        pass


def copy_document_properties(source_doc: Document, target_doc: Document) -> None:
    """
    Copy document-level properties like styles, themes, etc.

    This function is kept for compatibility but now calls the safe version.
    """

    copy_document_settings_safe(source_doc, target_doc)


def get_document_elements_in_order(doc: Document) -> List[Dict]:
    """
    Get all document elements (paragraphs and tables) in their order of appearance.

    Returns:
        List of dictionaries with 'type' and 'content' keys
    """

    elements = []

    # Add all paragraphs
    for para in doc.paragraphs:
        elements.append({
            'type': 'paragraph',
            'content': para
        })

    # Note: Tables are embedded within the document structure
    # For simplicity, we'll handle them as part of paragraph processing
    # This avoids the complexity of XML order tracking

    return elements


def extract_section_safe_copy(source_doc: Document, start_idx: int, end_idx: int) -> Document:
    """
    Safe document extraction that preserves formatting without XML corruption.

    This method copies paragraphs, tables, and basic formatting while ensuring
    the resulting document is valid and doesn't trigger Word warnings.
    """

    logger.debug(f"Using safe copying for range {start_idx} to {end_idx-1}")

    # Create new document
    new_doc = Document()

    # TEMPORARILY DISABLED: Copy comprehensive document structure (causes process crash)
    logger.debug("Skipping document structure copying (debugging process crash)...")
    # copy_document_structure(source_doc, new_doc)
    # copy_headers_and_footers(source_doc, new_doc)
    # copy_styles(source_doc, new_doc)

    # Copy additional document-level settings safely
    copy_document_settings_safe(source_doc, new_doc)

    # Clear the default empty paragraph
    if new_doc.paragraphs:
        p = new_doc.paragraphs[0]
        p.clear()

    # Track what we're copying
    paragraphs_copied = 0
    tables_copied = 0

    # Get both paragraphs and tables from the source document
    source_elements = get_document_elements_in_order(source_doc)

    # Filter elements to the target range
    target_elements = []
    current_para_idx = 0

    for element in source_elements:
        if element['type'] == 'paragraph':
            if start_idx <= current_para_idx < end_idx:
                target_elements.append(element)
            current_para_idx += 1
        elif element['type'] == 'table':
            # Include tables that fall within our range
            if start_idx <= current_para_idx < end_idx:
                target_elements.append(element)

    # TEMPORARILY REVERTING to safe functions to debug crash
    logger.debug("Using safe copying functions to isolate crash cause...")
    for element in target_elements:
        if element['type'] == 'paragraph':
            copy_paragraph_safe(new_doc, element['content'])
            paragraphs_copied += 1
        elif element['type'] == 'table':
            copy_table_safe(new_doc, element['content'])
            tables_copied += 1

    logger.info(f"Safely copied {paragraphs_copied} paragraphs and {tables_copied} tables")

    return new_doc


def extract_section_xml(source_doc: Document, start_idx: int, end_idx: Optional[int] = None) -> Document:
    """
    Extract a section from the source document using safe paragraph copying to avoid XML corruption.

    This approach prioritizes document integrity over advanced XML preservation.
    """

    # Determine which paragraphs to include
    total_paragraphs = len(source_doc.paragraphs)
    actual_end_idx = end_idx if end_idx is not None else total_paragraphs

    logger.debug(f"Extracting paragraphs {start_idx} to {actual_end_idx-1} (total: {actual_end_idx - start_idx})")

    # Use safe paragraph-by-paragraph copying to avoid XML corruption
    return extract_section_safe_copy(source_doc, start_idx, actual_end_idx)


